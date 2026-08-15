from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from korgan_legal_ai.blueprints.models import DocumentBlueprint
from korgan_legal_ai.blueprints.registry import CLAIM_DEBT_RECOVERY

# Verbatim norm quotes and their attribution are set in italics, and the citation list in the legal
# grounds follows the same treatment. This is the KORGAN house style for quoted law.
_QUOTE_LINE = re.compile(r'^[«"]')
_QUOTE_ATTRIBUTION = re.compile(r"^\(.*стать[яеи].*\)$", re.IGNORECASE)
_CITATION_LINE = re.compile(r"^-\s+.*,\s*статья\s", re.IGNORECASE)
_PENALTY_PERIOD_LINE = re.compile(r"^(за период с|применено договорное ограничение)", re.IGNORECASE)
_REQUISITE_LINE = re.compile(r"^(БИН/ИИН|Адрес|Кому)\s*:")


def _set_run_font(run, *, size: int = 12, bold: bool | None = None, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.italic = italic


def build_document_docx(
    document_text: str,
    *,
    blueprint: DocumentBlueprint | None = None,
    title: str | None = None,
) -> bytes:
    """Build a clean Word file from the already QA-approved document text.

    This function is presentation-only: it does not calculate, rewrite or add legal content.
    NEEDS_VERIFICATION markers are preserved verbatim and emphasized rather than hidden. Structure
    is recognised from the blueprint of the document that was actually drafted, so a claim, a
    pretrial demand and an appeal each get their own headings without the exporter knowing anything
    about the legal content.
    """
    spec = blueprint or CLAIM_DEBT_RECOVERY
    headings = set(spec.headings())
    document_title = spec.title
    party_labels = {spec.parties.author_label, spec.parties.opponent_label}

    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0

    props = doc.core_properties
    props.title = title or spec.title.capitalize()
    props.subject = "KORGAN Legal AI lawyer-review draft"
    props.author = ""
    props.last_modified_by = ""
    props.comments = "Generated from QA-approved legal draft text."

    lines = document_text.replace("\r\n", "\n").split("\n")
    after_title = False
    for raw in lines:
        line = raw.rstrip()
        if not line:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            after_title = False
            continue

        stripped = line.strip()
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.keep_together = False
        paragraph.paragraph_format.keep_with_next = False

        if stripped in headings:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run(stripped)
            _set_run_font(run, size=14 if stripped == document_title else 12, bold=True)
            after_title = stripped == document_title
            continue

        # The subject line directly under the document title stays centred with it.
        if after_title:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(stripped)
            _set_run_font(run, size=12, bold=False)
            after_title = False
            continue

        if "NEEDS_VERIFICATION" in stripped:
            run = paragraph.add_run(stripped)
            _set_run_font(run, bold=True)
            continue

        if _QUOTE_LINE.match(stripped) or _QUOTE_ATTRIBUTION.match(stripped):
            paragraph.paragraph_format.left_indent = Cm(1)
            run = paragraph.add_run(stripped)
            _set_run_font(run, italic=True)
            continue

        if _CITATION_LINE.match(stripped):
            run = paragraph.add_run(stripped)
            _set_run_font(run, italic=True)
            continue

        if stripped.startswith("ЦЕНА ИСКА:"):
            run = paragraph.add_run(stripped)
            _set_run_font(run, bold=True)
            continue

        label, separator, value = stripped.partition(":")
        if separator and label.strip() in party_labels:
            # Party requisites carry the reader's eye through the header block, so the side and its
            # name are bold while the rest of the block stays regular weight.
            _set_run_font(paragraph.add_run(f"{label}:"), bold=True)
            _set_run_font(paragraph.add_run(value), bold=True)
            continue

        if _REQUISITE_LINE.match(stripped):
            run = paragraph.add_run(stripped)
            _set_run_font(run)
            continue

        if _PENALTY_PERIOD_LINE.match(stripped):
            paragraph.paragraph_format.left_indent = Cm(1)
            run = paragraph.add_run(stripped)
            _set_run_font(run)
            continue

        run = paragraph.add_run(stripped)
        _set_run_font(run)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def build_claim_docx(document_text: str, *, title: str = "Исковое заявление") -> bytes:
    """Backwards-compatible entry point for the debt-recovery claim."""
    return build_document_docx(
        document_text,
        blueprint=CLAIM_DEBT_RECOVERY,
        title=title,
    )

from __future__ import annotations

import base64
from dataclasses import dataclass
from io import BytesIO
from typing import Literal, cast

from docx import Document
from openai import OpenAI
from pydantic import BaseModel, Field
from pypdf import PdfReader


SUPPORTED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "image/jpeg",
    "image/png",
}


class DocumentExtractionError(RuntimeError):
    """A document cannot be safely converted into evidence text."""


class VisualPageTranscription(BaseModel):
    page_number: int = Field(ge=1)
    text: str
    uncertain_fragments: list[str] = Field(default_factory=list)


class VisualTranscription(BaseModel):
    pages: list[VisualPageTranscription]


@dataclass(frozen=True)
class ExtractedDocument:
    source_id: str
    mime_type: str
    method: Literal["docx_text", "pdf_text", "vision"]
    text: str
    uncertain_fragments: tuple[str, ...] = ()

    def as_case_context(self) -> str:
        """Serialize evidence text for Fact Lock without exposing the original filename.

        Text-layer DOCX/PDF content is deterministic extraction. Visual-only content is marked as a
        transcription that requires confirmation for material legal facts; this prevents a confident
        OCR/model error from silently becoming a locked party, amount, date or contract number.
        """
        lines = [f"[DOCUMENT {self.source_id}]"]
        if self.method == "vision":
            lines.append("[VISUAL_TRANSCRIPTION_REQUIRES_CONFIRMATION]")
        lines.append(self.text.strip())
        if self.uncertain_fragments:
            lines.append("[OCR_UNCERTAIN]")
            lines.extend(f"- {item}" for item in self.uncertain_fragments if item.strip())
        return "\n".join(line for line in lines if line)


_TRANSCRIPTION_SYSTEM = """You are a legal-document transcription module.
Your only task is to read the supplied scan/image/PDF and transcribe what is visibly present.
Do NOT summarize, interpret, infer, repair, calculate or complete missing text.
Preserve names, identifiers, dates, amounts, percentages and contract numbers exactly as visible.
If any fragment is blurred, cut off or genuinely uncertain, copy the best visible fragment into
uncertain_fragments and do not silently guess it. For a single image use page_number=1.
Return only the supplied structured schema."""


class OpenAIDocumentExtractor:
    """Extract source text while keeping legal interpretation in Fact Lock.

    DOCX and text PDFs are parsed deterministically in-process. OpenAI vision is used only when a
    source has no reliable embedded text (images and scan-only PDFs). Model output is transcription,
    never a legal conclusion. Responses are sent with store=False. Material facts from visual-only
    transcription remain confirmation-gated by Fact Lock.
    """

    def __init__(
        self,
        *,
        api_key: str,
        model: str = "gpt-5-mini",
        max_text_chars: int = 50000,
        client: OpenAI | None = None,
    ) -> None:
        if not api_key.strip() and client is None:
            raise ValueError("OpenAI API key is required for visual document extraction")
        if max_text_chars <= 0:
            raise ValueError("max_text_chars must be positive")
        self.client = client or OpenAI(api_key=api_key)
        self.model = model
        self.max_text_chars = max_text_chars

    @staticmethod
    def _clean_text(value: str) -> str:
        lines = [line.rstrip() for line in value.replace("\x00", "").splitlines()]
        cleaned: list[str] = []
        blank = False
        for line in lines:
            if line.strip():
                cleaned.append(line.strip())
                blank = False
            elif not blank and cleaned:
                cleaned.append("")
                blank = True
        return "\n".join(cleaned).strip()

    def _bounded(self, value: str) -> str:
        value = self._clean_text(value)
        if not value:
            raise DocumentExtractionError("document contains no readable text")
        if len(value) > self.max_text_chars:
            raise DocumentExtractionError("extracted document text exceeds configured limit")
        return value

    def _extract_docx(self, data: bytes) -> str:
        try:
            document = Document(BytesIO(data))
        except Exception as exc:
            raise DocumentExtractionError("DOCX cannot be parsed") from exc

        blocks: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                blocks.append(text)
        for table_index, table in enumerate(document.tables, start=1):
            blocks.append(f"[TABLE {table_index}]")
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    blocks.append(" | ".join(cells))
        return self._bounded("\n".join(blocks))

    def _extract_pdf_text(self, data: bytes) -> str | None:
        try:
            reader = PdfReader(BytesIO(data))
        except Exception as exc:
            raise DocumentExtractionError("PDF cannot be parsed") from exc

        pages: list[str] = []
        visible_chars = 0
        for page_number, page in enumerate(reader.pages, start=1):
            try:
                text = self._clean_text(page.extract_text() or "")
            except Exception:
                text = ""
            if text:
                visible_chars += len(text)
                pages.append(f"[PAGE {page_number}]\n{text}")

        # A tiny extraction is commonly a scan with a stray metadata/footer string. Treat it as
        # scan-only rather than pretending we have captured the evidentiary content.
        if visible_chars < 120:
            return None
        return self._bounded("\n\n".join(pages))

    @staticmethod
    def _parsed(response) -> VisualTranscription:
        parsed = getattr(response, "output_parsed", None)
        if parsed is not None:
            return cast(VisualTranscription, parsed)
        for item in getattr(response, "output", []):
            if getattr(item, "type", None) != "message":
                continue
            for content in getattr(item, "content", []):
                parsed = getattr(content, "parsed", None)
                if parsed is not None:
                    return cast(VisualTranscription, parsed)
        raise DocumentExtractionError("vision model returned no structured transcription")

    def _vision(self, data: bytes, mime_type: str) -> tuple[str, tuple[str, ...]]:
        encoded = base64.b64encode(data).decode("ascii")
        if mime_type == "application/pdf":
            source_item = {
                "type": "input_file",
                "filename": "evidence.pdf",
                "file_data": encoded,
            }
        else:
            source_item = {
                "type": "input_image",
                "image_url": f"data:{mime_type};base64,{encoded}",
                "detail": "high",
            }

        try:
            response = self.client.responses.parse(
                model=self.model,
                store=False,
                input=[
                    {"role": "system", "content": _TRANSCRIPTION_SYSTEM},
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "input_text",
                                "text": "Transcribe this evidence document exactly as visible.",
                            },
                            source_item,
                        ],
                    },
                ],
                text_format=VisualTranscription,
            )
        except Exception as exc:
            raise DocumentExtractionError("visual document transcription failed") from exc

        transcription = self._parsed(response)
        page_blocks: list[str] = []
        uncertain: list[str] = []
        for page in transcription.pages:
            page_text = self._clean_text(page.text)
            if page_text:
                page_blocks.append(f"[PAGE {page.page_number}]\n{page_text}")
            uncertain.extend(item.strip() for item in page.uncertain_fragments if item.strip())
        return self._bounded("\n\n".join(page_blocks)), tuple(dict.fromkeys(uncertain))

    def extract(self, *, source_id: str, data: bytes, mime_type: str) -> ExtractedDocument:
        if mime_type not in SUPPORTED_MIME_TYPES:
            raise DocumentExtractionError("unsupported document type")
        if not data:
            raise DocumentExtractionError("empty document")

        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return ExtractedDocument(
                source_id=source_id,
                mime_type=mime_type,
                method="docx_text",
                text=self._extract_docx(data),
            )

        if mime_type == "application/pdf":
            text = self._extract_pdf_text(data)
            if text is not None:
                return ExtractedDocument(
                    source_id=source_id,
                    mime_type=mime_type,
                    method="pdf_text",
                    text=text,
                )
            visual_text, uncertain = self._vision(data, mime_type)
            return ExtractedDocument(
                source_id=source_id,
                mime_type=mime_type,
                method="vision",
                text=visual_text,
                uncertain_fragments=uncertain,
            )

        visual_text, uncertain = self._vision(data, mime_type)
        return ExtractedDocument(
            source_id=source_id,
            mime_type=mime_type,
            method="vision",
            text=visual_text,
            uncertain_fragments=uncertain,
        )

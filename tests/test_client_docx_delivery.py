from __future__ import annotations

import base64
import io

import pytest
from docx import Document
from docx.oxml.ns import qn

from korgan.client_docx import clean_client_docx


def _fixture() -> bytes:
    doc = Document()
    doc.add_paragraph("KORGAN QA STATUS: LAWYER-REVIEW DRAFT")
    doc.add_paragraph().add_run("Взыскать 150 000 тенге по статье 272 ГК РК.").bold = True
    doc.add_paragraph("[ТРЕБУЕТ УТОЧНЕНИЯ: адрес ответчика]")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "verification_notes: source-bound"
    table.cell(0, 1).text = "Неустойка: 3 000 тенге"
    doc.sections[0].header.paragraphs[0].text = "PRELIMINARY DRAFT"
    doc.sections[0].footer.paragraphs[0].text = (
        "Проект сформирован KORGAN Legal AI на основании материалов пользователя. "
        "Перед подачей необходимо проверить реквизиты."
    )
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    return "\n".join(
        node.text or ""
        for part in doc.part.package.parts
        if getattr(part, "element", None) is not None
        for node in part.element.iter(qn("w:t"))
    )


def test_delivery_strips_internal_stories_but_preserves_legal_content_and_formatting():
    original = _fixture()
    cleaned = clean_client_docx(original)
    text = _text(cleaned)
    for marker in ("QA STATUS", "PRELIMINARY", "verification_notes", "source-bound", "KORGAN Legal AI"):
        assert marker not in text
    assert "150 000 тенге по статье 272 ГК РК" in text
    assert "[ТРЕБУЕТ УТОЧНЕНИЯ: адрес ответчика]" in text
    assert "Неустойка: 3 000 тенге" in text
    doc = Document(io.BytesIO(cleaned))
    legal_paragraph = next(p for p in doc.paragraphs if "150 000" in p.text)
    assert legal_paragraph.runs[0].bold is True
    assert len(doc.tables[0].rows[0].cells) == 2
    assert "QA STATUS" in _text(original), "внутренний оригинал нельзя перезаписывать"
    assert clean_client_docx(cleaned) == cleaned


@pytest.mark.parametrize("document_type", ["claim", "contract", "response", "pretrial", "pretrial_response"])
def test_existing_downloads_are_cleaned_for_every_document_type(document_type):
    from korgan.miniapp_document_access import _decode_document

    case = {"document_type": document_type, "document_base64": base64.b64encode(_fixture()).decode()}
    assert "QA STATUS" not in _text(_decode_document(case))
    assert "QA STATUS" in _text(base64.b64decode(case["document_base64"]))

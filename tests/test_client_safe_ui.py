from __future__ import annotations

import io
from pathlib import Path

from docx import Document

from korgan.client_safe_ui import _lookup_local, clean_client_docx, sanitize_client_text
from korgan.legal.corpus import ACT_GK_SPECIAL, KNOWN_ACTS, LegalCorpus


def test_verification_dialogue_is_never_client_facing() -> None:
    raw = (
        "Не понял, что сделать с замечаниями. Можно ответить: "
        "пометь статью 715 ГК РК как NEEDS_VERIFICATION и продолжи. "
        "Что именно нужно сделать?"
    )
    result = sanitize_client_text(raw)
    assert result is not None
    assert "715" not in result
    assert "NEEDS_VERIFICATION" not in result
    assert "замечани" not in result.lower()
    assert "что именно нужно сделать" not in result.lower()


def test_document_caption_hides_internal_status_and_checklist() -> None:
    raw = (
        "⚠️ NEEDS_VERIFICATION\nГотовый проект иска — файл Word (.docx).\n\n"
        "Перед подачей проверьте:\n• статья 722 ГК РК: текста нормы нет в корпусе KORGAN"
    )
    result = sanitize_client_text(raw)
    assert result is not None
    assert "NEEDS_VERIFICATION" not in result
    assert "722" not in result
    assert "корпус" not in result.lower()
    assert "Проект иска сформирован" in result


def test_help_copy_cannot_leak_machine_label() -> None:
    result = sanitize_client_text("Неподтверждённые данные отмечаются как NEEDS_VERIFICATION.")
    assert result is not None
    assert "NEEDS_VERIFICATION" not in result
    assert result == "Неподтверждённые данные отмечаются как «требует проверки»."


def test_the_substituted_label_reads_as_russian_wherever_it_lands() -> None:
    """Метка встаёт в чужую фразу, склонять её не по чему — значит, кавычки.

    Подстановка именной группой давала «статус дополнительной проверкой
    системы»: падеж от исходного «как NEEDS_VERIFICATION» не наследуется.
    Клиент читает это как ошибку продукта, а не как скрытый служебный текст.
    """
    result = sanitize_client_text("Документ получил статус NEEDS_VERIFICATION.")
    assert result == "Документ получил статус «требует проверки»."


def test_internal_stage_label_is_cut_out_whole_not_by_its_prefix() -> None:
    """Снятие одного префикса оставляло «: LAWYER-REVIEW DRAFT» с двоеточием."""
    result = sanitize_client_text("KORGAN QA STATUS: LAWYER-REVIEW DRAFT")
    assert result is not None
    assert "LAWYER-REVIEW" not in result
    assert "KORGAN QA STATUS" not in result
    assert ":" not in result


def test_clean_docx_removes_internal_qa_status() -> None:
    document = Document()
    document.add_paragraph("KORGAN QA STATUS: LAWYER-REVIEW DRAFT")
    document.add_paragraph("ИСКОВОЕ ЗАЯВЛЕНИЕ")
    source = io.BytesIO()
    document.save(source)

    cleaned = clean_client_docx(source.getvalue())
    reopened = Document(io.BytesIO(cleaned))
    text = "\n".join(paragraph.text for paragraph in reopened.paragraphs)
    assert "KORGAN QA STATUS" not in text
    assert "LAWYER-REVIEW DRAFT" not in text
    assert "ИСКОВОЕ ЗАЯВЛЕНИЕ" in text


def test_legacy_citation_audit_can_read_verified_sqlite_corpus(tmp_path: Path) -> None:
    db = tmp_path / "corpus.sqlite3"
    adilet_id, title = KNOWN_ACTS[ACT_GK_SPECIAL]
    with LegalCorpus(db) as corpus:
        corpus.upsert_act(
            ACT_GK_SPECIAL,
            adilet_id,
            title,
            "https://adilet.zan.kz/rus/docs/K990000409_",
            "2026-08-17",
            "2026-08-17",
        )
        corpus.upsert_provision(
            act_id=ACT_GK_SPECIAL,
            article_no="715",
            item_no=None,
            heading="Договор займа",
            body=(
                "По договору займа одна сторона передает другой стороне деньги, "
                "а заемщик обязуется своевременно возвратить полученную сумму денег."
            ),
            edition_date="2026-08-17",
            url="https://adilet.zan.kz/rus/docs/K990000409_#z715",
            sort_key=1,
        )

    record = _lookup_local("ГК РК", "715", db_path=db)
    assert record is not None
    assert record.level == "VERIFIED"
    assert record.article == "715"
    assert "заемщик" in record.text.lower()
    assert "adilet.zan.kz" in record.source_url

from __future__ import annotations

import io

from aiogram.types import BufferedInputFile
from docx import Document

from korgan.claim_intent import is_claim_drafting_request
from korgan.contract_intent import is_contract_drafting_request
from korgan.i18n import KK, button
from korgan.language_context import _CURRENT_LANGUAGE
from korgan.legal_safety import privacy_text, terms_text
from korgan.localized_transport import _localize_docx
from korgan.response_intent import is_response_to_claim_request
from korgan.ui import language_menu, main_menu


def _reply_labels() -> list[str]:
    return [item.text for row in main_menu(KK).keyboard for item in row]


def test_kazakh_main_menu_is_fully_kazakh() -> None:
    labels = _reply_labels()
    for key in ("consultation", "document", "prices", "case", "lawyer", "help", "support", "feedback", "language", "delete"):
        assert button(KK, key) in labels
    assert "⚖️ Консультация" not in labels


def test_language_selector_has_both_languages() -> None:
    markup = language_menu()
    labels = [item.text for row in markup.inline_keyboard for item in row]
    callbacks = [item.callback_data for row in markup.inline_keyboard for item in row]
    assert labels == ["🇰🇿 Қазақша", "🇷🇺 Русский"]
    assert callbacks == ["lang:kk", "lang:ru"]


def test_kazakh_legal_intents_use_existing_generators() -> None:
    assert is_claim_drafting_request("Қарызды өндіріп алу туралы талап қою арызын дайында")
    assert not is_claim_drafting_request("Талап қою арызын қалай дайындауға болады?")
    assert is_contract_drafting_request("Қызмет көрсету шартын дайында")
    assert is_response_to_claim_request("Талап қою арызына пікір дайында")


def test_kazakh_terms_and_privacy_are_localized() -> None:
    assert "пайдалану шарттары" in terms_text(KK)
    assert "Дербес деректер" in privacy_text(KK)


def test_kazakh_docx_transport_localizes_fixed_court_labels() -> None:
    doc = Document()
    doc.add_paragraph("В суд: Районный суд")
    doc.add_paragraph("Истец: Иванов")
    doc.add_paragraph("Ответчик: Петров")
    doc.add_paragraph("Цена иска: 100 000 тенге")
    doc.add_paragraph("Госпошлина: 1 000 тенге")
    doc.add_paragraph("Правовое обоснование")
    doc.add_paragraph("На основании изложенного ПРОШУ СУД:")
    doc.add_paragraph("Приложения:")
    raw = io.BytesIO()
    doc.save(raw)

    token = _CURRENT_LANGUAGE.set(KK)
    try:
        localized = _localize_docx(BufferedInputFile(raw.getvalue(), filename="claim.docx"))
    finally:
        _CURRENT_LANGUAGE.reset(token)

    rendered = Document(io.BytesIO(localized.data))
    text = "\n".join(p.text for p in rendered.paragraphs)
    assert "Сотқа:" in text
    assert "Талап қоюшы:" in text
    assert "Жауапкер:" in text
    assert "Талап қою бағасы:" in text
    assert "Мемлекеттік баж:" in text
    assert "Құқықтық негіздеме" in text
    assert "СОТТАН СҰРАЙМЫН" in text
    assert "Қосымшалар:" in text

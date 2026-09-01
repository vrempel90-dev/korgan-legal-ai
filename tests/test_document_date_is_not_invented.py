"""Дату документа ставит тот, кто его подписывает, а не генератор.

Иск, претензия и ответ на претензию выходят из KORGAN с пустой строкой подписи:
документ ещё не подписан. Но рядом печаталась конкретная дата — день, когда
проект был сформирован. Это утверждение о факте, которого не было: документ в
этот день никто не подписал и никуда не направил.

Для досудебной претензии цена ошибки процессуальная. Дата претензии — точка
отсчёта срока ответа и доказательство соблюдения досудебного порядка. Если
клиент распечатает и отправит её через неделю, дата на документе разойдётся с
почтовой квитанцией, и оппонент получит готовый довод. Сам KORGAN это уже
знает: ``flag_same_day_pretrial_risk`` предупреждает, когда претензия в
материалах датирована днём подготовки иска, — и одновременно генератор
проставлял ровно такую дату на претензии, которую сам же и составил.

Для договора дата заключения — существенный реквизит, а не оформление. Когда
модель не вернула место и дату, подставлялся сегодняшний день, причём место
честно помечалось ``[ТРЕБУЕТ УТОЧНЕНИЯ]``, а дата молча выдумывалась.

Отзыв на иск уже сделан правильно — он оставляет дату подписанту. Здесь это
закрепляется как обязательное поведение всех четырёх остальных экспортов.
"""

from __future__ import annotations

import io
from datetime import datetime
from zoneinfo import ZoneInfo

import pytest
from docx import Document

from korgan.claim_docx import build_claim_docx
from korgan.contract_docx import build_contract_docx
from korgan.i18n import KK
from korgan.language_context import _CURRENT_LANGUAGE
from korgan.legal_types import (
    ClaimDraft,
    ContractDraft,
    ContractSection,
    VerificationStatus,
)
from korgan.pretrial import PretrialDraft, build_pretrial_docx
from korgan.pretrial_response import PretrialResponseDraft, build_pretrial_response_docx
from korgan.response_docx import build_response_to_claim_docx
from korgan.response_types import ResponseToClaimDraft

TODAY_KZ = datetime.now(ZoneInfo("Asia/Almaty")).strftime("%d.%m.%Y")

BLANK = "____________________"


def _text(blob: bytes) -> str:
    document = Document(io.BytesIO(blob))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _claim() -> ClaimDraft:
    return ClaimDraft(
        status=VerificationStatus.VERIFIED,
        title="ИСКОВОЕ ЗАЯВЛЕНИЕ о взыскании задолженности",
        court="Медеуский районный суд города Алматы",
        claimant=["Ахметов Руслан Маратович, ИИН 900101300123, г. Алматы, ул. Абая, 150"],
        defendant=["ТОО «Компания», БИН 210987654321, г. Алматы, ул. Розыбакиева, 10"],
        price_of_claim="2 300 000 тенге",
        state_duty="23 000 тенге",
        facts=["Между сторонами заключён договор № 12 от 15.01.2026."],
        legal_basis=["Обязательство должно исполняться надлежащим образом."],
        requests=["Взыскать с ответчика 2 300 000 тенге основного долга."],
        attachments=["Копия договора № 12 от 15.01.2026"],
        verification_notes=[],
        source_urls=["https://adilet.zan.kz/rus/docs/K940001000_"],
    )


def _pretrial() -> PretrialDraft:
    return PretrialDraft(
        status=VerificationStatus.VERIFIED,
        title="ДОСУДЕБНАЯ ПРЕТЕНЗИЯ",
        sender=["ТОО «Астана Строй», БИН 123456789012, г. Астана, ул. Кенесары, 40"],
        recipient=["ТОО «Заказчик», БИН 210987654321, г. Астана, ул. Абая, 15"],
        facts=["Между сторонами заключён договор подряда № 12 от 15.01.2026."],
        legal_basis=[],
        demands=["Оплатить задолженность в размере 2 300 000 тенге."],
        deadline="10 календарных дней с даты получения настоящей претензии",
        consequences=["При неисполнении требований спор будет передан на разрешение суда."],
        attachments=["Копия договора подряда № 12 от 15.01.2026"],
        verification_notes=[],
        source_urls=[],
        calculation=["Основной долг: 2 300 000 тенге."],
    )


def _pretrial_response() -> PretrialResponseDraft:
    return PretrialResponseDraft(
        status=VerificationStatus.VERIFIED,
        title="ОТВЕТ НА ПРЕТЕНЗИЮ",
        sender=["ТОО «Заказчик», БИН 210987654321, г. Астана, ул. Абая, 15"],
        recipient=["ТОО «Астана Строй», БИН 123456789012, г. Астана, ул. Кенесары, 40"],
        reference="претензия от 05.03.2026 № 7",
        claim_summary=["Заявлено требование об оплате 2 300 000 тенге."],
        disputed_circumstances=["Оспаривается объём принятых работ по акту от 20.02.2026."],
        position=["Требование признаётся в части 1 400 000 тенге."],
        objections=["Работы на сумму 900 000 тенге не приняты."],
        calculation_review=["Неустойка исчислена до наступления срока оплаты."],
        legal_basis=[],
        response_terms=["Готовы оплатить признанную часть 1 400 000 тенге."],
        attachments=["Копия акта от 20.02.2026 с замечаниями"],
        verification_notes=[],
        source_urls=[],
    )


def _response_to_claim() -> ResponseToClaimDraft:
    return ResponseToClaimDraft(
        status=VerificationStatus.VERIFIED,
        title="ОТЗЫВ НА ИСКОВОЕ ЗАЯВЛЕНИЕ",
        court="Специализированный межрайонный экономический суд города Астаны",
        case_number="7199-26-00-2/1234",
        claimant=["ТОО «Астана Строй», БИН 123456789012"],
        defendant=["ТОО «Заказчик», БИН 210987654321"],
        claim_summary=["Истец просит взыскать 2 300 000 тенге долга."],
        disputed_circumstances=["Оспаривается объём принятых работ по акту от 20.02.2026."],
        position=["Иск подлежит частичному удовлетворению."],
        objections=["Работы на сумму 900 000 тенге не приняты."],
        calculation_review=["Неустойка исчислена до наступления срока оплаты."],
        legal_basis=[],
        requests=["Отказать в удовлетворении исковых требований в части 900 000 тенге."],
        attachments=["Копия акта от 20.02.2026 с замечаниями"],
        verification_notes=[],
        source_urls=[],
    )


def _contract(*, place_and_date: str) -> ContractDraft:
    return ContractDraft(
        status=VerificationStatus.VERIFIED,
        contract_type="Договор подряда",
        title="ДОГОВОР ПОДРЯДА № 12",
        place_and_date=place_and_date,
        party_a=["ТОО «Астана Строй», БИН 123456789012"],
        party_b=["ТОО «Заказчик», БИН 210987654321"],
        preamble=["Стороны заключили настоящий договор о нижеследующем."],
        sections=[
            ContractSection(
                heading="Предмет договора",
                clauses=["Подрядчик обязуется выполнить работы, а заказчик — принять и оплатить их."],
            )
        ],
        requisites_a=["г. Алматы, ул. Кенесары, 40"],
        requisites_b=["г. Алматы, ул. Абая, 15"],
        verification_notes=[],
        source_urls=["https://adilet.zan.kz/rus/docs/K990000409_"],
    )


_SIGNED_DOCUMENTS = (
    ("иск", lambda: build_claim_docx(_claim())),
    ("претензия", lambda: build_pretrial_docx(_pretrial())),
    ("ответ на претензию", lambda: build_pretrial_response_docx(_pretrial_response())),
    ("отзыв на иск", lambda: build_response_to_claim_docx(_response_to_claim())),
)


@pytest.mark.parametrize(("kind", "build"), _SIGNED_DOCUMENTS, ids=[item[0] for item in _SIGNED_DOCUMENTS])
def test_unsigned_document_does_not_stamp_the_generation_day(kind: str, build) -> None:
    """День формирования проекта — не дата подписания документа."""
    text = _text(build())

    assert f"Дата: {TODAY_KZ}" not in text, kind
    assert TODAY_KZ not in text, kind


@pytest.mark.parametrize(("kind", "build"), _SIGNED_DOCUMENTS, ids=[item[0] for item in _SIGNED_DOCUMENTS])
def test_date_line_is_left_blank_next_to_the_blank_signature(kind: str, build) -> None:
    """Пустая подпись и заполненная дата не могут стоять рядом."""
    text = _text(build())

    assert f"Подпись: {BLANK}" in text or f"Дата: {BLANK}" in text, kind
    assert f"Дата: {BLANK}" in text, kind


def test_kazakh_export_leaves_the_date_blank_too() -> None:
    token = _CURRENT_LANGUAGE.set(KK)
    try:
        text = _text(build_claim_docx(_claim()))
    finally:
        _CURRENT_LANGUAGE.reset(token)

    assert TODAY_KZ not in text
    assert f"Күні: {BLANK}" in text


def test_contract_without_a_place_and_date_does_not_invent_the_date() -> None:
    """Дата заключения договора — существенный реквизит, а не оформление."""
    text = _text(build_contract_docx(_contract(place_and_date="")))

    assert TODAY_KZ not in text
    assert "[ТРЕБУЕТ УТОЧНЕНИЯ" in text


def test_contract_missing_its_place_and_date_is_not_released_as_ready() -> None:
    """Пропуск в шапке договора обязан понижать статус готовности.

    Заглушка места и даты появляется только при рендере, поэтому в проверку
    статуса она передаётся отдельно: иначе договор без даты заключения зависел
    бы в вопросе готовности от того, оказался ли пропуск заодно и в преамбуле.
    """
    text = _text(build_contract_docx(_contract(place_and_date="")))

    assert "KORGAN QA STATUS: PRELIMINARY DRAFT" in text


def test_contract_with_a_real_place_and_date_keeps_it_verbatim() -> None:
    text = _text(build_contract_docx(_contract(place_and_date="г. Алматы, 15 января 2026 года")))

    assert "г. Алматы, 15 января 2026 года" in text
    assert "[ТРЕБУЕТ УТОЧНЕНИЯ: место и дата заключения договора]" not in text

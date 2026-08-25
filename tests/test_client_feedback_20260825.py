from __future__ import annotations

import io

from docx import Document

from korgan.client_feedback_20260825 import _clean_response_sentence, install_special_part_research_guard
from korgan.contract_docx import build_contract_docx
from korgan.legal_types import ContractClause, ContractDraft, ContractSection, VerificationStatus


def _doc_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def test_bad_response_opening_is_repaired_deterministically() -> None:
    source = "На Рассмотрев Вашу досудебную претензию от 24.08.2026 сообщаем следующее."
    fixed = _clean_response_sentence(source)
    assert fixed.startswith("Рассмотрев Вашу досудебную претензию")
    assert "На Рассмотрев" not in fixed


def test_research_prompt_requires_verified_special_part_for_named_contracts() -> None:
    from korgan import fast_professional_litigation as litigation

    install_special_part_research_guard()
    prompt = litigation._professional_research_prompt(
        "Спор по договору поставки товара между двумя ТОО",
        max_chars=5000,
        checked_on="2026-08-25",
    )
    assert "Особенной части ГК РК" in prompt
    assert "Номер статьи по памяти запрещён" in prompt


def test_contract_combined_requisites_are_not_duplicated_and_spec_is_present() -> None:
    draft = ContractDraft(
        status=VerificationStatus.VERIFIED,
        contract_type="договор поставки",
        title="Договор поставки",
        place_and_date="г. Астана, 25.08.2026",
        party_a=["ТОО «Поставщик», БИН 111111111111"],
        party_b=["ТОО «Покупатель», БИН 222222222222"],
        preamble=[
            "ТОО «Поставщик», именуемое в дальнейшем «Поставщик», и ТОО «Покупатель», именуемое в дальнейшем «Покупатель», заключили настоящий Договор."
        ],
        sections=[
            ContractSection(
                heading="Предмет Договора",
                clauses=[
                    ContractClause("Наименование, характеристики и количество Товара определяются в Спецификации, являющейся неотъемлемой частью Договора."),
                ],
            ),
            ContractSection(
                heading="Заключительные положения и реквизиты сторон",
                clauses=[
                    ContractClause("Настоящий Договор вступает в силу с момента подписания Сторонами."),
                    ContractClause("ТОО «Поставщик», БИН 111111111111, БИК ABCDKZKX, Банк АО «Банк»."),
                    ContractClause("ТОО «Покупатель», БИН 222222222222, БИК EFGHKZKX, Банк АО «Банк 2»."),
                ],
            ),
        ],
        requisites_a=["ТОО «Поставщик»", "БИН 111111111111", "БИК ABCDKZKX"],
        requisites_b=["ТОО «Покупатель»", "БИН 222222222222", "БИК EFGHKZKX"],
        verification_notes=[],
        source_urls=[],
    )

    text = _doc_text(build_contract_docx(draft))
    lowered = text.casefold()

    assert lowered.count("реквизиты и подписи сторон") == 1
    assert "Заключительные положения и реквизиты сторон" not in text
    assert "Заключительные положения" in text
    # BИН legitimately appears once in the party identification/preamble and once
    # in the canonical requisites table. Bank identifiers exist only in requisites
    # and therefore must occur exactly once.
    assert text.count("БИК ABCDKZKX") == 1
    assert text.count("БИК EFGHKZKX") == 1
    assert "СПЕЦИФИКАЦИЯ" in text
    assert "Наименование" in text
    assert "Количество" in text
    assert text.count("Подпись: ____________________") == 2

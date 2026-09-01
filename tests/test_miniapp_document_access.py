from __future__ import annotations

import io

import pytest
from docx import Document
from fastapi import HTTPException

from korgan.miniapp_document_access import (
    _content_disposition,
    _make_token,
    _read_token,
    _render_docx_html,
)


def test_downloaded_file_keeps_its_russian_name() -> None:
    """Имя файла на кириллице обязано доезжать до пользователя.

    Заголовок содержал только ASCII-вариант имени, а он собирался отбрасыванием
    всех не-ASCII символов. Для «Исковое заявление.docx» от имени не оставалось
    ничего, и каждый документ сохранялся под одним и тем же бессмысленным
    именем: найти нужный среди скачанных было невозможно.
    """
    header = _content_disposition("Исковое заявление.docx")

    assert header.startswith("attachment; ")
    # RFC 6266: браузер и Telegram WebView берут filename*, а filename остаётся
    # запасным вариантом для клиентов, которые его не понимают.
    assert "filename*=UTF-8''" in header
    assert "%D0%98" in header  # «И» в percent-encoding
    assert 'filename="' in header


def test_ascii_name_is_not_mangled() -> None:
    header = _content_disposition("KORGAN_claim.docx")

    assert 'filename="KORGAN_claim.docx"' in header


def test_filename_cannot_inject_a_header_or_escape_the_directory() -> None:
    """Имя приходит из состояния дела и заголовок ломать не должно."""
    header = _content_disposition('../../etc/passwd\r\nX-Injected: 1"; drop.docx')

    # Перевод строки закончил бы заголовок и начал новый — его быть не должно.
    assert "\r" not in header
    assert "\n" not in header
    # Кавычка закрыла бы имя досрочно: их ровно две, обе — границы ASCII-имени.
    assert header.count('"') == 2
    # Разделители пути не выходят за пределы имени файла.
    assert "/" not in header
    assert "\\" not in header


def test_a_name_without_the_docx_extension_gets_one() -> None:
    header = _content_disposition("Отзыв на иск")

    assert 'filename="KORGAN_document.docx"' in header
    assert header.rstrip().endswith(".docx")


def test_document_access_token_roundtrip_and_expiry() -> None:
    token = _make_token("12345", "KOR-ABC123", 1_000, "test-secret")
    assert _read_token(token, "test-secret", now=900) == ("12345", "KOR-ABC123")

    with pytest.raises(HTTPException) as expired:
        _read_token(token, "test-secret", now=1_001)
    assert expired.value.status_code == 403


def test_document_access_token_rejects_tampering() -> None:
    token = _make_token("12345", "KOR-ABC123", 1_000, "test-secret")
    payload, signature = token.split(".", 1)
    tampered = f"{payload[:-1]}A.{signature}"
    with pytest.raises(HTTPException) as invalid:
        _read_token(tampered, "test-secret", now=900)
    assert invalid.value.status_code == 403


def test_document_preview_is_first_party_html_and_escapes_text() -> None:
    document = Document()
    document.add_heading("Исковое заявление", level=1)
    document.add_paragraph("Факт <не должен> превращаться в HTML & script")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Истец"
    table.cell(0, 1).text = "ТОО <Альфа>"
    buffer = io.BytesIO()
    document.save(buffer)

    rendered = _render_docx_html(buffer.getvalue(), "KORGAN <preview>")

    assert "Исковое заявление" in rendered
    assert "&lt;не должен&gt;" in rendered
    assert "ТОО &lt;Альфа&gt;" in rendered
    assert "KORGAN &lt;preview&gt;" in rendered
    assert "<script" not in rendered.lower()

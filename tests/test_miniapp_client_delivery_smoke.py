"""Оплаченный документ проходит через собранное Mini App и signed delivery.

Этот smoke не дублирует отдельные тесты банковской/OFD/Tole верификации. Его
задача — связать публичный клиентский контракт: дело → payment gate → только
approved order → сохраняемая generation job → настоящий DOCX → signed download
→ повторное открытие Mini App. Устаревшая отправка документа через Telegram
обязана оставаться закрытой с 410.
"""

from __future__ import annotations

import asyncio
import base64
import hashlib
import hmac
import io
import json
import threading
import time
from contextlib import asynccontextmanager
from dataclasses import replace
from typing import Any
from urllib.parse import urlencode, urlsplit

from docx import Document
from fastapi.testclient import TestClient

from korgan import miniapp_api
from korgan import miniapp_api_v2 as core
from korgan import miniapp_generation_api as generation_api
from korgan import miniapp_generation_jobs as jobs
from korgan.miniapp_api_recovery_cors import app
from korgan.miniapp_document_payments import DocumentPaymentOrder
from korgan.miniapp_generation_jobs import GenerationJob

TERMS_VERSION = "2026-08-16-v1"
USER_ID = -900000821
ORDER_ID = 821
JOB_ID = "00000000-0000-0000-0000-000000000821"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _init_data(user_id: int) -> str:
    pairs = {
        "auth_date": str(int(time.time())),
        "query_id": "korgan-client-delivery-smoke",
        "user": json.dumps(
            {"id": user_id, "first_name": "KORGAN", "language_code": "ru"},
            ensure_ascii=False,
            separators=(",", ":"),
        ),
    }
    data_check_string = "\n".join(f"{key}={pairs[key]}" for key in sorted(pairs))
    secret_key = hmac.new(
        b"WebAppData",
        miniapp_api.settings.telegram_bot_token.encode(),
        hashlib.sha256,
    ).digest()
    pairs["hash"] = hmac.new(secret_key, data_check_string.encode(), hashlib.sha256).hexdigest()
    return urlencode(pairs)


def _headers(user_id: int) -> dict[str, str]:
    return {"X-Telegram-Init-Data": _init_data(user_id)}


def _docx() -> bytes:
    document = Document()
    document.add_heading("Исковое заявление", level=1)
    document.add_paragraph("О взыскании задолженности по договору подряда.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_paid_document_survives_reopen_and_uses_signed_delivery(monkeypatch) -> None:
    owner_headers = _headers(USER_ID)
    state_by_identity: dict[str, dict[str, Any]] = {}
    order: DocumentPaymentOrder | None = None
    job: GenerationJob | None = None
    case_id = ""
    worker_finished = threading.Event()
    document_bytes = _docx()

    def payment(status: str) -> DocumentPaymentOrder:
        return DocumentPaymentOrder(
            id=ORDER_ID,
            user_key=core.store.user_key(str(USER_ID)),
            case_id=case_id,
            case_fingerprint="scope-smoke",
            document_type="claim",
            language="ru",
            amount_kzt=1000,
            status=status,
            transaction_id="KASPI-SMOKE-821" if status == "approved" else "",
            receipt_check={},
            decision_note="",
        )

    async def create_order(**_kwargs):
        nonlocal order
        if order is None:
            order = payment("pending_receipt")
        return order

    async def get_scope_order(**_kwargs):
        return order

    async def get_order(order_id: int, *, user_key: str | None = None):
        if order is None or order_id != ORDER_ID:
            return None
        if user_key is not None and user_key != order.user_key:
            return None
        return order

    async def create_job(**_kwargs):
        nonlocal job
        if job is None:
            job = GenerationJob(
                id=JOB_ID,
                payment_order_id=ORDER_ID,
                user_key=core.store.user_key(str(USER_ID)),
                case_id=case_id,
                status="queued",
                stage="queued",
                progress=0,
                error_detail="",
            )
        return job

    async def latest_job(**kwargs):
        if job is None or kwargs["case_id"] != case_id:
            return None
        return job

    async def require_job(job_id: str, *, user_key: str):
        if job is None or job.id != job_id or job.user_key != user_key:
            raise AssertionError("smoke requested another or missing generation job")
        return job

    async def run_worker(started_job, **kwargs) -> None:
        nonlocal job, order
        job = replace(started_job, status="running", stage="legal_research", progress=20)
        await asyncio.sleep(0)
        state = await core.store.load(kwargs["identity"])
        state["cases"][started_job.case_id].update(
            {
                "status": "document_ready",
                "title": "Исковое заявление",
                "filename": "KORGAN_claim.docx",
                "document_base64": base64.b64encode(document_bytes).decode("ascii"),
                "filing_ready": False,
                "release_status": "preliminary",
                "verification_status": "needs_verification",
                "verification_notes": ["Требуется финальная проверка юристом"],
                "quality_score": 9.0,
                "quality_issues": ["Проверить приложения"],
            }
        )
        await core.store.save(kwargs["identity"], state)
        assert order is not None
        order = replace(order, status="consumed")
        job = replace(job, status="succeeded", stage="completed", progress=100)
        worker_finished.set()

    monkeypatch.setattr(generation_api.settings, "payments_enabled", True)
    monkeypatch.setattr(generation_api.settings, "kaspi_payment_url", "https://pay.korgan.test")
    monkeypatch.setattr(generation_api.settings, "document_price_kzt", 1000)

    async def noop_store(*args, **kwargs) -> None:
        return None

    monkeypatch.setattr(generation_api.v5.v4, "init_document_payment_store", noop_store)
    monkeypatch.setattr(generation_api.v5.v4, "close_document_payment_store", noop_store)
    monkeypatch.setattr(generation_api.jobs, "init_generation_job_store", noop_store)
    monkeypatch.setattr(generation_api.jobs, "close_generation_job_store", noop_store)

    @asynccontextmanager
    async def noop_payment_lock(*args, **kwargs):
        yield

    monkeypatch.setattr(generation_api, "payment_operation_lock", noop_payment_lock)
    monkeypatch.setattr(generation_api.document_store, "_require_pool", lambda: object())
    monkeypatch.setattr(generation_api.document_store, "create_document_order", create_order)
    monkeypatch.setattr(generation_api.document_store, "get_scope_order", get_scope_order)
    monkeypatch.setattr(generation_api.document_store, "get_document_order", get_order)
    monkeypatch.setattr(generation_api.v5.v4, "_document_scope", lambda *args: "scope-smoke")
    monkeypatch.setattr(jobs, "create_or_get_job", create_job)
    monkeypatch.setattr(jobs, "latest_job_for_case", latest_job)
    monkeypatch.setattr(jobs, "require_job", require_job)
    monkeypatch.setattr(jobs, "run_job", run_worker)

    async def load(identity: str) -> dict[str, Any]:
        return state_by_identity.setdefault(identity, {"consent": None, "cases": {}})

    async def save(identity: str, state: dict[str, Any]) -> None:
        state_by_identity[identity] = state

    async def delete(identity: str) -> None:
        state_by_identity.pop(identity, None)

    monkeypatch.setattr(core.legacy, "_state", load)
    monkeypatch.setattr(core.store, "load", load)
    monkeypatch.setattr(core.store, "save", save)
    monkeypatch.setattr(core.store, "delete", delete)

    with TestClient(app) as client:
        accepted = client.post(
            "/miniapp/consent",
            headers=owner_headers,
            json={"accepted": True, "terms_version": TERMS_VERSION},
        )
        assert accepted.status_code == 200

        created = client.post(
            "/miniapp/cases",
            headers=owner_headers,
            json={
                "description": "Заказчик не оплатил выполненные работы по договору.",
                "document_type": "claim",
                "language": "ru",
            },
        )
        assert created.status_code == 200
        case_id = created.json()["case"]["id"]

        # No legal generation may start before the order is approved.
        payment_requested = client.post(
            "/miniapp/documents/generate",
            headers=owner_headers,
            json={"case_id": case_id, "document_type": "claim", "language": "ru"},
        )
        assert payment_requested.status_code == 200
        assert payment_requested.json()["payment_required"] is True
        assert payment_requested.json()["generation_started"] is False
        assert payment_requested.json()["payment"]["status"] == "pending_receipt"
        assert job is None

        # Bank/provider verification is tested separately; this smoke begins the
        # generation half only after the same persisted order is approved.
        assert order is not None
        order = replace(order, status="approved", transaction_id="KASPI-SMOKE-821")

        started = client.post(
            "/miniapp/documents/generate",
            headers=owner_headers,
            json={"case_id": case_id, "document_type": "claim", "language": "ru"},
        )
        assert started.status_code == 200
        assert started.json()["payment_required"] is False
        assert started.json()["job"]["job_id"] == JOB_ID

        deadline = time.monotonic() + 1.0
        while not worker_finished.is_set() and time.monotonic() < deadline:
            time.sleep(0.001)
        assert worker_finished.is_set(), "generation worker did not complete"

        completed = client.get(f"/miniapp/documents/generation/{JOB_ID}", headers=owner_headers)
        assert completed.status_code == 200
        assert completed.json()["job"]["document_ready"] is True
        assert completed.json()["document"]["filename"] == "KORGAN_claim.docx"
        assert "document_base64" not in completed.json()["document"]

        issued = client.post(f"/miniapp/cases/{case_id}/document/access", headers=owner_headers)
        assert issued.status_code == 200
        links = issued.json()
        download = urlsplit(links["download_url"])
        preview_url = urlsplit(links["preview_url"])
        downloaded = client.get(download.path + "?" + download.query)
        preview = client.get(preview_url.path + "?" + preview_url.query)
        assert downloaded.status_code == 200
        assert downloaded.headers["content-type"] == DOCX_MIME
        assert downloaded.content == document_bytes
        assert Document(io.BytesIO(downloaded.content)).paragraphs[0].text == "Исковое заявление"
        assert preview.status_code == 200
        assert "О взыскании задолженности" in preview.text

        # Retired transport stays fail-closed; private legal files use signed
        # access instead of being forwarded through the bot.
        retired = client.post(f"/miniapp/cases/{case_id}/document/telegram", headers=owner_headers)
        assert retired.status_code == 410

    # Reopen: server-side case and completed job survive loss of React state.
    with TestClient(app) as reopened:
        cases = reopened.get("/miniapp/cases", headers=owner_headers)
        assert cases.status_code == 200
        assert [item["id"] for item in cases.json()["cases"]] == [case_id]
        assert cases.json()["cases"][0]["has_document"] is True

        recovered = reopened.get(f"/miniapp/cases/{case_id}/generation", headers=owner_headers)
        assert recovered.status_code == 200
        assert recovered.json()["job"]["document_ready"] is True
        assert recovered.json()["document"]["filename"] == "KORGAN_claim.docx"

        opened_again = reopened.post(f"/miniapp/cases/{case_id}/document/access", headers=owner_headers)
        assert opened_again.status_code == 200

        deleted = reopened.delete(f"/miniapp/cases/{case_id}", headers=owner_headers)
        assert deleted.status_code == 200
        assert reopened.get("/miniapp/cases", headers=owner_headers).json()["cases"] == []
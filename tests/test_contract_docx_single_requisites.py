from __future__ import annotations

import io

from docx import Document

from korgan.contract_docx import build_contract_docx
from korgan.legal_types import ContractClause, ContractDraft, ContractSection, VerificationStatus


def _all_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def test_contract_renderer_outputs_requisites_and_signatures_once() -> None:
    draft = ContractDraft(
        status=VerificationStatus.VERIFIED,
        contract_type="договор поставки",
        title="Договор поставки",
        place_and_date="г. Алматы, 30.08.2026",
        party_a=["ТОО «Астана Трейд»"],
        party_b=["ТОО «Бизнес Партнер»"],
        preamble=[
            "ТОО «Астана Трейд», именуемое в дальнейшем «Поставщик», с одной стороны, и ТОО «Бизнес Партнер», именуемое в дальнейшем «Покупатель», с другой стороны, заключили настоящий Договор."
        ],
        sections=[
            ContractSection(
                heading="Предмет Договора",
                clauses=[ContractClause("Поставщик передает товар, а Покупатель принимает и оплачивает его.")],
            ),
            ContractSection(
                heading="Реквизиты и подписи Сторон",
                clauses=[
                    ContractClause("Поставщик: ТОО «Астана Трейд»"),
                    ContractClause("БИН: 220340078912"),
                    ContractClause("Юридический адрес: __________________"),
                    ContractClause("Покупатель: ТОО «Бизнес Партнер»"),
                    ContractClause("БИН: 230540067891"),
                    ContractClause("Подпись: ____________________"),
                ],
            ),
        ],
        requisites_a=[
            "ТОО «Астана Трейд»",
            "БИН 220340078912",
            "Юридический адрес: __________________",
            "Подпись: ____________________",
        ],
        requisites_b=[
            "ТОО «Бизнес Партнер»",
            "БИН 230540067891",
            "Юридический адрес: __________________",
            "Подпись: ____________________",
        ],
        verification_notes=[],
        source_urls=[],
    )

    data = build_contract_docx(draft)
    doc = Document(io.BytesIO(data))
    text = _all_text(data)
    lowered = text.casefold()

    assert lowered.count("реквизиты и подписи сторон") == 1
    assert "12.1." not in text
    assert "12.2." not in text
    assert text.count("Подпись: ____________________") == 2
    assert text.count("БИН 220340078912") == 1
    assert text.count("БИН 230540067891") == 1

    signing_tables = [table for table in doc.tables if len(table.columns) == 2]
    assert signing_tables, "Final requisites/signatures must be rendered as a two-column table"


def test_combined_final_section_keeps_legal_clauses_but_not_requisites_rows() -> None:
    draft = ContractDraft(
        status=VerificationStatus.VERIFIED,
        contract_type="договор",
        title="Договор",
        place_and_date="г. Алматы, 30.08.2026",
        party_a=["ТОО «А»"],
        party_b=["ТОО «Б»"],
        preamble=["ТОО «А» и ТОО «Б» заключили настоящий Договор."],
        sections=[
            ContractSection(
                heading="Заключительные положения и реквизиты сторон",
                clauses=[
                    ContractClause("Договор вступает в силу с момента подписания Сторонами."),
                    ContractClause("БИН: 123456789012"),
                    ContractClause("Банк: __________________"),
                ],
            )
        ],
        requisites_a=["ТОО «А»", "БИН 123456789012"],
        requisites_b=["ТОО «Б»", "БИН 210987654321"],
        verification_notes=[],
        source_urls=[],
    )

    text = _all_text(build_contract_docx(draft))
    assert "Заключительные положения" in text
    assert "Договор вступает в силу с момента подписания Сторонами." in text
    assert text.count("БИН 123456789012") == 1
    assert "Банк: __________________" not in text

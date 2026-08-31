"""Ответ на претензию и отзыв на иск — позиция стороны, а не отписка.

Профессиональный ответ разбирает требования оппонента по существу: отделяет
признаваемое от оспариваемого, проверяет его расчёт и обосновывает каждое
возражение. Шаблонное «с требованиями не согласны» без разбора — не документ.

Здесь фиксируется структура, которая делает такой разбор обязательным и
проверяемым, и запрет на признания, которых доверитель не делал.
"""

from __future__ import annotations

from io import BytesIO

from docx import Document

from korgan.document_quality import assess_document_quality
from korgan.legal_types import LegalResearch, VerificationStatus
from korgan.pretrial_response import (
    PretrialResponseDraft,
    build_pretrial_response_docx,
    pretrial_response_quality_issues,
)
from korgan.provision_check import verified_claim_line
from korgan.response_docx import build_response_to_claim_docx
from korgan.response_types import ResponseToClaimDraft

GK_URL = "https://adilet.zan.kz/rus/docs/K990000409_"
CASE_CONTEXT = (
    "Истец: ТОО «Астана Строй», БИН 123456789012.\n"
    "Ответчик: ТОО «Заказчик», БИН 210987654321.\n"
    "Иск подан в Специализированный межрайонный экономический суд города Астаны, "
    "дело № 7199-26-00-2/1234.\n"
    "Договор подряда № 12 от 15.01.2026, акт от 20.02.2026 подписан с замечаниями."
)
ARTICLE_623 = (
    "Заказчик обязан уплатить подрядчику обусловленную цену после окончательной сдачи "
    "результатов работы при условии, что работа выполнена надлежащим образом и в согласованный срок."
)


def _research() -> LegalResearch:
    return LegalResearch(
        status=VerificationStatus.VERIFIED,
        applicable_law=[],
        procedural_requirements=[],
        verified_claims=[
            verified_claim_line(
                "Заказчик обязан оплатить принятые работы",
                "статья 623 ГК РК",
                ARTICLE_623,
                GK_URL,
            )
        ],
        unverified_claims=[],
        source_urls=[GK_URL],
        notes=[],
    )


# --- ответ на досудебную претензию ---------------------------------------


def _pretrial_response(**overrides) -> PretrialResponseDraft:
    data = dict(
        status=VerificationStatus.VERIFIED,
        title="ОТВЕТ НА ДОСУДЕБНУЮ ПРЕТЕНЗИЮ",
        sender=['ТОО «Заказчик», БИН 210987654321, г. Астана, ул. Абая, 15'],
        recipient=['ТОО «Астана Строй», БИН 123456789012, г. Астана, ул. Кенесары, 40'],
        reference="претензия от 05.03.2026 № 7",
        claim_summary=[
            "Заявлено требование об оплате 2 300 000 тенге задолженности и 71 300 тенге неустойки.",
        ],
        admitted_circumstances=[
            "Факт заключения договора подряда № 12 от 15.01.2026 не оспаривается.",
        ],
        disputed_circumstances=[
            "Оспаривается объём принятых работ: акт от 20.02.2026 подписан с замечаниями.",
        ],
        position=[
            "Требование признаётся в части 1 400 000 тенге и оспаривается в остальной части.",
        ],
        objections=[
            "Работы на сумму 900 000 тенге не приняты: в акте от 20.02.2026 зафиксированы недостатки, "
            "не устранённые на дату настоящего ответа.",
        ],
        calculation_review=[
            "Неустойка исчислена с 01.03.2026, тогда как срок оплаты по пункту 4.2 договора наступает "
            "20.03.2026; начисление за 19 дней является необоснованным.",
        ],
        legal_basis=[f"{ARTICLE_623} Правовое основание: статья 623 ГК РК."],
        settlement_offer="Готовы оплатить признанную часть 1 400 000 тенге в течение 10 рабочих дней.",
        response_terms=["Оплата признанной части будет произведена в согласованный сторонами срок."],
        attachments=["Копия акта от 20.02.2026 с замечаниями"],
        verification_notes=[],
        source_urls=[GK_URL],
    )
    data.update(overrides)
    return PretrialResponseDraft(**data)


def test_pretrial_response_separates_admitted_from_disputed() -> None:
    draft = _pretrial_response()
    assert draft.admitted_circumstances
    assert draft.disputed_circumstances
    assert pretrial_response_quality_issues(draft, _research()) == []


def test_bare_disagreement_without_reasoning_is_blocked() -> None:
    draft = _pretrial_response(
        position=["С требованиями не согласны."],
        objections=["Требования не признаём."],
        disputed_circumstances=[],
        calculation_review=[],
    )
    issues = pretrial_response_quality_issues(draft, _research())
    assert any("не обоснован" in issue.lower() or "без обоснования" in issue.lower() for issue in issues)


def test_admissions_are_never_manufactured() -> None:
    """Пустое признание допустимо: доверитель не обязан признавать ничего."""
    draft = _pretrial_response(admitted_circumstances=[])
    issues = pretrial_response_quality_issues(draft, _research())
    assert not any("призна" in issue.lower() for issue in issues)

    report = assess_document_quality("pretrial_response", CASE_CONTEXT, _research(), draft)
    assert report.ready is True


def test_money_claim_requires_a_calculation_review() -> None:
    draft = _pretrial_response(calculation_review=[])
    issues = pretrial_response_quality_issues(draft, _research())
    assert any("расч" in issue.lower() for issue in issues)


def test_pretrial_response_word_shows_admitted_disputed_and_calculation_review() -> None:
    document = Document(BytesIO(build_pretrial_response_docx(_pretrial_response())))
    body = "\n".join(p.text for p in document.paragraphs)

    assert "Признаваемые обстоятельства" in body
    assert "Оспариваемые обстоятельства" in body
    assert "Разбор расчёта" in body
    assert "не оспаривается" in body
    assert "начисление за 19 дней" in body


def test_pretrial_response_word_omits_empty_sections() -> None:
    draft = _pretrial_response(admitted_circumstances=[], calculation_review=[])
    document = Document(BytesIO(build_pretrial_response_docx(draft)))
    body = "\n".join(p.text for p in document.paragraphs)

    assert "Признаваемые обстоятельства" not in body
    assert "Разбор расчёта" not in body
    assert "Оспариваемые обстоятельства" in body


# --- отзыв на иск ---------------------------------------------------------


def _response(**overrides) -> ResponseToClaimDraft:
    data = dict(
        status=VerificationStatus.VERIFIED,
        title="ОТЗЫВ НА ИСКОВОЕ ЗАЯВЛЕНИЕ",
        court="Специализированный межрайонный экономический суд города Астаны",
        case_number="дело № 7199-26-00-2/1234",
        claimant=['ТОО «Астана Строй», БИН 123456789012'],
        defendant=['ТОО «Заказчик», БИН 210987654321'],
        claim_summary=["Истец просит взыскать 2 300 000 тенге долга и 71 300 тенге неустойки."],
        admitted_circumstances=["Факт заключения договора подряда № 12 от 15.01.2026 не оспаривается."],
        disputed_circumstances=["Оспаривается объём принятых работ по акту от 20.02.2026."],
        position=["Иск подлежит частичному удовлетворению в размере 1 400 000 тенге."],
        objections=[
            "Работы на сумму 900 000 тенге не приняты: акт от 20.02.2026 подписан с замечаниями.",
        ],
        calculation_review=[
            "Неустойка исчислена с 01.03.2026 при сроке оплаты 20.03.2026 по пункту 4.2 договора.",
        ],
        legal_basis=[f"{ARTICLE_623} Правовое основание: статья 623 ГК РК."],
        requests=["Отказать в удовлетворении исковых требований в части 900 000 тенге."],
        attachments=["Копия акта от 20.02.2026 с замечаниями"],
        verification_notes=[],
        source_urls=[GK_URL],
    )
    data.update(overrides)
    return ResponseToClaimDraft(**data)


def test_response_to_claim_separates_admitted_from_disputed() -> None:
    draft = _response()
    report = assess_document_quality("response_to_claim", CASE_CONTEXT, _research(), draft)
    assert report.ready is True
    assert draft.admitted_circumstances
    assert draft.disputed_circumstances


def test_response_to_claim_word_shows_the_new_sections() -> None:
    document = Document(BytesIO(build_response_to_claim_docx(_response())))
    body = "\n".join(p.text for p in document.paragraphs)

    assert "Признаваемые обстоятельства" in body
    assert "Оспариваемые обстоятельства" in body
    assert "Разбор расчёта истца" in body


def test_response_to_claim_word_omits_empty_sections() -> None:
    draft = _response(admitted_circumstances=[], calculation_review=[])
    document = Document(BytesIO(build_response_to_claim_docx(draft)))
    body = "\n".join(p.text for p in document.paragraphs)

    assert "Признаваемые обстоятельства" not in body
    assert "Разбор расчёта истца" not in body


def test_limitation_period_is_only_raised_when_dates_support_it() -> None:
    """Исковая давность не заявляется ради объёма."""
    draft = _response(objections=["Истёк срок исковой давности."], disputed_circumstances=[])
    report = assess_document_quality("response_to_claim", CASE_CONTEXT, _research(), draft)
    assert report.ready is False

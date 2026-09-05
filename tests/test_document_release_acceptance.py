"""Приёмка выпуска: три документа доходят до клиента каждый своим конвейером.

Боевой дефект, ради которого написан этот файл: заявка «Досудебная претензия»
выдала ``KORGAN_otvet_na_pretenziyu.docx`` с заголовком «ОТВЕТ НА ПРЕТЕНЗИЮ» и
перевёрнутыми ролями. Проверяется весь путь настоящим кодом — исследование,
черновик, финальная проверка, рендер Word, сохранение состояния, переход задачи
в READY, — подменён только провайдер модели.
"""

from __future__ import annotations

import asyncio
import base64
import io
import json
from typing import Any

import pytest
from docx import Document

from korgan import generation_progress as progress
from korgan import miniapp_generation_jobs as jobs
from tests.test_document_generation_e2e import (  # noqa: F401 — фикстуры переиспользуются
    FakePool,
    FakeStore,
    Recorder,
    _fill,
    _job,
    provider,
    wired,
)

CREDITOR = "ТОО «Альфа Трейд», БИН 123456789012, г. Алматы, ул. Абая, 10"
DEBTOR = "ТОО «Бета Снаб», БИН 210987654321, г. Алматы, ул. Толе би, 20"

PRETRIAL_CONTEXT = (
    f"Кредитор {CREDITOR}. Должник {DEBTOR}. "
    "10.01.2026 заключён договор поставки №12. 15.01.2026 поставлен товар на 4 500 000 тенге, "
    "поставка подтверждена договором, товарной накладной и актом приёма-передачи, товар принят "
    "без замечаний. Срок оплаты — 15 календарных дней после поставки. Оплата не произведена. "
    "Нужна обычная досудебная претензия от кредитора должнику."
)
PRETRIAL_RESPONSE_CONTEXT = (
    f"Претензию направило {CREDITOR}, получатель — {DEBTOR}. "
    "Требуют оплатить 4 500 000 тенге по договору поставки №12 от 10.01.2026. "
    "Нужен ответ на претензию от имени получателя."
)
CLAIM_CONTEXT = (
    f"Истец {CREDITOR}. Ответчик {DEBTOR}. "
    "10.01.2026 заключён договор поставки №12, 15.01.2026 поставлен товар на 4 500 000 тенге, "
    "принят по акту приёма-передачи без замечаний. Срок оплаты — 15 календарных дней после "
    "поставки, оплата не произведена. 15.02.2026 направлена претензия, получена 17.02.2026, "
    "ответа нет. Прошу взыскать основной долг 4 500 000 тенге и государственную пошлину. "
    "Моральный вред, неустойку и проценты не заявлять."
)

#: Заголовки, которые модель обязана вернуть для каждого документа. Стаб
#: провайдера из e2e-набора отвечает обобщённо; здесь заголовок важен именно
#: потому, что боевой дефект был виден в нём.
_TITLES = {
    "korgan_pretrial_demand": "ПРЕТЕНЗИЯ",
    "korgan_10_of_10_pretrial": "ПРЕТЕНЗИЯ",
    "korgan_pretrial_response": "ОТВЕТ НА ПРЕТЕНЗИЮ",
    "korgan_10_of_10_pretrial_response": "ОТВЕТ НА ПРЕТЕНЗИЮ",
}


class _Response:
    def __init__(self, text: str) -> None:
        self.output_text = text
        self.output: list[Any] = []
        self.usage = None


class _TitledResponses:
    """Тот же стаб, но с осмысленным заголовком по имени схемы."""

    def __init__(self, owner: "TitledProvider") -> None:
        self._owner = owner

    async def create(self, **kwargs: Any) -> _Response:
        self._owner.calls += 1
        fmt = (kwargs.get("text") or {}).get("format") or {}
        schema = fmt.get("schema") or {"type": "object", "properties": {}}
        payload = _fill(schema)
        title = _TITLES.get(str(fmt.get("name") or ""))
        if title and isinstance(payload, dict) and "title" in payload:
            payload["title"] = title
        if isinstance(payload, dict):
            if "sender" in payload:
                payload["sender"] = [self._owner.sender]
            if "recipient" in payload:
                payload["recipient"] = [self._owner.recipient]
        return _Response(json.dumps(payload, ensure_ascii=False))


class TitledProvider:
    def __init__(self, *, sender: str, recipient: str) -> None:
        self.calls = 0
        self.sender = sender
        self.recipient = recipient
        self.responses = _TitledResponses(self)


@pytest.fixture()
def titled_provider(monkeypatch: pytest.MonkeyPatch):
    from korgan import openai_legal
    from tests.test_document_generation_e2e import _service_objects

    def _install(*, sender: str, recipient: str) -> TitledProvider:
        fake = TitledProvider(sender=sender, recipient=recipient)
        monkeypatch.setattr(openai_legal, "build_legal_client", lambda settings: (fake, "fake"))
        for target in _service_objects():
            monkeypatch.setattr(target, "client", fake, raising=False)
        return fake

    return _install


def _run(store, *, document_type: str, context: str, job_id: str = "job-1", order: int = 91) -> None:
    asyncio.run(
        jobs.run_job(
            _job(job_id, order=order),
            identity="identity",
            store=store,
            document_type=document_type,
            context=context,
            language="ru",
        )
    )


def _docx_text(store, case_id: str = "case-1") -> tuple[str, dict[str, Any]]:
    case = store.state["cases"][case_id]
    data = base64.b64decode(case["document_base64"])
    assert data, "документ пустой"
    document = Document(io.BytesIO(data))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            text += "\n" + " ".join(cell.text for cell in row.cells)
    return text, case


# --- три документа доходят до клиента --------------------------------------

def test_claim_generates_and_opens(provider, wired) -> None:
    store, recorder, consumed = wired
    _run(store, document_type="claim", context=CLAIM_CONTEXT)

    assert ("succeeded", "completed") in recorder.transitions
    text, case = _docx_text(store)
    assert text.strip()
    assert case["filename"] == "KORGAN_iskovoe_zayavlenie.docx"
    assert consumed == [91]


def test_pretrial_generates_with_its_own_filename_and_roles(titled_provider, wired) -> None:
    """Обычная претензия: файл претензии, отправитель — кредитор."""
    titled_provider(sender=CREDITOR, recipient=DEBTOR)
    store, recorder, _ = wired
    _run(store, document_type="pretrial", context=PRETRIAL_CONTEXT)

    assert ("succeeded", "completed") in recorder.transitions
    text, case = _docx_text(store)
    assert case["filename"] == "KORGAN_dosudebnaya_pretenziya.docx"

    upper = text.upper()
    assert "ОТВЕТ НА ПРЕТЕНЗИЮ" not in upper, "претензия вышла как ответ на претензию"
    assert "ПРЕТЕНЗИЯ" in upper

    # Роли не перевёрнуты. Шапка делового письма печатает сначала адресата
    # («Кому:»), затем отправителя («От:»), поэтому проверяется именно связка
    # метки со стороной, а не порядок имён в тексте.
    recipient_block = text.split("От:", 1)[0]
    sender_block = text.split("От:", 1)[1] if "От:" in text else ""
    assert "Бета Снаб" in recipient_block, "адресатом претензии оказался не должник"
    assert "Альфа Трейд" in sender_block, "отправителем претензии оказался не кредитор"

    lowered = text.lower()
    for admission in ("не оспариваем задолженность", "готовы обсудить погашение"):
        assert admission not in lowered, "в претензии кредитора появилось признание должника"


def test_pretrial_response_generates_with_its_own_filename(titled_provider, wired) -> None:
    titled_provider(sender=DEBTOR, recipient=CREDITOR)
    store, recorder, _ = wired
    _run(store, document_type="pretrial_response", context=PRETRIAL_RESPONSE_CONTEXT)

    assert ("succeeded", "completed") in recorder.transitions
    text, case = _docx_text(store)
    assert case["filename"] == "KORGAN_otvet_na_pretenziyu.docx"
    assert "ОТВЕТ НА ПРЕТЕНЗИЮ" in text.upper()


def test_pretrial_and_pretrial_response_never_share_a_filename(provider, wired) -> None:
    """Две заявки подряд не должны выдать один и тот же документ."""
    store, _, _ = wired
    _run(store, document_type="pretrial", context=PRETRIAL_CONTEXT)
    first = store.state["cases"]["case-1"]["filename"]
    _run(store, document_type="pretrial_response", context=PRETRIAL_RESPONSE_CONTEXT, job_id="job-2", order=92)
    second = store.state["cases"]["case-1"]["filename"]
    assert first == "KORGAN_dosudebnaya_pretenziya.docx"
    assert second == "KORGAN_otvet_na_pretenziyu.docx"


# --- прогресс и идемпотентность --------------------------------------------

def test_progress_advances_through_the_real_pipeline_stages(provider, wired, monkeypatch) -> None:
    """Стадии приходят из конвейера, а не из таймера на экране."""
    store, _, _ = wired
    seen: list[tuple[str, int]] = []

    async def record(job_id: str, *, stage: str, progress_value: int) -> None:
        seen.append((stage, progress_value))

    monkeypatch.setattr(jobs, "advance_stage", record)
    _run(store, document_type="claim", context=CLAIM_CONTEXT)

    stages = [stage for stage, _ in seen]
    assert progress.LEGAL_RESEARCH in stages
    assert progress.DRAFTING in stages
    assert progress.LEGAL_QA in stages
    assert progress.DOCX_RENDER in stages
    assert stages == sorted(stages, key=progress.progress_for), "стадии пришли не по порядку"
    assert [value for _, value in seen] == sorted(value for _, value in seen)


def test_repeated_job_run_does_not_duplicate_document_or_payment(provider, wired) -> None:
    store, _, consumed = wired
    _run(store, document_type="claim", context=CLAIM_CONTEXT)
    saves_after_first = store.saves

    # Тот же job_id: переход состояния в базе выигрывает только один исполнитель.
    _run(store, document_type="claim", context=CLAIM_CONTEXT)

    assert store.saves == saves_after_first, "второй запуск записал документ второй раз"
    assert consumed == [91], "оплата списана более одного раза"


def test_public_job_reports_ready_only_with_a_saved_document() -> None:
    running = jobs.GenerationJob(
        id="job-1",
        payment_order_id=1,
        user_key="user-key",
        case_id="case-1",
        status="running",
        stage=progress.DRAFTING,
        progress=progress.progress_for(progress.DRAFTING),
        error_detail="",
    )
    assert jobs.public_job(running)["document_ready"] is False
    assert jobs.public_job(running)["stage"] == progress.DRAFTING

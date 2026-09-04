"""Сквозная генерация через НАСТОЯЩИЙ production-конвейер.

Существующий `test_miniapp_client_delivery_smoke` подменяет `jobs.run_job`
целиком: он проверяет HTTP-контракт выдачи, но не то, что юридический конвейер
действительно производит открывающийся Word. Здесь подменяется только внешний
провайдер модели — единственная зависимость, которой в тестовом окружении нет,
— а исследование, черновик, финальная проверка, рендер DOCX, сохранение
состояния и переход задачи в READY выполняются настоящим кодом.

Проверяется то, из-за чего документ не доходит до клиента на практике: файл
не создан, файл пустой, файл не открывается, READY выставлен раньше файла,
повтор порождает второй документ, упавшая задача отравляет следующую.
"""

from __future__ import annotations

import asyncio
import base64
import io
import json
from contextlib import asynccontextmanager
from typing import Any

import pytest
from docx import Document

from korgan import miniapp_generation_jobs as jobs

CONTEXT = (
    "Истец ТОО «Алтын Курылыс», БИН 123456789012, г. Алматы, ул. Абая, 10. "
    "Ответчик ТОО «Мега Строй», БИН 210987654321, г. Алматы, пр. Достык, 55. "
    "12.01.2025 заключён договор поставки № 17. Истец поставил товар на 4 500 000 тенге, "
    "ответчик принял его по накладной № 44 от 20.01.2025 без замечаний и не оплатил. "
    "Претензия направлена 03.03.2025, ответа нет. Прошу взыскать основной долг 4 500 000 тенге."
)


# --------------------------------------------------------------------------- #
# Стаб провайдера: отвечает по фактической схеме запроса, факты дела сохраняет.
# --------------------------------------------------------------------------- #
class _Response:
    def __init__(self, text: str) -> None:
        self.output_text = text
        self.output: list[Any] = []
        self.usage = None


def _fill(schema: dict[str, Any], path: str = "") -> Any:
    kind = schema.get("type")
    if isinstance(kind, list):
        kind = next((item for item in kind if item != "null"), "string")
    if "enum" in schema:
        return schema["enum"][0]
    if kind == "object":
        return {key: _fill(value, f"{path}.{key}") for key, value in (schema.get("properties") or {}).items()}
    if kind == "array":
        return [_fill(schema.get("items") or {"type": "string"}, path + "[]")]
    if kind == "integer":
        return 0
    if kind == "number":
        return 0.0
    if kind == "boolean":
        return False
    return _string_for(path)


def _string_for(path: str) -> str:
    key = path.lower()
    if "истец" in key or "claimant" in key:
        return "ТОО «Алтын Курылыс», БИН 123456789012, г. Алматы, ул. Абая, 10"
    if "ответчик" in key or "defendant" in key:
        return "ТОО «Мега Строй», БИН 210987654321, г. Алматы, пр. Достык, 55"
    if "court" in key or "суд" in key:
        return "Специализированный межрайонный экономический суд города Алматы"
    if "price" in key or "цена" in key:
        return "4 500 000 тенге"
    return "По материалам дела: поставка по накладной № 44 от 20.01.2025 не оплачена."


class _Responses:
    def __init__(self, owner: "FakeProvider") -> None:
        self._owner = owner

    async def create(self, **kwargs: Any) -> _Response:
        self._owner.calls += 1
        if self._owner.fail_with is not None:
            raise self._owner.fail_with
        fmt = (kwargs.get("text") or {}).get("format") or {}
        schema = fmt.get("schema") or {"type": "object", "properties": {}}
        return _Response(json.dumps(_fill(schema), ensure_ascii=False))


class FakeProvider:
    def __init__(self) -> None:
        self.calls = 0
        self.fail_with: BaseException | None = None
        self.responses = _Responses(self)


def _service_objects() -> list[Any]:
    """Все объекты, у которых конвейер держит клиент провайдера.

    Слоёв сервиса несколько (`miniapp_api.service`, `miniapp_api_v2.service`,
    адаптер pipeline v2, мультиплексор исковых методов), и какой из них
    отвечает на конкретный тип документа, зависит от того, какие рантайм-слои
    успели установиться. Подменять клиента у одного объекта значит рисковать
    настоящим сетевым вызовом из теста, поэтому стаб ставится всюду, где
    клиент уже есть.
    """
    from korgan import miniapp_api as legacy
    from korgan import miniapp_api_v2 as core

    seen: list[Any] = []
    pending = [legacy.service, core.service]
    while pending:
        candidate = pending.pop()
        if candidate is None or any(candidate is item for item in seen):
            continue
        seen.append(candidate)
        for attr in ("inner", "_inner", "service", "_service", "stable", "_claim"):
            nested = getattr(candidate, attr, None)
            if nested is not None and not isinstance(nested, (str, bytes)):
                pending.append(nested)
    return [item for item in seen if hasattr(item, "client")]


@pytest.fixture()
def provider(monkeypatch: pytest.MonkeyPatch) -> FakeProvider:
    """Единственная внешняя зависимость конвейера — провайдер модели.

    Мультиплексор исковых методов создаёт свой сервис лениво, уже после начала
    теста, и вместе с ним — собственный клиент. Поэтому недостаточно подменить
    клиента у существующих объектов: подменяется и фабрика клиента, иначе
    очередной ленивый сервис ушёл бы в сеть по-настоящему.
    """
    from korgan import openai_legal

    fake = FakeProvider()
    monkeypatch.setattr(openai_legal, "build_legal_client", lambda settings: (fake, "fake"))
    for target in _service_objects():
        monkeypatch.setattr(target, "client", fake, raising=False)
    return fake


# --------------------------------------------------------------------------- #
# Минимальные двойники хранилищ: настоящие требуют Postgres.
# --------------------------------------------------------------------------- #
class FakeStore:
    def __init__(self, case_id: str = "case-1") -> None:
        self.state: dict[str, Any] = {"cases": {case_id: {"id": case_id}}}
        self.saves = 0
        self.fail_next_save = False

    def user_key(self, _identity: str) -> str:
        return "user-key"

    async def load(self, _identity: str) -> dict[str, Any]:
        return self.state

    async def save(self, _identity: str, state: dict[str, Any]) -> None:
        if self.fail_next_save:
            self.fail_next_save = False
            raise RuntimeError("state storage unavailable")
        self.saves += 1
        self.state = state


class FakePool:
    @asynccontextmanager
    async def acquire(self):
        yield self

    async def execute(self, *args):
        return "UPDATE 1"

    async def fetchrow(self, *args):
        return None


def _job(job_id: str = "job-1", *, order: int = 91, case_id: str = "case-1") -> jobs.GenerationJob:
    return jobs.GenerationJob(
        id=job_id,
        payment_order_id=order,
        user_key="user-key",
        case_id=case_id,
        status="queued",
        stage="queued",
        progress=0,
        error_detail="",
    )


class Recorder:
    """Переходы состояния задачи в том порядке, в каком их увидел бы клиент."""

    def __init__(self) -> None:
        self.transitions: list[tuple[str, str]] = []
        self.document_present_at: dict[str, bool] = {}


@pytest.fixture()
def wired(monkeypatch: pytest.MonkeyPatch):
    store = FakeStore()
    recorder = Recorder()
    claimed: set[str] = set()
    consumed: list[int] = []

    async def fake_update(job_id: str, *, status: str, stage: str, progress: int, error_detail: str = ""):
        recorder.transitions.append((status, stage))
        case = (store.state.get("cases") or {}).get("case-1") or {}
        recorder.document_present_at[status] = bool(case.get("document_base64"))

    async def fake_claim(job_id: str):
        if job_id in claimed:
            return None
        claimed.add(job_id)
        return _job(job_id)

    async def fake_consume(order_id: int, *, user_key: str) -> bool:
        consumed.append(order_id)
        return True

    monkeypatch.setattr(jobs, "_POOL", FakePool())
    monkeypatch.setattr(jobs, "claim_job", fake_claim)
    monkeypatch.setattr(jobs, "update_job", fake_update)
    monkeypatch.setattr(jobs.document_store, "consume_document_order", fake_consume)
    return store, recorder, consumed


def _document_from(store: FakeStore, case_id: str = "case-1") -> tuple[bytes, dict[str, Any]]:
    case = store.state["cases"][case_id]
    return base64.b64decode(case["document_base64"]), case


def test_end_to_end_generation_delivers_an_openable_docx(provider, wired) -> None:
    """Полный путь: запрос → конвейер → DOCX → сохранение → выдача клиенту."""
    store, recorder, consumed = wired

    asyncio.run(
        jobs.run_job(
            _job(),
            identity="identity",
            store=store,
            document_type="claim",
            context=CONTEXT,
            language="ru",
        )
    )

    assert provider.calls > 0, "юридический конвейер не обращался к модели"
    assert ("succeeded", "completed") in recorder.transitions

    data, case = _document_from(store)
    assert data, "документ пустой"
    assert len(data) > 5000, f"Word подозрительно мал: {len(data)} байт"

    document = Document(io.BytesIO(data))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert text.strip(), "Word открылся, но в нём нет текста"

    assert case["status"] == "document_ready"
    assert case["filename"].endswith(".docx")
    assert consumed == [91], "оплата не списана ровно один раз"


def test_ready_is_never_reported_before_the_document_exists(provider, wired) -> None:
    """READY разрешён только после того, как файл действительно сохранён."""
    store, recorder, _ = wired

    asyncio.run(
        jobs.run_job(
            _job(),
            identity="identity",
            store=store,
            document_type="claim",
            context=CONTEXT,
            language="ru",
        )
    )

    assert recorder.document_present_at.get("succeeded") is True
    for status, present in recorder.document_present_at.items():
        if status == "running":
            assert present is False, "клиент увидел бы готовность до сохранения файла"


def test_internal_service_text_never_reaches_the_word_file(provider, wired) -> None:
    """Служебные пометки конвейера не должны попадать в файл клиента."""
    store, _, _ = wired

    asyncio.run(
        jobs.run_job(
            _job(),
            identity="identity",
            store=store,
            document_type="claim",
            context=CONTEXT,
            language="ru",
        )
    )

    data, _case = _document_from(store)
    document = Document(io.BytesIO(data))
    body = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            body += "\n" + " ".join(cell.text for cell in row.cells)

    for leak in (
        "FILING_ACTION:",
        "SENIOR_PREFLIGHT_SCORE:",
        "LEGAL_GROUNDING:",
        "CLAIM_FAIL",
        "Traceback",
        "korgan_fast_professional",
    ):
        assert leak not in body, f"в Word попал служебный текст: {leak}"


def test_drafting_failure_does_not_poison_the_next_job(provider, wired) -> None:
    """Упавшая задача не должна ломать следующую генерацию."""
    store, recorder, _ = wired

    provider.fail_with = RuntimeError("Error code: 503 provider unavailable")
    with pytest.raises(RuntimeError):
        asyncio.run(
            jobs.run_job(
                _job("job-failed"),
                identity="identity",
                store=store,
                document_type="claim",
                context=CONTEXT,
                language="ru",
            )
        )
    assert ("failed", "failed") in recorder.transitions
    assert not (store.state["cases"]["case-1"].get("document_base64") or ""), "упавшая задача оставила документ"

    provider.fail_with = None
    asyncio.run(
        jobs.run_job(
            _job("job-next"),
            identity="identity",
            store=store,
            document_type="claim",
            context=CONTEXT,
            language="ru",
        )
    )

    data, _case = _document_from(store)
    assert Document(io.BytesIO(data)).paragraphs, "следующая задача не выдала документ"


def test_repeated_run_of_the_same_job_does_not_produce_a_second_document(provider, wired) -> None:
    """Повторный запуск той же задачи не создаёт второй документ и вторую оплату."""
    store, _recorder, consumed = wired
    job = _job()

    asyncio.run(
        jobs.run_job(job, identity="identity", store=store, document_type="claim", context=CONTEXT, language="ru")
    )
    first_saves = store.saves
    first_document, _ = _document_from(store)

    # Второй запуск проигрывает переход состояния в базе и обязан уйти молча.
    asyncio.run(
        jobs.run_job(job, identity="identity", store=store, document_type="claim", context=CONTEXT, language="ru")
    )

    assert store.saves == first_saves, "повторный запуск сохранил документ второй раз"
    assert consumed == [91], "повторный запуск списал оплату второй раз"
    again, _ = _document_from(store)
    assert again == first_document


def test_delivery_failure_is_not_hidden_behind_success(provider, wired) -> None:
    """Сбой сохранения — это отказ, а не готовый документ."""
    store, recorder, _ = wired
    store.fail_next_save = True

    with pytest.raises(RuntimeError):
        asyncio.run(
            jobs.run_job(
                _job(),
                identity="identity",
                store=store,
                document_type="claim",
                context=CONTEXT,
                language="ru",
            )
        )

    assert ("succeeded", "completed") not in recorder.transitions
    assert ("failed", "failed") in recorder.transitions


def test_generation_timeout_reports_a_reason_the_client_can_act_on(provider, wired, monkeypatch) -> None:
    """Превышение бюджета доходит до клиента причиной, а не общим отказом."""
    import korgan.document_latency_budget_runtime as budget

    store, _recorder, _ = wired
    monkeypatch.setattr(budget, "document_generation_timeout_seconds", lambda: 0.01)

    original = budget._ORIGINAL_GENERATE

    async def slow(document_type: str, context: str, language: str):
        await asyncio.sleep(5)
        return await original(document_type, context, language)

    monkeypatch.setattr(budget, "_ORIGINAL_GENERATE", slow)

    captured: list[str] = []

    async def capture_update(job_id: str, *, status: str, stage: str, progress: int, error_detail: str = ""):
        if error_detail:
            captured.append(error_detail)

    monkeypatch.setattr(jobs, "update_job", capture_update)

    with pytest.raises(budget.DocumentGenerationTimeout):
        asyncio.run(
            jobs.run_job(
                _job(),
                identity="identity",
                store=store,
                document_type="claim",
                context=CONTEXT,
                language="ru",
            )
        )

    assert captured, "клиенту не сказали ничего"
    assert "не уложился" in captured[-1]
    assert "новая оплата не потребуется" in captured[-1]

"""Golden regeneration: pinned clauses must survive, numbering must stay clean.

This is the regression net for the failure mode where a fix lands in one
document type and the same class of bug resurfaces in another, or where a clause
present in one version quietly disappears in the next.

Every pinned document type is regenerated through its real exporter and checked
for two things at once:

* every clause pinned in ``korgan.golden_documents`` is still there;
* no structural number is doubled and no narrative paragraph carries a number.
"""

from __future__ import annotations

import io
import re

import pytest
from docx import Document
from docx.oxml.ns import qn

from korgan.claim_docx import build_claim_docx
from korgan.contract_docx import build_contract_docx
from korgan.citation_audit import audit_citations
from korgan.document_release import review_document
from korgan.golden_documents import SPECS, golden_cases, missing_required_clauses
from korgan.text_integrity import integrity_findings
from korgan.response_docx import build_response_docx
from tests import golden_fixtures

CONTRACT_TYPES = sorted(key for key in SPECS if key.startswith("contract_"))
ALL_TYPES = ["claim", "response", *CONTRACT_TYPES]

_DOUBLED_NUMBER = re.compile(r"^\s*\d+(?:\.\d+)*\.\s+\d+(?:\.\d+)*\.")
_LEADING_NUMBER = re.compile(r"^\s*\d+(?:\.\d+)*\.\s")


def _build(document_type: str) -> bytes:
    if document_type == "claim":
        return build_claim_docx(golden_fixtures.claim_draft())
    if document_type == "response":
        return build_response_docx(golden_fixtures.response_draft())
    return build_contract_docx(golden_fixtures.contract_draft(document_type))


def _document(document_type: str) -> Document:
    return Document(io.BytesIO(_build(document_type)))


def _text(document_type: str) -> str:
    doc = _document(document_type)
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            lines.extend(cell.text for cell in row.cells)
    return "\n".join(lines)


def _is_listed(paragraph) -> bool:
    """True when Word will render an automatic number in front of the paragraph."""
    if paragraph._p.find(f'{qn("w:pPr")}/{qn("w:numPr")}') is not None:
        return True
    return (paragraph.style.name or "").startswith("List Number")


@pytest.mark.parametrize("document_type", ALL_TYPES)
def test_pinned_clauses_survive_regeneration(document_type: str) -> None:
    missing = missing_required_clauses(document_type, _text(document_type))
    assert not missing, f"{document_type}: пропали обязательные положения: {missing}"


@pytest.mark.parametrize("document_type", ALL_TYPES)
def test_no_structural_number_is_doubled(document_type: str) -> None:
    doubled = [line for line in _text(document_type).splitlines() if _DOUBLED_NUMBER.match(line)]
    assert not doubled, f"{document_type}: задвоенная нумерация: {doubled}"


@pytest.mark.parametrize("document_type", ALL_TYPES)
def test_narrative_paragraphs_never_join_a_numbered_list(document_type: str) -> None:
    """Only просительная часть and приложения are Word lists — nothing else."""
    doc = _document(document_type)
    listed = [p.text for p in doc.paragraphs if _is_listed(p) and p.text.strip()]

    body = "\n".join(p.text for p in doc.paragraphs)
    allowed: set[str] = set()
    for section in ("ПРОШУ СУД:", "Приложения:"):
        if section in body:
            tail = body.split(section, 1)[1]
            allowed.update(line.strip() for line in tail.splitlines() if line.strip())

    unexpected = [text for text in listed if text.strip() not in allowed]
    assert not unexpected, f"{document_type}: свободный текст попал в нумерованный список: {unexpected}"


def test_response_objections_are_numbered_but_their_prose_is_not() -> None:
    """The exact defect: narrative under an objection must stay unnumbered."""
    doc = _document("response")
    lines = [p.text for p in doc.paragraphs if p.text.strip()]

    assert "1. Размер задолженности определён Истцом без учёта фактически принятого объёма услуг." in lines
    assert "2. Истцом не соблюдён согласованный сторонами порядок приёмки услуг." in lines
    assert "2.1. акт направлен без приложения отчёта об оказанных услугах;" in lines

    for prose in lines:
        if prose.startswith("Согласно позиции Ответчика") or prose.startswith("С учётом принятых работ"):
            assert not _LEADING_NUMBER.match(prose), f"проза получила номер: {prose}"
            break
    else:  # pragma: no cover - fixture guarantees the paragraphs exist
        pytest.fail("в отзыве не найдены повествовательные абзацы возражений")

    numbered = [line for line in lines if _LEADING_NUMBER.match(line)]
    assert not any("Согласно позиции Ответчика" in line for line in numbered)
    assert not any("С учётом принятых работ" in line for line in numbered)


@pytest.mark.parametrize("document_type", ALL_TYPES)
def test_every_golden_document_passes_the_release_gate(document_type: str) -> None:
    """Citations and text integrity are checked for every type, not just one."""
    report = review_document(_text(document_type))
    assert report.released, f"{document_type}: релизный гейт заблокировал: {report.blocking}"


@pytest.mark.parametrize("document_type", ALL_TYPES)
def test_no_golden_document_carries_glued_or_truncated_text(document_type: str) -> None:
    findings = integrity_findings(_text(document_type))
    assert not findings, f"{document_type}: обрывы текста: {[f.as_note() for f in findings]}"


def test_regeneration_is_stable_across_repeated_runs() -> None:
    """Repeated regeneration must not introduce truncation artefacts."""
    for document_type in ALL_TYPES:
        renders = [_text(document_type) for _ in range(4)]
        assert len(set(renders)) == 1, f"{document_type}: регенерация нестабильна"
        assert not integrity_findings(renders[0])


def test_response_case_exercises_three_codes_and_judges_each_citation() -> None:
    audit = audit_citations(_text("response"))
    acts = {finding.act for finding in audit.findings}
    assert {"ГПК РК", "ГК РК", "НК РК"} <= acts
    # часть 2 статьи 166 ГПК РК is the amended article the corpus pins.
    part_two = next(f for f in audit.findings if f.act == "ГПК РК" and f.part == "2")
    assert part_two.quoted
    assert not part_two.blocks_release


def test_golden_registry_covers_every_pinned_type() -> None:
    cases = {case.document_type: case for case in golden_cases()}
    assert set(cases) == set(SPECS)
    response = cases["response"]
    assert response.cites_law
    # Fails loudly once a date is set but goes stale is a manual QA concern; here
    # we only assert the registry records the obligation at all.
    assert "ГПК РК" in response.cited_acts


def test_it_contract_keeps_the_personal_data_clause() -> None:
    """The clause that vanished between v3 and v4 is pinned and asserted."""
    text = _text("contract_it_services").lower()
    assert "персональн" in text
    assert "обработка персональных данных" in text
    assert not missing_required_clauses("contract_it_services", _text("contract_it_services"))


def test_dropping_a_pinned_clause_fails_the_golden_check() -> None:
    """The net has to actually catch a removal, not just pass on good input."""
    payload = golden_fixtures.contract_payload("contract_it_services")
    payload["sections"] = [
        section
        for section in payload["sections"]
        if "персональных данных" not in section["heading"].lower()
    ]

    from korgan.legal_types import ContractDraft, VerificationStatus

    draft = ContractDraft.from_payload(
        status=VerificationStatus.NEEDS_VERIFICATION,
        source_urls=[],
        payload=payload,
    )
    doc = Document(io.BytesIO(build_contract_docx(draft)))
    text = "\n".join(p.text for p in doc.paragraphs)

    missing = missing_required_clauses("contract_it_services", text)
    assert "обработка персональных данных пользователей" in missing

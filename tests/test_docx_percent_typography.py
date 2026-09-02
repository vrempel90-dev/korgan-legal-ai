from __future__ import annotations

import io

from docx import Document

from korgan.claim_docx import build_claim_docx
from korgan.legal_types import ClaimDraft, VerificationStatus


def _docx_text(payload: bytes) -> str:
    document = Document(io.BytesIO(payload))
    lines = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                lines.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(lines)


def _draft() -> ClaimDraft:
    return ClaimDraft(
        status=VerificationStatus.VERIFIED,
        title="ИСКОВОЕ ЗАЯВЛЕНИЕ",
        court="Специализированный межрайонный экономический суд",
        claimant=["ТОО «Истец», БИН 010140001230, г. Алматы"],
        defendant=["ТОО «Ответчик», БИН 020240005675, г. Алматы"],
        price_of_claim="1 100 000 тенге",
        facts=["Договором установлена неустойка 0.1% за каждый день просрочки."],
        legal_basis=["Требование основано на условиях договора."],
        requests=["Взыскать договорную неустойку по ставке 0.1% в день."],
        attachments=["Расчёт неустойки по ставке 0.01%."],
        verification_notes=[],
        source_urls=[],
        calculation=[
            "Договорная неустойка: база 1 000 000 тенге × 0.1% × 10 дн. = 10 000 тенге."
        ],
        late_interest="Ставка 0.01% за каждый день просрочки.",
    )


def test_docx_uses_decimal_comma_for_percentage_rates() -> None:
    text = _docx_text(build_claim_docx(_draft()))

    assert "0,1%" in text
    assert "0,01%" in text
    assert "0.1%" not in text
    assert "0.01%" not in text


def test_percentage_typography_does_not_change_dates_or_amounts() -> None:
    draft = _draft()
    draft.facts.append("Период: 01.03.2026—10.03.2026; сумма 1 000 000 тенге.")

    text = _docx_text(build_claim_docx(draft))

    assert "01.03.2026—10.03.2026" in text
    assert "1 000 000 тенге" in text

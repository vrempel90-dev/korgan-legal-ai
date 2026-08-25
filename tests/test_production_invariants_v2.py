from __future__ import annotations

import io
from datetime import date

from docx import Document

from korgan.legal_types import LegalResearch, VerificationStatus
from korgan.production_invariants_v2 import (
    BlockerClass,
    _repair_signature,
    append_review_markers,
    build_money_ledger,
    calc_contractual_penalty,
    canonicalize_research,
    classify_issue,
    review_marker,
)


MONEY_FIXTURE = (
    "Основной долг 3 250 000 тенге. "
    "По договору неустойка 0,2% в день, потолок 20%. "
    "Просрочка с 08.03.2026."
)


def test_i3_internal_citation_defect_is_not_user_blocker() -> None:
    issue = (
        "статья 469 ГК РК: пересказ обобщает узкое условие нормы: "
        "норма формулирует право, а не обязанность"
    )
    classified = classify_issue(issue)
    assert classified.blocker_class == BlockerClass.INTERNAL_QUALITY
    assert "KORGAN" in classified.action


def test_i3_missing_sender_is_user_resolvable() -> None:
    classified = classify_issue("не указан отправитель претензии")
    assert classified.blocker_class == BlockerClass.NEEDS_USER_DATA
    assert "отправителя" in classified.action


def test_i4_marker_keeps_exact_article_diagnosis() -> None:
    issue = "статья 469 ГК РК: пересказ обобщает узкое условие нормы"
    marker = review_marker(issue)
    assert marker == "[СВЕРИТЬ: статья 469 ГК РК: пересказ обобщает узкое условие нормы]"


def test_i2_internal_warning_is_written_into_docx() -> None:
    source = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Тестовый документ")
    doc.save(source)

    output = append_review_markers(
        source.getvalue(),
        ["статья 469 ГК РК: пересказ требует проверки"],
    )
    rendered = Document(io.BytesIO(output))
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "[СВЕРИТЬ: статья 469 ГК РК: пересказ требует проверки]" in text


def test_i5_contractual_penalty_fixture_reaches_cap_on_15_june() -> None:
    penalty = calc_contractual_penalty(
        3_250_000,
        "0.2",
        date(2026, 3, 8),
        date(2026, 8, 25),
        cap_percent="20",
    )
    assert penalty.daily_amount == 6_500
    assert penalty.days == 171
    assert penalty.cap_amount == 650_000
    assert penalty.cap_reached_on == date(2026, 6, 15)
    assert penalty.amount == 650_000


def test_i5_money_ledger_never_zero_when_money_is_in_input() -> None:
    ledger = build_money_ledger(MONEY_FIXTURE, as_of=date(2026, 8, 25))
    assert ledger.input_amounts == (3_250_000,)
    assert ledger.principal == 3_250_000
    assert ledger.penalty is not None
    assert ledger.penalty.amount == 650_000
    assert ledger.total == 3_900_000
    assert ledger.total > 0


def test_i7_repair_blocker_signature_ignores_order_and_duplicates() -> None:
    left = _repair_signature(["статья 469 требует проверки", "неверный пересказ", "статья 469 требует проверки"])
    right = _repair_signature(["неверный пересказ", "статья 469 требует проверки"])
    assert left == right


def test_i9_research_output_is_canonical_and_reproducible() -> None:
    research = LegalResearch(
        status=VerificationStatus.NEEDS_VERIFICATION,
        applicable_law=["ГК РК", "ГПК РК", "ГК РК"],
        procedural_requirements=["B", "A"],
        verified_claims=[
            "Вывод по статье 469 ГК РК [основание: статья 469 ГК РК; источник: https://example/469]",
            "Вывод по статье 439 ГК РК [основание: статья 439 ГК РК; источник: https://example/439]",
        ],
        unverified_claims=["Z", "A", "Z"],
        source_urls=["https://example/z", "https://example/a", "https://example/z"],
        notes=["REMEDY: B", "REMEDY: A"],
    )
    first = canonicalize_research(research)
    snapshot = (
        tuple(first.verified_claims),
        tuple(first.unverified_claims),
        tuple(first.source_urls),
        tuple(first.applicable_law),
        tuple(first.procedural_requirements),
        tuple(first.notes),
    )
    second = canonicalize_research(first)
    assert snapshot == (
        tuple(second.verified_claims),
        tuple(second.unverified_claims),
        tuple(second.source_urls),
        tuple(second.applicable_law),
        tuple(second.procedural_requirements),
        tuple(second.notes),
    )
    assert "439" in second.verified_claims[0]
    assert "469" in second.verified_claims[1]


def test_research_invariant_keeps_needs_verification_when_unverified_dominates() -> None:
    research = LegalResearch(
        status=VerificationStatus.VERIFIED,
        applicable_law=[],
        procedural_requirements=[],
        verified_claims=["одна подтвержденная норма"],
        unverified_claims=["риск 1", "риск 2"],
        source_urls=["https://example/source"],
        notes=[],
    )
    canonicalize_research(research)
    assert research.status == VerificationStatus.NEEDS_VERIFICATION

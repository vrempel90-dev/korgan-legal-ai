from __future__ import annotations

import inspect
import io

from docx import Document

from korgan.claim_docx import build_claim_docx
from korgan.claim_intent import is_claim_drafting_request
from korgan.i18n import KK, button, tr
from korgan.kazakh_article_forms import install_kazakh_article_forms
from korgan.kazakh_legal_bridge import install_kazakh_legal_bridge
from korgan.language_context import _CURRENT_LANGUAGE
from korgan.legal_safety import confirm_claim
from korgan.legal_types import ClaimDraft, VerificationStatus
from korgan.ui import main_menu


def _kk_draft() -> ClaimDraft:
    return ClaimDraft(
        status=VerificationStatus.VERIFIED,
        title="Қарыз сомасын өндіріп алу туралы ТАЛАП ҚОЮ АРЫЗЫ",
        court="Алмалы ауданының аудандық соты",
        claimant=["Иванов Иван Иванович, ЖСН 900101300001, Алматы қ., Абай көш., 10"],
        defendant=["Петров Петр Петрович, ЖСН 900101300002, Алматы қ., Төле би көш., 20"],
        price_of_claim="1 000 000 теңге",
        state_duty="10 000 тенге (1% от цены иска, статья 665 Налогового кодекса РК (Кодекс РК № 214-VIII))",
        facts=[
            "Талап қоюшы жауапкерге 1 000 000 теңге қарыз берген.",
            "Қарыз беру қолхатпен расталады.",
            "Қайтару мерзімі өтіп кеткен, ақша қайтарылмаған.",
        ],
        legal_basis=["ҚР АК 722-бабы бойынша қарыз алушы қарыз сомасын белгіленген мерзімде қайтаруға міндетті."],
        requests=["Жауапкерден талап қоюшының пайдасына 1 000 000 теңге қарыз сомасын өндіріп алу."],
        attachments=["Қолхат көшірмесі", "Банк құжаты"],
        verification_notes=[],
        source_urls=["https://adilet.zan.kz/rus/docs/K990000409_"],
    )


def test_kazakh_main_menu_is_fully_kazakh() -> None:
    labels = [item.text for row in main_menu(KK).keyboard for item in row]
    assert button(KK, "consultation") in labels
    assert button(KK, "document") in labels
    assert button(KK, "language") in labels
    assert "⚖️ Консультация" not in labels
    assert "NEEDS_VERIFICATION" not in tr(KK, "welcome")


def test_kazakh_claim_intent_routes_to_document_generation() -> None:
    for text in (
        "Маған қарызды өндіру туралы талап қою арызын дайында",
        "Талап арыз жасап бер",
        "Талап қою арызын құрастыр",
    ):
        assert is_claim_drafting_request(text), text


def test_kazakh_advice_question_is_not_forced_to_claim() -> None:
    assert not is_claim_drafting_request("Талап қою арызын қалай дайындауға болады?")


def test_kazakh_article_morphology_is_verified_by_same_citation_parser() -> None:
    install_kazakh_legal_bridge()
    install_kazakh_article_forms()
    from korgan.citation_audit import extract_references
    from korgan import document_quality

    refs = extract_references("ҚР Азаматтық кодексінің 722-бабы бойынша қарыз қайтарылуға тиіс.")
    assert any(ref.act == "ГК РК" and ref.article == "722" for ref in refs)
    assert document_quality._ARTICLE_RE.search("ҚР АК 722-бабы")
    assert document_quality._ARTICLE_RE.search("ҚР АК 722-бап")


def test_kazakh_claim_docx_has_kazakh_filing_labels_and_attachment_restart() -> None:
    token = _CURRENT_LANGUAGE.set(KK)
    try:
        data = build_claim_docx(_kk_draft())
    finally:
        _CURRENT_LANGUAGE.reset(token)

    document = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Талап қоюшы:" in text
    assert "Жауапкер:" in text
    assert "Талап қою бағасы:" in text
    assert "Мемлекеттік баж:" in text
    assert "Құқықтық негіздеме" in text
    assert "СОТТАН СҰРАЙМЫН" in text
    assert "Қосымшалар:" in text
    assert "Күні:" in text
    assert "Қолы:" in text
    assert "Цена иска:" not in text
    assert "Госпошлина:" not in text

    # Word numbering is represented in OOXML, not paragraph.text. The
    # attachment list must have its own numbering id so it restarts at 1.
    attachments_index = next(i for i, p in enumerate(document.paragraphs) if p.text.strip() == "Қосымшалар:")
    first_attachment = document.paragraphs[attachments_index + 1]
    num_pr = first_attachment._p.pPr.numPr
    assert num_pr is not None
    assert num_pr.numId is not None


def test_claim_confirmation_never_calls_legacy_questionnaire_handler() -> None:
    source = inspect.getsource(confirm_claim)
    assert "_generate_now" in source
    assert "base_bot.claim_handler" not in source

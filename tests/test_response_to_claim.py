from __future__ import annotations

from io import BytesIO

from docx import Document

from korgan.legal_types import VerificationStatus
from korgan.response_docx import build_response_to_claim_docx
from korgan.response_intent import is_response_to_claim_request
from korgan.response_types import ResponseObjection, ResponseToClaimDraft


def test_response_intent_does_not_steal_feedback() -> None:
    assert is_response_to_claim_request("Подготовь отзыв на иск")
    assert is_response_to_claim_request("Нужны возражения на исковое заявление")
    assert not is_response_to_claim_request("Хочу оставить отзыв о KORGAN")
    assert not is_response_to_claim_request("Подготовь иск о взыскании долга")


def test_response_docx_contains_court_structure_without_internal_qa() -> None:
    draft = ResponseToClaimDraft(
        status=VerificationStatus.VERIFIED,
        court="Бостандыкский районный суд города Алматы",
        case_number="1234-26-00-2/999",
        claimant=["Иванов Иван Иванович, г. Алматы"],
        defendant=["Петров Петр Петрович, г. Алматы"],
        claim_summary=["Истец просит взыскать 500 000 тенге."],
        position=["Ответчик иск не признает."],
        objections=["Требование о взыскании суммы оспаривается по основаниям, изложенным ниже."],
        legal_basis=["Отзыв подается в порядке статьи 166 ГПК Республики Казахстан."],
        requests=["Отказать в удовлетворении иска."],
        attachments=["Копия документа, подтверждающего доводы ответчика."],
    )

    payload = build_response_to_claim_docx(draft)
    doc = Document(BytesIO(payload))
    text = "\n".join(p.text for p in doc.paragraphs)

    assert "ОТЗЫВ НА ИСК" in text
    assert "статьи 166" in text
    assert "ПРОШУ СУД" in text
    assert "KORGAN QA STATUS" not in text
    assert "https://" not in text


def test_response_draft_normalizes_raw_objection_dicts() -> None:
    draft = ResponseToClaimDraft(
        status=VerificationStatus.VERIFIED,
        objections=[
            {
                "text": "1. Основное возражение",
                "subclauses": ["1.1. Частный довод"],
                "prose": ["Дополнительное пояснение."],
            }
        ],
    )

    assert len(draft.objections) == 1
    objection = draft.objections[0]
    assert isinstance(objection, ResponseObjection)
    assert objection.text == "Основное возражение"
    assert objection.subclauses == ["Частный довод"]
    assert objection.prose == ["Дополнительное пояснение."]
    assert objection.body_lines() == [
        "Основное возражение",
        "Частный довод",
        "Дополнительное пояснение.",
    ]


def test_response_docx_uses_a4() -> None:
    draft = ResponseToClaimDraft(status=VerificationStatus.NEEDS_VERIFICATION)
    doc = Document(BytesIO(build_response_to_claim_docx(draft)))
    section = doc.sections[0]
    assert abs(section.page_width.mm - 210) < 0.2
    assert abs(section.page_height.mm - 297) < 0.2

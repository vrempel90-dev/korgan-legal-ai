"""Нормы корпуса попадают в документ без повторного даунгрейда в NEEDS_VERIFICATION."""

import io
from datetime import date

from docx import Document

from korgan.claim_docx import build_claim_docx
from korgan.legal_calc import NEEDS_RATE_MARKER
from korgan.legal_corpus import CORPUS, GENERAL_TERRITORIAL_JURISDICTION, LATE_PAYMENT_INTEREST
from korgan.legal_types import ClaimDraft, VerificationStatus
from korgan.production_legal import (
    COURT_PLACEHOLDER,
    GESV_NOTE,
    INTEREST_DATE_NOTE,
    INTEREST_RATE_NOTE,
    _apply_individual_loan_check,
    _apply_jurisdiction_provision,
    _apply_late_interest,
    _corpus_pending_note,
)

# Дело № 3: расписка на 800 000 тенге между физлицами, срок возврата 10.04.2026.
DELO_3_CONTEXT = (
    "Стороны: Займодавец (истец): Ахметов Руслан Маратович; Заёмщик (ответчик): Садыков Тимур Ерланович\n"
    "Идентификаторы: ИИН 000000000101; ИИН 000000000202\n"
    "Суммы: 800 000 тенге — сумма займа по расписке\n"
    "Даты: 10.04.2026 — срок возврата\n"
    "Доказательства: расписка от 10.01.2026\n"
)

FILING_DATE = date(2026, 8, 15)


def _draft(**overrides) -> ClaimDraft:
    base = dict(
        status=VerificationStatus.NEEDS_VERIFICATION,
        title="ИСКОВОЕ ЗАЯВЛЕНИЕ о взыскании долга по расписке",
        court=COURT_PLACEHOLDER,
        claimant=["Ахметов Руслан Маратович, ИИН 000000000101"],
        defendant=["Садыков Тимур Ерланович, ИИН 000000000202"],
        price_of_claim="800 000 тенге",
        facts=["10.01.2026 ответчик получил в долг 800 000 тенге, что подтверждается распиской"],
        legal_basis=["В соответствии со статьёй 715 ГК РК заёмщик обязан возвратить предмет займа."],
        requests=["Взыскать основной долг в размере 800 000 тенге"],
        attachments=["Копия расписки от 10.01.2026"],
        verification_notes=[],
        source_urls=[],
    )
    base.update(overrides)
    return ClaimDraft(**base)


def test_interest_provision_and_amount_reach_the_draft() -> None:
    draft = _draft()

    _apply_late_interest(draft, "10.04.2026", True, FILING_DATE)

    assert any("353" in line for line in draft.legal_basis)
    assert "50 104 тенге" in draft.late_interest
    assert "18%" in draft.late_interest


def test_interest_request_states_amount_and_continuing_accrual() -> None:
    draft = _draft()

    _apply_late_interest(draft, "10.04.2026", True, FILING_DATE)

    interest_request = [line for line in draft.requests if "неустойку" in line]
    assert len(interest_request) == 1
    assert "50 104 тенге" in interest_request[0]
    assert "по день фактического исполнения" in interest_request[0]


def test_no_interest_claimed_leaves_draft_untouched() -> None:
    draft = _draft()

    _apply_late_interest(draft, "10.04.2026", False, FILING_DATE)

    assert draft.late_interest == ""
    assert not any("353" in line for line in draft.legal_basis)
    assert len(draft.requests) == 1


def test_unknown_due_date_keeps_the_provision_and_flags_the_date() -> None:
    draft = _draft()

    _apply_late_interest(draft, "", True, FILING_DATE)

    assert any("353" in line for line in draft.legal_basis)
    assert INTEREST_DATE_NOTE in draft.verification_notes


def test_unknown_rate_flags_the_rate_not_the_provision() -> None:
    """Просрочка началась до появления ставки в справочнике."""
    draft = _draft()

    _apply_late_interest(draft, "01.02.2025", True, date(2025, 6, 1))

    assert any("353" in line for line in draft.legal_basis)
    assert draft.late_interest == NEEDS_RATE_MARKER
    assert INTEREST_RATE_NOTE in draft.verification_notes


def test_jurisdiction_article_is_added_once() -> None:
    draft = _draft()

    _apply_jurisdiction_provision(draft)
    _apply_jurisdiction_provision(draft)

    jurisdiction_lines = [line for line in draft.legal_basis if "29" in line]
    assert len(jurisdiction_lines) == 1
    assert "месту жительства ответчика" in jurisdiction_lines[0]


def test_individual_loan_raises_gesv_check() -> None:
    draft = _draft()

    _apply_individual_loan_check(DELO_3_CONTEXT, draft)

    assert GESV_NOTE in draft.verification_notes
    # Норма может обессилить договор истца — в тело иска она не попадает.
    assert not any("725-1" in line for line in draft.legal_basis)


def test_gesv_check_is_skipped_when_materials_state_the_rate() -> None:
    draft = _draft()

    _apply_individual_loan_check(DELO_3_CONTEXT + "Важные факты: ГЭСВ 0 процентов\n", draft)

    assert GESV_NOTE not in draft.verification_notes


def test_gesv_check_is_skipped_for_legal_entity_lender() -> None:
    draft = _draft()

    _apply_individual_loan_check(DELO_3_CONTEXT + "Идентификаторы: ТОО «Альфа», БИН 000000000303\n", draft)

    assert GESV_NOTE not in draft.verification_notes


def test_corpus_provisions_carry_a_pending_official_check_note() -> None:
    draft = _draft()
    _apply_late_interest(draft, "10.04.2026", True, FILING_DATE)
    _apply_jurisdiction_provision(draft)

    _corpus_pending_note(draft)

    notes = " ".join(draft.verification_notes)
    assert "ст. 353" in notes and "ст. 29" in notes


def test_corpus_entries_expose_article_and_source() -> None:
    for provision in CORPUS.values():
        assert provision.article
        assert provision.source_url.startswith("https://adilet.zan.kz/")
        assert provision.as_verified_claim().startswith(provision.summary[:20])

    assert LATE_PAYMENT_INTEREST.article == "353"
    assert GENERAL_TERRITORIAL_JURISDICTION.article == "29"


def test_docx_renders_the_interest_calculation() -> None:
    draft = _draft()
    _apply_late_interest(draft, "10.04.2026", True, FILING_DATE)

    text = "\n".join(p.text for p in Document(io.BytesIO(build_claim_docx(draft))).paragraphs)

    assert "Расчёт неустойки за просрочку" in text
    assert "50 104 тенге" in text


def test_docx_omits_the_interest_block_when_not_claimed() -> None:
    text = "\n".join(p.text for p in Document(io.BytesIO(build_claim_docx(_draft()))).paragraphs)

    assert "Расчёт неустойки за просрочку" not in text


def test_consultation_citations_are_extracted() -> None:
    from korgan.legal_corpus import extract_cited_articles

    answer = (
        "По вашему делу применяются статьи 715 и 722 ГК РК. "
        "За просрочку начисляется неустойка по ст. 353 ГК РК. "
        "Иск подаётся по ст. 29 ГПК РК."
    )

    assert extract_cited_articles(answer) == ["ст. 715 ГК", "ст. 722 ГК", "ст. 353 ГК", "ст. 29 ГПК"]


def test_answer_without_citations_yields_nothing() -> None:
    from korgan.legal_corpus import extract_cited_articles

    assert extract_cited_articles("Норму нельзя назвать без официального источника.") == []

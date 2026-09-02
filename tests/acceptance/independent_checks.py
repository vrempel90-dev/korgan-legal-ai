"""Независимая проверка готового DOCX. Ничего из кода генерации не использует.

Зачем модуль написан заново
---------------------------
Проверять документ теми же функциями, которыми он собран, значит сверять
реализацию саму с собой: общая ошибка в разборе сумм или в формуле неустойки
одинаково исказит и документ, и проверку, и обе сойдутся. Поэтому здесь всё
своё — свой разбор сумм, свой разбор ссылок на нормы, свой расчёт неустойки по
интервалам. Совпадение результатов двух независимо написанных реализаций и есть
то, что подтверждает число.

Единственное, что модуль берёт извне, — сам файл документа и заявленные
исходные данные дела. Ни ``legal_calc``, ни ``penalty_engine``, ни
``document_linter`` он не импортирует.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from datetime import date, timedelta
from decimal import Decimal, ROUND_HALF_UP

from docx import Document

# --------------------------------------------------------------------------
# Свой разбор документа
# --------------------------------------------------------------------------

# Сумма в тенге. Разряды разделены пробелом либо не разделены вовсе.
_AMOUNT_RE = re.compile(
    r"(?<!\d)(?<!\d )(\d{1,3}(?:[  ]\d{3})+|\d+)\s*(?:тенге|теңге|₸)",
    re.IGNORECASE,
)

# Ссылка на норму, включая перечисление: «ст. 178, 180 и 183 ГК РК».
_CITATION_RE = re.compile(
    r"(?:стать[а-яё]*|ст\.)\s*(?P<numbers>\d+(?:\s*(?:,|и)\s*\d+)*)\s*"
    r"(?P<act>ГК\s*РК|ГПК\s*РК|НК\s*РК|ТК\s*РК|ЗПП\s*РК|"
    r"Гражданского\s+кодекса[^,.;]{0,40}|"
    r"Гражданского\s+процессуального\s+кодекса[^,.;]{0,40}|"
    r"Налогового\s+кодекса[^,.;]{0,40}|"
    r"Трудового\s+кодекса[^,.;]{0,40})",
    re.IGNORECASE,
)
_NUMBER_RE = re.compile(r"\d+")

_ACT_CANONICAL: tuple[tuple[re.Pattern[str], str], ...] = (
    (re.compile(r"ГПК\s*РК|Гражданского\s+процессуального", re.IGNORECASE), "ГПК РК"),
    (re.compile(r"ГК\s*РК|Гражданского\s+кодекса", re.IGNORECASE), "ГК РК"),
    (re.compile(r"НК\s*РК|Налогового\s+кодекса", re.IGNORECASE), "НК РК"),
    (re.compile(r"ТК\s*РК|Трудового\s+кодекса", re.IGNORECASE), "ТК РК"),
    (re.compile(r"ЗПП\s*РК|защите\s+прав\s+потребител", re.IGNORECASE), "ЗПП РК"),
)


def read_paragraphs(payload: bytes) -> list[str]:
    """Все абзацы документа, включая таблицы."""
    document = Document(io.BytesIO(payload))
    lines = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                lines.extend(paragraph.text for paragraph in cell.paragraphs)
    return [line for line in lines if line and line.strip()]


def document_text(payload: bytes) -> str:
    return "\n".join(read_paragraphs(payload))


def amounts_in(text: str) -> list[int]:
    """Все суммы в тенге, в порядке появления."""
    values: list[int] = []
    for match in _AMOUNT_RE.finditer(text or ""):
        digits = re.sub(r"[  ]", "", match.group(1))
        if digits.isdigit():
            values.append(int(digits))
    return values


def citations_in(text: str) -> list[tuple[str, str]]:
    """Все упоминания норм как пары «акт, номер статьи»."""
    found: list[tuple[str, str]] = []
    for match in _CITATION_RE.finditer(text or ""):
        act_text = match.group("act")
        act = ""
        for pattern, canonical in _ACT_CANONICAL:
            if pattern.search(act_text):
                act = canonical
                break
        if not act:
            continue
        for number in _NUMBER_RE.findall(match.group("numbers")):
            pair = (act, number)
            if pair not in found:
                found.append(pair)
    return found


# --------------------------------------------------------------------------
# Свой расчёт неустойки по интервалам
# --------------------------------------------------------------------------


@dataclass(frozen=True, slots=True)
class Payment:
    on: date
    amount: int


@dataclass(frozen=True, slots=True)
class PenaltyExpectation:
    """Ожидаемый расчёт, выведенный независимо от кода продукта."""

    intervals: tuple[tuple[date, date, int, int, int], ...]
    raw_total: int
    total: int
    capped: bool


def expected_penalty(
    *,
    contract_value: int,
    payments: tuple[Payment, ...],
    rate_percent_per_day: str,
    start: date,
    end: date,
    cap_amount: int | None = None,
) -> PenaltyExpectation:
    """Посчитать неустойку по интервалам постоянного остатка долга.

    Реализация намеренно прямолинейная: остаток пересчитывается на каждую дату
    платежа, каждый отрезок считается отдельно, подытоги округляются построчно.
    Совпадение с продуктовым расчётом что-то значит именно потому, что здесь
    ничего от него не заимствовано.
    """
    rate = Decimal(rate_percent_per_day) / Decimal(100)

    borders = sorted({payment.on for payment in payments if start <= payment.on <= end})
    intervals: list[tuple[date, date, int, int, int]] = []

    cursor = start
    balance = contract_value - sum(p.amount for p in payments if p.on < start)
    for border in borders:
        if border > cursor:
            days = (border - timedelta(days=1) - cursor).days + 1
            subtotal = int(
                (Decimal(balance) * rate * Decimal(days)).to_integral_value(rounding=ROUND_HALF_UP)
            )
            intervals.append((cursor, border - timedelta(days=1), days, balance, subtotal))
        balance -= sum(p.amount for p in payments if p.on == border)
        cursor = border

    if cursor <= end:
        days = (end - cursor).days + 1
        subtotal = int(
            (Decimal(balance) * rate * Decimal(days)).to_integral_value(rounding=ROUND_HALF_UP)
        )
        intervals.append((cursor, end, days, balance, subtotal))

    raw_total = sum(item[4] for item in intervals)
    total = raw_total
    capped = False
    if cap_amount is not None and raw_total > cap_amount:
        total = cap_amount
        capped = True

    return PenaltyExpectation(
        intervals=tuple(intervals), raw_total=raw_total, total=total, capped=capped
    )


def expected_state_duty(claim_price: int, *, legal_entity: bool) -> int:
    """Госпошлина по статье 665 НК РК: 3% для юрлица, 1% для физлица."""
    rate = Decimal("0.03") if legal_entity else Decimal("0.01")
    return int((Decimal(claim_price) * rate).to_integral_value(rounding=ROUND_HALF_UP))


# --------------------------------------------------------------------------
# Проверки чистоты и структуры
# --------------------------------------------------------------------------

#: Служебные следы, которых в судебном документе быть не должно. Список написан
#: заново, а не импортирован из линтера: иначе проверка молчала бы ровно о тех
#: маркерах, которые линтер забыл перечислить.
FORBIDDEN_TRACES: tuple[str, ...] = (
    "verification_notes",
    "critical_errors",
    "unsupported_legal_claims",
    "missing_required_fields",
    "TODO",
    "FIXME",
    "DEBUG",
    "Traceback",
    "NEEDS_VERIFICATION",
    "PRELIMINARY",
    "KORGAN QUALITY",
    "SENIOR_PREFLIGHT_SCORE",
    "CLAIM_PIPELINE",
    "Детерминированный расчёт:",
    "Ссылка на норму:",
    "{{",
    "}}",
    "[ТРЕБУЕТ РАСЧ",
    "placeholder",
    "principal_amount",
    "penalty_amount",
    "claim_price",
    "state_duty",
    "total_claim",
)

#: Единственный разрешённый служебный штамп — видимый QA-статус в шапке.
ALLOWED_PREFIX = "KORGAN QA STATUS"


@dataclass
class CleanlinessReport:
    traces: list[tuple[str, str]] = field(default_factory=list)

    @property
    def clean(self) -> bool:
        return not self.traces


def check_cleanliness(payload: bytes) -> CleanlinessReport:
    report = CleanlinessReport()
    for line in read_paragraphs(payload):
        if line.strip().startswith(ALLOWED_PREFIX):
            continue
        for trace in FORBIDDEN_TRACES:
            if trace in line:
                report.traces.append((trace, line.strip()[:120]))
    return report


_MOTION_DEMAND_RE = re.compile(r"истреб\w*|запросить\s+у\s+истца", re.IGNORECASE)
_FROM_CLAIMANT_RE = re.compile(r"у\s+истца|от\s+истца", re.IGNORECASE)
_WORD_RE = re.compile(r"[а-яёa-z0-9]{4,}", re.IGNORECASE)
_STOPWORDS = frozenset({"копия", "копии", "документ", "документы", "истца", "истец",
                        "ответчика", "приложение", "оригинал", "также"})


def contradictory_motions(payload: bytes) -> list[str]:
    """Ходатайства, просящие суд истребовать у истца его же приложение."""
    lines = read_paragraphs(payload)
    # Приложения идут после своего заголовка и не нумеруются экспортёром, так
    # что опознаются по позиции, а не по номеру строки.
    attachments: list[str] = []
    collecting = False
    for line in lines:
        stripped = line.strip()
        if re.match(r"^(?:Приложения|Қосымшалар)\s*:?\s*$", stripped, re.IGNORECASE):
            collecting = True
            continue
        if collecting:
            if not stripped or re.match(r"^[А-ЯЁA-Z][^а-яё]{0,40}:$", stripped):
                collecting = False
                continue
            attachments.append(stripped)
    problems: list[str] = []
    for line in lines:
        if not (_MOTION_DEMAND_RE.search(line) and _FROM_CLAIMANT_RE.search(line)):
            continue
        motion_words = {w.lower() for w in _WORD_RE.findall(line)} - _STOPWORDS
        for attachment in attachments:
            attachment_words = {w.lower() for w in _WORD_RE.findall(attachment)} - _STOPWORDS
            if len(attachment_words) < 2:
                continue
            overlap = motion_words & attachment_words
            if len(overlap) >= 2 and len(overlap) / len(attachment_words) >= 0.5:
                problems.append(line.strip())
                break
    return problems

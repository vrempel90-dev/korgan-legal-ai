from __future__ import annotations

import asyncio
import io
from types import SimpleNamespace

import pytest
from docx import Document


def _docx_bytes(*paragraphs: str) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _meta() -> dict[str, object]:
    return {
        "filing_ready": True,
        "release_status": "verified",
        "quality_score": 10.0,
        "quality_issues": [],
        "verification_notes": [],
    }


@pytest.mark.parametrize(
    "document_type",
    ["claim", "response", "pretrial", "pretrial_response", "contract"],
)
def test_live_quote_mismatch_never_discards_generated_word(
    monkeypatch: pytest.MonkeyPatch,
    document_type: str,
) -> None:
    """Regression for the production 95% dead end seen after payment.

    A strict live mismatch must still be detected, but it may no longer turn an
    already generated DOCX into HTTP 422. The unsafe cited proposition is
    removed, the user's facts and requested relief survive, and the result is
    returned as a review draft for every document route using core._generate.
    """
    import korgan.live_article_release_runtime as runtime

    async def scenario() -> None:
        original_word = _docx_bytes(
            "Факты дела: ответчик получил товар и не оплатил его в срок.",
            (
                "Согласно статье 630 ГК РК: «Ответчик обязан уплатить штраф "
                "в размере пятидесяти процентов.»"
            ),
            "ПРОШУ: взыскать подтвержденную материалами дела задолженность.",
        )
        draft = SimpleNamespace(status=runtime.VerificationStatus.VERIFIED)

        async def original_generate(kind: str, context: str, language: str):
            assert kind == document_type
            return draft, original_word, f"{kind}.docx", _meta()

        async def strict_verifier(payload: bytes) -> None:
            text = runtime._docx_text(payload)
            if runtime.citation_audit.extract_references(text):
                raise runtime.LiveArticleVerificationError(
                    "дословная цитата статья 630 ГК РК не совпадает с живым текстом Adilet"
                )

        monkeypatch.setattr(runtime, "_ORIGINAL_GENERATE", original_generate)
        monkeypatch.setattr(runtime, "live_article_verification_enabled", lambda: True)
        monkeypatch.setattr(runtime, "verify_document_articles", strict_verifier)

        returned_draft, payload, filename, meta = await runtime._guarded_generate(
            document_type,
            "context",
            "ru",
        )

        Document(io.BytesIO(payload))  # valid DOCX, not an error payload
        text = runtime._docx_text(payload)
        assert "Факты дела" in text
        assert "ПРОШУ" in text
        assert "статье 630 ГК РК" not in text
        assert "ПЕРЕД ПОДАЧЕЙ" in text
        assert filename == f"{document_type}.docx"
        assert returned_draft.status == runtime.VerificationStatus.NEEDS_VERIFICATION
        assert meta["filing_ready"] is False
        assert meta["release_status"] == "preliminary"
        assert meta["live_article_verification"] == "repaired"
        assert "статья 630 ГК РК" in meta["live_article_removed"]
        assert any(
            str(note).startswith("FILING_ACTION:")
            for note in meta["verification_notes"]
        )

    asyncio.run(scenario())


def test_adilet_outage_strips_unverified_citations_and_returns_word(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """A temporary official-source outage must not strand a paid document."""
    import korgan.live_article_release_runtime as runtime

    async def scenario() -> None:
        original_word = _docx_bytes(
            "Факты дела и доказательства сохранены.",
            "На основании статьи 272 ГК РК обязательство должно быть исполнено.",
            "В порядке статьи 148 ГПК РК заявитель обращается в суд.",
            "ПРОШУ: удовлетворить требования в подтвержденной части.",
        )
        draft = SimpleNamespace(status=runtime.VerificationStatus.VERIFIED)

        async def original_generate(kind: str, context: str, language: str):
            return draft, original_word, "KORGAN_iskovoe_zayavlenie.docx", _meta()

        async def unavailable_verifier(payload: bytes) -> None:
            if runtime.citation_audit.extract_references(runtime._docx_text(payload)):
                raise runtime.LiveArticleVerificationError(
                    "не удалось открыть актуальную официальную редакцию K940001000_ на Adilet"
                )

        monkeypatch.setattr(runtime, "_ORIGINAL_GENERATE", original_generate)
        monkeypatch.setattr(runtime, "live_article_verification_enabled", lambda: True)
        monkeypatch.setattr(runtime, "verify_document_articles", unavailable_verifier)

        _, payload, _, meta = await runtime._guarded_generate("claim", "context", "ru")

        Document(io.BytesIO(payload))
        text = runtime._docx_text(payload)
        assert "Факты дела и доказательства сохранены" in text
        assert "ПРОШУ" in text
        assert runtime.citation_audit.extract_references(text) == []
        assert meta["filing_ready"] is False
        assert meta["release_status"] == "preliminary"
        assert meta["live_article_verification"] == "repaired"

    asyncio.run(scenario())


def test_verified_word_is_returned_unchanged(monkeypatch: pytest.MonkeyPatch) -> None:
    """The resilience path must not downgrade a document that passes live QA."""
    import korgan.live_article_release_runtime as runtime

    async def scenario() -> None:
        original_word = _docx_bytes("Проверенный документ без замечаний.")
        draft = SimpleNamespace(status=runtime.VerificationStatus.VERIFIED)
        original_meta = _meta()

        async def original_generate(kind: str, context: str, language: str):
            return draft, original_word, "verified.docx", original_meta

        async def verifier(payload: bytes) -> None:
            assert payload == original_word

        monkeypatch.setattr(runtime, "_ORIGINAL_GENERATE", original_generate)
        monkeypatch.setattr(runtime, "live_article_verification_enabled", lambda: True)
        monkeypatch.setattr(runtime, "verify_document_articles", verifier)

        returned_draft, payload, filename, meta = await runtime._guarded_generate(
            "response",
            "context",
            "ru",
        )

        assert returned_draft is draft
        assert payload == original_word
        assert filename == "verified.docx"
        assert meta == original_meta
        assert draft.status == runtime.VerificationStatus.VERIFIED

    asyncio.run(scenario())

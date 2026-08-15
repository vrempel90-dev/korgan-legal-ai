from __future__ import annotations

from io import BytesIO
from types import SimpleNamespace

from docx import Document

from korgan_legal_ai.documents.extractor import (
    ExtractedDocument,
    OpenAIDocumentExtractor,
    VisualPageTranscription,
    VisualTranscription,
)
from korgan_legal_ai.telegram.word_export import build_claim_docx


class NeverVisionClient:
    class Responses:
        def parse(self, **kwargs):
            raise AssertionError("vision must not be called for DOCX")

    responses = Responses()


class CapturingVisionClient:
    def __init__(self) -> None:
        self.calls = []
        self.responses = self

    def parse(self, **kwargs):
        self.calls.append(kwargs)
        return SimpleNamespace(
            output_parsed=VisualTranscription(
                pages=[
                    VisualPageTranscription(
                        page_number=1,
                        text="Договор №15 от 10.02.2026. Сумма 5 450 000 тенге.",
                        uncertain_fragments=["БИН: 1234?6789012"],
                    )
                ]
            )
        )


def test_docx_is_read_locally_including_tables() -> None:
    source = Document()
    source.add_paragraph("Договор поставки №15")
    table = source.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Сумма"
    table.cell(0, 1).text = "5 450 000 тенге"
    buffer = BytesIO()
    source.save(buffer)

    extractor = OpenAIDocumentExtractor(api_key="", client=NeverVisionClient())  # type: ignore[arg-type]
    result = extractor.extract(
        source_id="doc-1",
        data=buffer.getvalue(),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert result.method == "docx_text"
    assert "Договор поставки №15" in result.text
    assert "Сумма | 5 450 000 тенге" in result.text
    assert result.uncertain_fragments == ()
    assert "VISUAL_TRANSCRIPTION" not in result.as_case_context()


def test_image_transcription_is_store_false_and_confirmation_gated() -> None:
    client = CapturingVisionClient()
    extractor = OpenAIDocumentExtractor(api_key="", client=client)  # type: ignore[arg-type]

    result = extractor.extract(source_id="doc-2", data=b"fake-jpeg", mime_type="image/jpeg")
    context = result.as_case_context()

    assert result.method == "vision"
    assert "Договор №15" in result.text
    assert result.uncertain_fragments == ("БИН: 1234?6789012",)
    assert "[VISUAL_TRANSCRIPTION_REQUIRES_CONFIRMATION]" in context
    assert "[OCR_UNCERTAIN]" in context
    assert client.calls[0]["store"] is False
    image_item = client.calls[0]["input"][1]["content"][1]
    assert image_item["type"] == "input_image"
    assert image_item["image_url"].startswith("data:image/jpeg;base64,")


def test_case_context_uses_opaque_source_id_not_original_filename() -> None:
    extracted = ExtractedDocument(
        source_id="doc-7",
        mime_type="application/pdf",
        method="pdf_text",
        text="[PAGE 1]\nИстец: ТОО Alpha",
    )
    context = extracted.as_case_context()
    assert "[DOCUMENT doc-7]" in context
    assert "Alpha" in context


def test_word_export_preserves_legal_text_and_verification_markers() -> None:
    source = (
        "В: NEEDS_VERIFICATION — определить суд\n\n"
        "Истец: ТОО A\nОтветчик: ТОО B\n\n"
        "ИСКОВОЕ ЗАЯВЛЕНИЕ\nо взыскании задолженности\n\n"
        "ЦЕНА ИСКА: 5450000 KZT\n\n"
        "ПРОШУ СУД:\n1. Взыскать 5450000 KZT."
    )
    payload = build_claim_docx(source)

    doc = Document(BytesIO(payload))
    rendered = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "NEEDS_VERIFICATION — определить суд" in rendered
    assert "ИСКОВОЕ ЗАЯВЛЕНИЕ" in rendered
    assert "5450000 KZT" in rendered
    assert doc.core_properties.author in {None, ""}

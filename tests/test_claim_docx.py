import io

from docx import Document
from docx.oxml.ns import qn

from korgan.claim_docx import build_claim_docx
from korgan.legal_types import ClaimDraft, VerificationStatus


def _draft(**overrides) -> ClaimDraft:
    base = dict(
        status=VerificationStatus.NEEDS_VERIFICATION,
        title="ИСКОВОЕ ЗАЯВЛЕНИЕ",
        court="",
        claimant=[],
        defendant=[],
        price_of_claim="",
        facts=["Факт из материалов дела"],
        legal_basis=["Подтвержденная правовая основа"],
        requests=["Удовлетворить заявленное требование"],
        attachments=["Копия документа"],
        verification_notes=["Уточнить подсудность"],
        source_urls=["https://adilet.zan.kz/rus/docs/K1500000377"],
    )
    base.update(overrides)
    return ClaimDraft(**base)


def _header_text(payload: bytes) -> str:
    """Шапка иска, а не первый параграф: перед ней печатается KORGAN QA STATUS."""
    for paragraph in Document(io.BytesIO(payload)).paragraphs:
        if "В суд:" in paragraph.text:
            return paragraph.text
    raise AssertionError("В документе нет шапки с реквизитами сторон")


def _num_ids(payload: bytes) -> list[str | None]:
    doc = Document(io.BytesIO(payload))
    ids: list[str | None] = []
    for paragraph in doc.paragraphs:
        if paragraph.style.name != "List Number":
            continue
        found = paragraph._p.find(f'{qn("w:pPr")}/{qn("w:numPr")}/{qn("w:numId")}')
        ids.append(found.get(qn("w:val")) if found is not None else None)
    return ids


def test_claim_docx_is_valid_zip_container() -> None:
    payload = build_claim_docx(_draft(court="", claimant=[], defendant=[], price_of_claim=""))

    assert payload.startswith(b"PK")
    assert len(payload) > 1000


def test_party_label_is_not_duplicated() -> None:
    """The model may return «Истец: ...»; the template already prints the label."""
    payload = build_claim_docx(
        _draft(
            claimant=["Истец: Иванов Иван Иванович", "ИИН 800101300123"],
            defendant=["Ответчик: Петров Пётр Петрович"],
        )
    )

    header = _header_text(payload)
    assert header.count("Истец:") == 1
    assert header.count("Ответчик:") == 1
    assert "Иванов Иван Иванович" in header
    assert "Петров Пётр Петрович" in header


def test_court_and_price_labels_are_not_duplicated() -> None:
    payload = build_claim_docx(
        _draft(
            court="В суд: Медеуский районный суд города Алматы",
            price_of_claim="Цена иска: 1 500 000 тенге",
        )
    )

    header = _header_text(payload)
    assert header.count("В суд:") == 1
    assert header.count("Цена иска:") == 1
    assert "Медеуский районный суд города Алматы" in header


def test_missing_party_still_renders_placeholder() -> None:
    payload = build_claim_docx(_draft(claimant=[], defendant=["Ответчик:"]))

    header = _header_text(payload)
    assert "[ДАННЫЕ: данные истца]" in header
    assert "[ДАННЫЕ: ИИН ответчика]" in header
    assert "[ДАННЫЕ: адрес ответчика]" in header


def test_attachments_restart_numbering_from_one() -> None:
    """Annexes are an independent list, not a continuation of the prayer for relief."""
    payload = build_claim_docx(
        _draft(
            requests=["Взыскать основной долг", "Взыскать судебные расходы"],
            attachments=["Копия договора займа", "Копия платежного поручения"],
        )
    )

    request_ids = _num_ids(payload)[:2]
    annex_ids = _num_ids(payload)[2:]

    assert len(annex_ids) == 2
    assert len(set(annex_ids)) == 1
    assert annex_ids[0] is not None
    assert annex_ids[0] not in request_ids


def test_annex_numbering_declares_start_override() -> None:
    payload = build_claim_docx(
        _draft(requests=["Взыскать основной долг"], attachments=["Копия договора займа"])
    )

    doc = Document(io.BytesIO(payload))
    annex_num_id = _num_ids(payload)[-1]
    numbering = doc.part.numbering_part.element

    overrides = [
        num
        for num in numbering.findall(qn("w:num"))
        if num.get(qn("w:numId")) == annex_num_id
        and num.find(f'{qn("w:lvlOverride")}/{qn("w:startOverride")}') is not None
    ]
    assert len(overrides) == 1
    start = overrides[0].find(f'{qn("w:lvlOverride")}/{qn("w:startOverride")}')
    assert start.get(qn("w:val")) == "1"

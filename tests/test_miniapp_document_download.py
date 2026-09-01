"""Скачивание документа проверяется скачиванием документа.

Ссылку на файл выдаёт отдельный шаг: `document/access` возвращает подписанный
адрес, по которому Telegram WebView качает файл своим механизмом. Проверено при
этом было только то, что поддельный токен отвергается. Всё остальное —
существует ли файл по выданной ссылке, приходит ли он как DOCX, под каким
именем его сохранит клиент, те ли это байты, что лежат в деле, — не проверялось
ни разу.

Для пользователя это последний шаг всей работы, и «скачал, а там ноль байт»
или «скачал, а оно не открывается» неотличимо от несделанной работы. Поэтому
здесь ссылка действительно используется, а ответ разбирается целиком.
"""

from __future__ import annotations

import base64
import hashlib
import hmac
import io
import json
import time
from typing import Any
from urllib.parse import urlencode

import pytest
from docx import Document
from fastapi.testclient import TestClient

from korgan import miniapp_api
from korgan import miniapp_api_v2 as core
from korgan import miniapp_document_access as access
from korgan.miniapp_api_recovery_cors import app

CASE_ID = "KOR-DOWNLOAD-0001"
FILENAME = "KORGAN_iskovoe_zayavlenie.docx"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Исковое заявление", level=1)
    document.add_paragraph("О взыскании задолженности по договору подряда.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _init_data(user_id: int = -900000901) -> str:
    pairs = {
        "auth_date": str(int(time.time())),
        "query_id": "korgan-document-download",
        "user": json.dumps(
            {"id": user_id, "first_name": "KORGAN", "language_code": "ru"},
            ensure_ascii=False,
            separators=(",", ":"),
        ),
    }
    data_check_string = "\n".join(f"{key}={pairs[key]}" for key in sorted(pairs))
    secret_key = hmac.new(
        b"WebAppData",
        miniapp_api.settings.telegram_bot_token.encode(),
        hashlib.sha256,
    ).digest()
    pairs["hash"] = hmac.new(secret_key, data_check_string.encode(), hashlib.sha256).hexdigest()
    return urlencode(pairs)


@pytest.fixture()
def ready_document(monkeypatch) -> bytes:
    """Дело с готовым документом, каким его видит выдача файла."""
    payload = _docx_bytes()
    state: dict[str, Any] = {
        "consent": {"accepted": True, "terms_version": "2026-08-16-v1"},
        "cases": {
            CASE_ID: {
                "id": CASE_ID,
                "title": "Исковое заявление",
                "filename": FILENAME,
                "document_base64": base64.b64encode(payload).decode("ascii"),
            }
        },
    }

    async def load(_user_id: str) -> dict[str, Any]:
        return state

    monkeypatch.setattr(core.legacy, "_state", load)
    return payload


def test_issued_link_returns_the_document_itself(ready_document: bytes) -> None:
    headers = {"X-Telegram-Init-Data": _init_data()}

    with TestClient(app) as client:
        issued = client.post(f"/miniapp/cases/{CASE_ID}/document/access", headers=headers)
        assert issued.status_code == 200
        link = issued.json()

        assert link["filename"] == FILENAME
        assert link["download_url"].startswith("http")

        downloaded = client.get(link["download_url"])

    assert downloaded.status_code == 200
    # Байты дела, а не пересобранный где-то по дороге файл.
    assert downloaded.content == ready_document
    # DOCX — это ZIP: клиент, открывающий пустышку, увидит поломанный файл.
    assert downloaded.content[:4] == b"PK\x03\x04"
    assert Document(io.BytesIO(downloaded.content)).paragraphs[0].text == "Исковое заявление"


def test_downloaded_file_arrives_as_a_named_word_document(ready_document: bytes) -> None:
    headers = {"X-Telegram-Init-Data": _init_data()}

    with TestClient(app) as client:
        link = client.post(f"/miniapp/cases/{CASE_ID}/document/access", headers=headers).json()
        downloaded = client.get(link["download_url"])

    assert downloaded.headers["content-type"] == DOCX_MIME
    disposition = downloaded.headers["content-disposition"]
    # Файл сохраняется, а не открывается вкладкой, и имя дойдёт до обоих
    # видов клиентов: ASCII-запасное и настоящее в UTF-8.
    assert disposition.startswith("attachment;")
    assert f'filename="{FILENAME}"' in disposition
    assert f"filename*=UTF-8''{FILENAME}" in disposition
    assert downloaded.headers["x-content-type-options"] == "nosniff"


def test_the_same_link_downloads_more_than_once(ready_document: bytes) -> None:
    headers = {"X-Telegram-Init-Data": _init_data()}

    with TestClient(app) as client:
        link = client.post(f"/miniapp/cases/{CASE_ID}/document/access", headers=headers).json()
        first = client.get(link["download_url"])
        second = client.get(link["download_url"])

    assert first.status_code == 200
    assert second.status_code == 200
    assert second.content == ready_document


def test_preview_link_opens_the_document_text(ready_document: bytes) -> None:
    headers = {"X-Telegram-Init-Data": _init_data()}

    with TestClient(app) as client:
        link = client.post(f"/miniapp/cases/{CASE_ID}/document/access", headers=headers).json()
        preview = client.get(link["preview_url"])

    assert preview.status_code == 200
    assert preview.headers["content-type"].startswith("text/html")
    assert "Исковое заявление" in preview.text
    # Просмотр остаётся внутри KORGAN и ничего не подгружает со стороны.
    assert "default-src 'none'" in preview.headers["content-security-policy"]


def test_expired_link_stops_working(ready_document: bytes) -> None:
    expired = access._make_token(
        "-900000901",
        CASE_ID,
        int(time.time()) - 1,
        miniapp_api.settings.telegram_bot_token,
    )

    with TestClient(app) as client:
        response = client.get("/miniapp/document/download", params={"token": expired})

    assert response.status_code == 403


def test_link_signed_for_another_case_is_refused(ready_document: bytes) -> None:
    foreign = access._make_token(
        "-900000901",
        "KOR-SOMEONE-ELSE",
        int(time.time()) + 120,
        miniapp_api.settings.telegram_bot_token,
    )

    with TestClient(app) as client:
        response = client.get("/miniapp/document/download", params={"token": foreign})

    assert response.status_code == 404


def test_tampered_link_is_refused(ready_document: bytes) -> None:
    headers = {"X-Telegram-Init-Data": _init_data()}

    with TestClient(app) as client:
        link = client.post(f"/miniapp/cases/{CASE_ID}/document/access", headers=headers).json()
        token = link["download_url"].split("token=", 1)[1]
        payload, signature = token.split("%2E" if "%2E" in token else ".", 1)
        # Подпись шестнадцатеричная, поэтому «0» вместо «1» гарантированно её меняет.
        forged = f"{payload}.{signature[:-1]}{'0' if signature[-1] != '0' else '1'}"
        assert forged != token
        response = client.get("/miniapp/document/download", params={"token": forged})

    assert response.status_code == 403

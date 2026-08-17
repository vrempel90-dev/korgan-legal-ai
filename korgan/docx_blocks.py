"""Typed DOCX blocks and deterministic numbering for KORGAN documents."""

from __future__ import annotations

from dataclasses import dataclass, field

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

from korgan.contract_numbering import strip_leading_number


@dataclass(slots=True)
class Heading:
    text: str
    centered: bool = False
    space_before: int = 10


@dataclass(slots=True)
class Prose:
    text: str
    indent_levels: int = 0
    first_line_indent: bool = True
    centered: bool = False
    bold: bool = False


@dataclass(slots=True)
class NumberedItem:
    text: str
    level: int = 0
    bold: bool = False


@dataclass(slots=True)
class AutoNumberedList:
    """A real Word numbered list.

    ``restart=True`` creates a fresh numbering definition with start=1. This is
    used for «Приложения», so it can never continue the prayer-for-relief list.
    """
    items: list[str] = field(default_factory=list)
    restart: bool = False


@dataclass(slots=True)
class Spacer:
    pass


Block = Heading | Prose | NumberedItem | AutoNumberedList | Spacer


def advance(counters: list[int], level: int) -> str:
    if level < 0:
        raise ValueError("numbering level must not be negative")
    while len(counters) <= level:
        counters.append(0)
    del counters[level + 1:]
    counters[level] += 1
    return ".".join(str(value) for value in counters[: level + 1]) + "."


def _next_numeric_id(elements, attr: str) -> int:
    used: set[int] = set()
    for element in elements:
        raw = element.get(qn(attr))
        if raw and raw.isdigit():
            used.add(int(raw))
    return max(used, default=0) + 1


def _fresh_numbering_from_one(doc) -> int:
    """Create a self-contained decimal Word list starting at 1.

    We do not clone the document's ``List Number`` style because some DOCX
    viewers keep the inherited list state and continue numbering across legal
    sections. A dedicated abstractNum + numId makes the restart explicit in the
    file itself and interoperable across Word/LibreOffice/mobile viewers.
    """
    numbering = doc.part.numbering_part.element
    abstract_id = _next_numeric_id(numbering.findall(qn("w:abstractNum")), "w:abstractNumId")
    num_id = _next_numeric_id(numbering.findall(qn("w:num")), "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)

    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)

    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    reference = OxmlElement("w:abstractNumId")
    reference.set(qn("w:val"), str(abstract_id))
    num.append(reference)

    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def _apply_num_id(paragraph: Paragraph, num_id: int) -> None:
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = num_id


def render_blocks(doc, blocks: list[Block]) -> None:
    counters: list[int] = []

    for block in blocks:
        if isinstance(block, Spacer):
            doc.add_paragraph()
            continue

        if isinstance(block, Heading):
            paragraph = doc.add_paragraph()
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if block.centered else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.space_before = Pt(block.space_before)
            paragraph.add_run(block.text).bold = True
            continue

        if isinstance(block, Prose):
            paragraph = doc.add_paragraph()
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if block.centered else WD_ALIGN_PARAGRAPH.JUSTIFY
            )
            if block.first_line_indent:
                paragraph.paragraph_format.first_line_indent = Cm(0.8)
            if block.indent_levels:
                paragraph.paragraph_format.left_indent = Cm(0.8 * block.indent_levels)
            paragraph.add_run(block.text).bold = block.bold
            continue

        if isinstance(block, NumberedItem):
            number = advance(counters, block.level)
            paragraph = doc.add_paragraph()
            if block.level == 0 and block.bold:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(10)
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.paragraph_format.first_line_indent = Cm(0.8)
                if block.level >= 2:
                    paragraph.paragraph_format.left_indent = Cm(0.8)
            run = paragraph.add_run(f"{number} {strip_leading_number(block.text)}")
            run.bold = block.bold
            continue

        if isinstance(block, AutoNumberedList):
            num_id = _fresh_numbering_from_one(doc) if block.restart else None
            for item in block.items:
                paragraph = doc.add_paragraph(strip_leading_number(item), style="List Number")
                if num_id is not None:
                    _apply_num_id(paragraph, num_id)
            continue

        raise TypeError(f"unsupported document block: {type(block).__name__}")

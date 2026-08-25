"""Production invariants v2 for the KORGAN legal-document pipeline.

This module is intentionally installed last by ``strict_bot``.  The older
production layers remain available for rollback, while this adapter enforces
cross-cutting release invariants which cannot safely live in one document
renderer or one model prompt.

The rules are operational, not cosmetic:
* a quality issue is either user-resolvable (NEEDS_USER_DATA) or internal
  (INTERNAL_QUALITY);
* only NEEDS_USER_DATA may stop delivery;
* remaining INTERNAL_QUALITY issues are written into the DOCX as [СВЕРИТЬ: ...]
  and repeated verbatim in the Telegram response;
* every monetary input produces a deterministic ledger; contractual daily
  penalties are calculated in code, including a percentage cap;
* the research pass uses high web context and its output order is canonical;
* identical input is stable within a running production process;
* an unchanged blocker set never causes another repair model call;
* deterministic finalization runs before claim quality/repair;
* a newer Telegram update cancels a stale in-flight update for the same chat.
"""

from __future__ import annotations

import asyncio
import contextvars
import copy
import hashlib
import io
import json
import logging
import re
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from decimal import Decimal, ROUND_CEILING, ROUND_HALF_UP
from enum import StrEnum
from typing import Any, Awaitable, Callable
from zoneinfo import ZoneInfo

from aiogram import BaseMiddleware
from aiogram.types import BufferedInputFile, CallbackQuery, Message, TelegramObject
from docx import Document

from korgan.legal_calc import format_kzt, parse_all_amounts_kzt
from korgan.legal_types import ClaimDraft, LegalResearch, VerificationStatus

LOGGER = logging.getLogger(__name__)


class BlockerClass(StrEnum):
    NEEDS_USER_DATA = "NEEDS_USER_DATA"
    INTERNAL_QUALITY = "INTERNAL_QUALITY"


@dataclass(frozen=True, slots=True)
class ClassifiedIssue:
    text: str
    blocker_class: BlockerClass
    action: str


@dataclass(frozen=True, slots=True)
class ContractualPenalty:
    principal: int
    daily_rate_percent: Decimal
    cap_percent: Decimal | None
    start: date
    as_of: date
    days: int
    daily_amount: int
    uncapped_amount: int
    amount: int
    cap_amount: int | None
    cap_reached_on: date | None


@dataclass(frozen=True, slots=True)
class MoneyLedger:
    input_amounts: tuple[int, ...]
    principal: int | None
    penalty: ContractualPenalty | None
    rows: tuple[tuple[str, int], ...]
    unresolved: tuple[str, ...]

    @property
    def total(self) -> int:
        return sum(value for _, value in self.rows if value > 0)


_USER_DATA_PATTERNS: tuple[re.Pattern[str], ...] = (
    re.compile(r"(?i)не\s+(?:указан|указана|указаны|заполнен[аы]?)\s+(?:отправител|адресат|истец|ответчик|данн|фио|бин|иин|адрес)"),
    re.compile(r"(?i)(?:данные|реквизиты)\s+(?:истца|ответчика|отправителя|адресата).*?(?:нет|отсутств|уточн)"),
    re.compile(r"(?i)не\s+хватает.*?(?:даты|срока|суммы|адреса|фио|бин|иин|реквизит)"),
    re.compile(r"(?i)(?:дата\s+начала\s+просрочки|срок\s+исполнения).*?(?:не\s+удал|не\s+установ|уточн|отсутств)"),
    re.compile(r"(?i)нет\s+фактическ\w*\s+(?:основан|обстоятельств)"),
    re.compile(r"(?i)отсутствует\s+просительная\s+часть"),
)
_INTERNAL_FORCE_PATTERNS: tuple[re.Pattern[str], ...] = (
    re.compile(r"(?i)стать(?:я|и|е|ю|ёй|ей)|\bст\.\s*\d"),
    re.compile(r"(?i)правов\w*\s+(?:основан|ссыл|норм)"),
    re.compile(r"(?i)пересказ|source-bound|verified|цитат|норм\w*\s+права"),
    re.compile(r"(?i)служебн\w*\s+фраз|целостност|поврежден|повреждён|экспорт"),
    re.compile(r"(?i)цена\s+денежн\w*\s+иск\w*.*известн\w*\s+сумм"),
)

_DAILY_RATE_RE = re.compile(
    r"(?<!\d)(?P<rate>\d{1,3}(?:[.,]\d{1,4})?)\s*%\s*(?:/|в\s+|за\s+)?(?:день|дн(?:я|ь)?|сутк\w*)",
    re.IGNORECASE,
)
_CAP_PERCENT_RE = re.compile(
    r"(?:потолок|предел|максимум|не\s+более|огранич\w*)[^%\n]{0,55}?(?P<cap>\d{1,3}(?:[.,]\d{1,4})?)\s*%",
    re.IGNORECASE,
)
_PENALTY_RE = re.compile(r"(?i)неустойк\w*|пен[яию]\b|штрафн\w*|просроч\w*")
_DATE = r"(?:\d{1,2}[./-]\d{1,2}[./-]\d{4})"
_OVERDUE_START_RE = re.compile(
    rf"(?:просроч\w*[^.\n]{{0,45}}?\sс|начиная\s+с)\s*(?P<date>{_DATE})",
    re.IGNORECASE,
)
_EXCLUDE_MONEY_ROW_RE = re.compile(r"(?i)госпошлин|государственн\w*\s+пошлин|судебн\w*\s+расход")

_RUN_ID: contextvars.ContextVar[str] = contextvars.ContextVar("korgan_invariant_run_id", default="")
_REPAIR_BLOCKER_SETS: contextvars.ContextVar[frozenset[str]] = contextvars.ContextVar(
    "korgan_repair_blocker_sets", default=frozenset()
)
_DELIVERED_KINDS: contextvars.ContextVar[frozenset[str]] = contextvars.ContextVar(
    "korgan_delivered_kinds", default=frozenset()
)
_ACTIVE_UPDATES: dict[int, asyncio.Task[Any]] = {}
_ACTIVE_LOCK = asyncio.Lock()
_RESEARCH_CACHE: dict[str, LegalResearch] = {}
_RESEARCH_CACHE_MAX = 128


def _compact(value: str) -> str:
    return " ".join(str(value or "").split()).strip()


def _today_kz() -> date:
    return datetime.now(ZoneInfo("Asia/Almaty")).date()


def _issue_action(text: str) -> str:
    lower = text.lower()
    if "отправител" in lower:
        return "укажите отправителя: ФИО/наименование и имеющиеся реквизиты"
    if "адресат" in lower:
        return "укажите адресата претензии: ФИО/наименование и известный адрес"
    if "истц" in lower:
        return "укажите данные истца, которых нет в материалах дела"
    if "ответчик" in lower:
        return "укажите данные ответчика, которых нет в материалах дела"
    if "дата" in lower or "просроч" in lower or "срок" in lower:
        return "укажите точную дату/срок из договора или иных материалов"
    if "сумм" in lower or "цен" in lower:
        return "укажите исходную денежную сумму, если её действительно нет в материалах"
    if "фактичес" in lower or "обстоятель" in lower:
        return "опишите недостающие обстоятельства одним сообщением"
    if "проситель" in lower or "требован" in lower:
        return "укажите, какого результата вы требуете от другой стороны/суда"
    return "добавьте недостающие фактические данные, указанные в причине"


def classify_issue(issue: str) -> ClassifiedIssue:
    text = _compact(issue)
    # Legal wording, citation drift and generated-text defects are never shifted
    # onto the user even when they contain words like "уточнить".
    if any(pattern.search(text) for pattern in _INTERNAL_FORCE_PATTERNS):
        return ClassifiedIssue(text, BlockerClass.INTERNAL_QUALITY, "KORGAN должен исправить/пометить это сам")
    if any(pattern.search(text) for pattern in _USER_DATA_PATTERNS):
        return ClassifiedIssue(text, BlockerClass.NEEDS_USER_DATA, _issue_action(text))
    return ClassifiedIssue(text, BlockerClass.INTERNAL_QUALITY, "KORGAN должен исправить/пометить это сам")


def classify_issues(issues: list[str] | tuple[str, ...]) -> tuple[list[ClassifiedIssue], list[ClassifiedIssue]]:
    unique = list(dict.fromkeys(_compact(item) for item in issues if _compact(item)))
    classified = [classify_issue(item) for item in unique]
    user = [item for item in classified if item.blocker_class == BlockerClass.NEEDS_USER_DATA]
    internal = [item for item in classified if item.blocker_class == BlockerClass.INTERNAL_QUALITY]
    return user, internal


def review_marker(issue: str) -> str:
    text = _compact(issue).strip(" .")
    if text.startswith("[") and text.endswith("]"):
        text = text[1:-1].strip()
    return f"[СВЕРИТЬ: {text}]"


def _user_data_message(kind: str, issues: list[ClassifiedIssue]) -> str:
    rows = [f"• {item.text}\n  Что прислать: {item.action}." for item in issues[:6]]
    return (
        f"Документ «{kind}» пока не выпущен: не хватает данных, которые может сообщить только пользователь.\n\n"
        + "\n".join(rows)
    )


def _internal_message(issues: list[ClassifiedIssue]) -> str:
    return "\n".join(f"• {item.text}" for item in issues[:6])


def append_review_markers(file_bytes: bytes, issues: list[str], *, language: str = "ru") -> bytes:
    markers = list(dict.fromkeys(review_marker(item) for item in issues if _compact(item)))
    if not markers:
        return file_bytes
    doc = Document(io.BytesIO(file_bytes))
    doc.add_paragraph()
    heading = doc.add_paragraph()
    heading.add_run("KORGAN: тексеруді қажет ететін тармақтар" if language == "kk" else "KORGAN: отметки, требующие проверки").bold = True
    for marker in markers:
        doc.add_paragraph(marker)
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _parse_date(raw: str) -> date | None:
    value = (raw or "").strip()
    for fmt in ("%d.%m.%Y", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return datetime.strptime(value, fmt).date()
        except ValueError:
            continue
    return None


def calc_contractual_penalty(
    principal: int,
    daily_rate_percent: Decimal | float | str,
    start: date,
    as_of: date,
    *,
    cap_percent: Decimal | float | str | None = None,
) -> ContractualPenalty:
    if principal <= 0:
        raise ValueError("principal must be positive")
    if as_of < start:
        raise ValueError("as_of precedes penalty start")
    rate = Decimal(str(daily_rate_percent)).copy_abs()
    if rate <= 0:
        raise ValueError("daily rate must be positive")
    cap = Decimal(str(cap_percent)).copy_abs() if cap_percent is not None else None
    if cap is not None and cap <= 0:
        raise ValueError("cap must be positive")

    days = (as_of - start).days + 1
    exact_daily = Decimal(principal) * rate / Decimal("100")
    daily_amount = int(exact_daily.quantize(Decimal("1"), rounding=ROUND_HALF_UP))
    uncapped_exact = exact_daily * Decimal(days)
    uncapped_amount = int(uncapped_exact.quantize(Decimal("1"), rounding=ROUND_HALF_UP))

    cap_amount: int | None = None
    cap_reached_on: date | None = None
    amount = uncapped_amount
    if cap is not None:
        cap_exact = Decimal(principal) * cap / Decimal("100")
        cap_amount = int(cap_exact.quantize(Decimal("1"), rounding=ROUND_HALF_UP))
        amount = min(uncapped_amount, cap_amount)
        days_to_cap = int((cap_exact / exact_daily).to_integral_value(rounding=ROUND_CEILING))
        cap_reached_on = start + timedelta(days=max(0, days_to_cap - 1))

    return ContractualPenalty(
        principal=principal,
        daily_rate_percent=rate,
        cap_percent=cap,
        start=start,
        as_of=as_of,
        days=days,
        daily_amount=daily_amount,
        uncapped_amount=uncapped_amount,
        amount=amount,
        cap_amount=cap_amount,
        cap_reached_on=cap_reached_on,
    )


def _contractual_penalty_from_context(case_context: str, principal: int | None, *, as_of: date) -> tuple[ContractualPenalty | None, list[str]]:
    if not principal or not _PENALTY_RE.search(case_context or ""):
        return None, []
    unresolved: list[str] = []
    rate_match = _DAILY_RATE_RE.search(case_context or "")
    start_match = _OVERDUE_START_RE.search(case_context or "")
    cap_match = _CAP_PERCENT_RE.search(case_context or "")
    if not rate_match:
        unresolved.append("для договорной неустойки не распознана дневная ставка")
    if not start_match:
        unresolved.append("для договорной неустойки не распознана дата начала просрочки")
    if unresolved:
        return None, unresolved
    start = _parse_date(start_match.group("date"))
    if start is None:
        return None, ["для договорной неустойки не удалось разобрать дату начала просрочки"]
    rate = Decimal(rate_match.group("rate").replace(",", "."))
    cap = Decimal(cap_match.group("cap").replace(",", ".")) if cap_match else None
    if as_of < start:
        return None, []
    return calc_contractual_penalty(principal, rate, start, as_of, cap_percent=cap), []


def build_money_ledger(case_context: str, draft: ClaimDraft | None = None, *, as_of: date | None = None) -> MoneyLedger:
    checked_on = as_of or _today_kz()
    input_amounts = tuple(dict.fromkeys(parse_all_amounts_kzt(case_context or "")))
    principal = max(input_amounts) if input_amounts else None

    rows: list[tuple[str, int]] = []
    if draft is not None:
        seen: set[int] = set()
        for request in draft.requests:
            if _EXCLUDE_MONEY_ROW_RE.search(str(request)):
                continue
            for amount in parse_all_amounts_kzt(str(request)):
                if amount > 0 and amount not in seen:
                    seen.add(amount)
                    rows.append(("document_demand", amount))
    if not rows and principal:
        # This fallback is deliberate: ledger_total=0 with a monetary input is a
        # pipeline failure, not a legitimate arithmetic result.  The source row
        # also makes the propagation defect observable in logs/tests.
        rows.append(("input_money_not_yet_propagated", principal))

    penalty, unresolved = _contractual_penalty_from_context(case_context, principal, as_of=checked_on)
    if penalty is not None and all(value != penalty.amount for _, value in rows):
        rows.append(("contractual_penalty", penalty.amount))

    ledger = MoneyLedger(
        input_amounts=input_amounts,
        principal=principal,
        penalty=penalty,
        rows=tuple(rows),
        unresolved=tuple(unresolved),
    )
    if ledger.input_amounts and ledger.total <= 0:
        LOGGER.error("PIPELINE_INVARIANT_VIOLATION invariant=I5 reason=money_input_but_zero_ledger")
        raise RuntimeError("I5: ledger_total=0 при непустом денежном входе")
    return ledger


def apply_money_ledger_to_claim(case_context: str, draft: ClaimDraft, *, as_of: date | None = None) -> MoneyLedger:
    ledger = build_money_ledger(case_context, draft, as_of=as_of)
    penalty = ledger.penalty
    if penalty is not None and _PENALTY_RE.search(case_context or ""):
        calculated = (
            f"Взыскать неустойку в размере {format_kzt(penalty.amount)} за период "
            f"с {penalty.start.strftime('%d.%m.%Y')} по {penalty.as_of.strftime('%d.%m.%Y')} "
            f"из расчёта {penalty.daily_rate_percent:g}% в день"
        )
        if penalty.cap_amount is not None and penalty.cap_reached_on is not None:
            calculated += (
                f"; договорный предел {penalty.cap_percent:g}% = {format_kzt(penalty.cap_amount)} "
                f"достигнут {penalty.cap_reached_on.strftime('%d.%m.%Y')}"
            )
        calculated += "."
        replaced = False
        requests: list[str] = []
        for request in draft.requests:
            if _PENALTY_RE.search(str(request)):
                if not replaced:
                    requests.append(calculated)
                    replaced = True
                continue
            requests.append(request)
        if not replaced:
            requests.append(calculated)
        draft.requests = requests
        if ledger.principal:
            draft.price_of_claim = format_kzt(ledger.principal + penalty.amount)
    LOGGER.info(
        "CLAIM_MONEY_AUTHORITY price=%r input_amounts=%d ledger_total=%d principal=%r penalty=%r unresolved=%d",
        draft.price_of_claim,
        len(ledger.input_amounts),
        ledger.total,
        ledger.principal,
        penalty.amount if penalty else None,
        len(ledger.unresolved),
    )
    return ledger


def apply_money_ledger_to_pretrial(case_context: str, draft: Any, *, as_of: date | None = None) -> MoneyLedger:
    ledger = build_money_ledger(case_context, None, as_of=as_of)
    penalty = ledger.penalty
    if penalty is not None and _PENALTY_RE.search(case_context or ""):
        line = (
            f"Уплатить неустойку {format_kzt(penalty.amount)} за период с {penalty.start.strftime('%d.%m.%Y')} "
            f"по {penalty.as_of.strftime('%d.%m.%Y')} ({penalty.daily_rate_percent:g}% в день"
        )
        if penalty.cap_amount is not None and penalty.cap_reached_on is not None:
            line += (
                f", предел {penalty.cap_percent:g}% = {format_kzt(penalty.cap_amount)}, "
                f"достигнут {penalty.cap_reached_on.strftime('%d.%m.%Y')}"
            )
        line += ")."
        demands: list[str] = []
        replaced = False
        for demand in list(getattr(draft, "demands", []) or []):
            if _PENALTY_RE.search(str(demand)):
                if not replaced:
                    demands.append(line)
                    replaced = True
                continue
            demands.append(str(demand))
        if not replaced:
            demands.append(line)
        draft.demands = demands
    LOGGER.info(
        "PRETRIAL_MONEY_AUTHORITY input_amounts=%d ledger_total=%d principal=%r penalty=%r unresolved=%d",
        len(ledger.input_amounts),
        ledger.total,
        ledger.principal,
        penalty.amount if penalty else None,
        len(ledger.unresolved),
    )
    return ledger


def _article_sort_key(value: str) -> tuple[str, int, str]:
    text = _compact(value)
    match = re.search(r"(?i)(?:стать(?:я|и|е|ю|ёй|ей)|ст\.)\s*(\d+)", text)
    number = int(match.group(1)) if match else 10**9
    act = ""
    for marker in ("ГК РК", "ГПК РК", "ТК РК", "НК РК", "АППК РК"):
        if marker.lower() in text.lower():
            act = marker
            break
    return act, number, text.lower()


def canonicalize_research(research: LegalResearch) -> LegalResearch:
    research.verified_claims = sorted(dict.fromkeys(map(_compact, research.verified_claims)), key=_article_sort_key)
    research.unverified_claims = sorted(dict.fromkeys(map(_compact, research.unverified_claims)), key=str.lower)
    research.source_urls = sorted(dict.fromkeys(_compact(x) for x in research.source_urls if _compact(x)))
    research.applicable_law = sorted(dict.fromkeys(map(_compact, research.applicable_law)), key=str.lower)
    research.procedural_requirements = sorted(
        dict.fromkeys(map(_compact, research.procedural_requirements)), key=str.lower
    )
    # REMEDY/VERIFIED_COURT notes are semantically independent rows. Stable order
    # removes model-order noise without changing the chosen legal conclusions.
    research.notes = sorted(dict.fromkeys(map(_compact, research.notes)), key=str.lower)
    if len(research.verified_claims) < len(research.unverified_claims):
        research.status = VerificationStatus.NEEDS_VERIFICATION
        LOGGER.warning(
            "RESEARCH_INVARIANT verified=%d unverified=%d status=VIOLATED",
            len(research.verified_claims),
            len(research.unverified_claims),
        )
    else:
        LOGGER.info(
            "RESEARCH_INVARIANT verified=%d unverified=%d status=OK",
            len(research.verified_claims),
            len(research.unverified_claims),
        )
    return research


def _research_key(case_context: str, language: str) -> str:
    normalized = _compact(case_context).lower()
    raw = f"{_today_kz().isoformat()}|{language}|{normalized}".encode("utf-8")
    return hashlib.sha256(raw).hexdigest()


def _mark_delivery_once(kind: str) -> bool:
    delivered = _DELIVERED_KINDS.get()
    if kind in delivered:
        LOGGER.error("PIPELINE_INVARIANT_VIOLATION invariant=I10 duplicate_finalization kind=%s run=%s", kind, _RUN_ID.get())
        return False
    _DELIVERED_KINDS.set(frozenset((*delivered, kind)))
    return True


def _repair_signature(issues: list[str]) -> str:
    normalized = sorted(dict.fromkeys(_compact(item).lower() for item in issues if _compact(item)))
    return hashlib.sha256("\n".join(normalized).encode("utf-8")).hexdigest()


class CancelStaleGenerationMiddleware(BaseMiddleware):
    """Cancel the previous in-flight update for a chat before handling the new one."""

    @staticmethod
    def _identity(event: TelegramObject) -> tuple[int | None, str]:
        if isinstance(event, Message):
            return event.chat.id, f"message:{event.message_id}"
        if isinstance(event, CallbackQuery) and event.message is not None:
            return event.message.chat.id, f"callback:{event.id}"
        return None, ""

    async def __call__(
        self,
        handler: Callable[[TelegramObject, dict[str, Any]], Awaitable[Any]],
        event: TelegramObject,
        data: dict[str, Any],
    ) -> Any:
        chat_id, event_id = self._identity(event)
        if chat_id is None:
            return await handler(event, data)

        current = asyncio.current_task()
        if current is None:
            return await handler(event, data)

        async with _ACTIVE_LOCK:
            previous = _ACTIVE_UPDATES.get(chat_id)
            if previous is not None and previous is not current and not previous.done():
                previous.cancel()
                LOGGER.info("STALE_RUN_CANCELLED chat_id=%s before=%s", chat_id, event_id)
            _ACTIVE_UPDATES[chat_id] = current

        run_token = _RUN_ID.set(f"{chat_id}:{event_id}")
        repair_token = _REPAIR_BLOCKER_SETS.set(frozenset())
        delivery_token = _DELIVERED_KINDS.set(frozenset())
        try:
            return await handler(event, data)
        except asyncio.CancelledError:
            LOGGER.info("STALE_RUN_STOPPED run=%s", _RUN_ID.get())
            return None
        finally:
            _RUN_ID.reset(run_token)
            _REPAIR_BLOCKER_SETS.reset(repair_token)
            _DELIVERED_KINDS.reset(delivery_token)
            async with _ACTIVE_LOCK:
                if _ACTIVE_UPDATES.get(chat_id) is current:
                    _ACTIVE_UPDATES.pop(chat_id, None)


async def _repair_pretrial_once(service: Any, case_context: str, research: LegalResearch, draft: Any, issues: list[str], language: str) -> Any:
    from korgan.pretrial import PretrialDraft, _PRETRIAL_SCHEMA, normalize_pretrial

    current = {
        "title": draft.title,
        "sender": draft.sender,
        "recipient": draft.recipient,
        "facts": draft.facts,
        "legal_basis": draft.legal_basis,
        "demands": draft.demands,
        "deadline": draft.deadline,
        "consequences": draft.consequences,
        "attachments": draft.attachments,
        "verification_notes": draft.verification_notes,
    }
    repair_method = getattr(service, "_quality_repair", None)
    if repair_method is None:
        return draft
    payload = await repair_method(
        schema_name="korgan_pretrial_invariant_repair",
        schema=_PRETRIAL_SCHEMA,
        case_context=case_context,
        research=research,
        current_payload=current,
        issues=issues,
        language=language,
        document_label="досудебную претензию после production quality gate",
        extra_rules=(
            "8. Исправь именно внутренние дефекты формулировок. Не проси пользователя исправлять созданный KORGAN пересказ нормы.\n"
            "9. Если VERIFIED не позволяет утверждать правовой вывод, не выдумывай норму: оставь нейтральную формулировку, а runtime добавит [СВЕРИТЬ].\n"
            "10. Денежное требование не заменяй фразой 'произвести расчет': арифметику выполняет deterministic money ledger."
        ),
    )
    repaired = PretrialDraft(status=research.status, source_urls=list(research.source_urls), **payload)
    normalize_pretrial(repaired)
    return repaired


def _install_repair_progress_guard(service_cls: type[Any]) -> None:
    original = getattr(service_cls, "_quality_repair", None)
    if original is None or getattr(service_cls, "_korgan_invariant_repair_guard", False):
        return

    async def guarded(self: Any, *args: Any, **kwargs: Any) -> dict[str, Any]:
        issues = [str(x) for x in kwargs.get("issues", []) or []]
        signature = _repair_signature(issues)
        seen = _REPAIR_BLOCKER_SETS.get()
        if signature in seen:
            LOGGER.warning(
                "REPAIR_NO_PROGRESS_STOP run=%s blockers=%s",
                _RUN_ID.get(),
                issues[:6],
            )
            current = kwargs.get("current_payload")
            return dict(current) if isinstance(current, dict) else {}
        _REPAIR_BLOCKER_SETS.set(frozenset((*seen, signature)))
        return await original(self, *args, **kwargs)

    service_cls._quality_repair = guarded
    service_cls._korgan_invariant_repair_guard = True


def _install_high_context_research(service_cls: type[Any]) -> None:
    if getattr(service_cls, "_korgan_invariant_structured", False):
        return
    original_structured = service_cls._structured_response

    async def structured(self: Any, *args: Any, **kwargs: Any):
        schema_name = str(kwargs.get("schema_name", ""))
        tools = kwargs.get("tools")
        if schema_name == "korgan_fast_professional_rk_research" and tools:
            rewritten = copy.deepcopy(tools)
            for tool in rewritten:
                if isinstance(tool, dict) and tool.get("type") == "web_search":
                    tool["search_context_size"] = "high"
            kwargs["tools"] = rewritten
            LOGGER.info("COST_SPEED research web_context=high reason=production_invariant_v2")
        return await original_structured(self, *args, **kwargs)

    service_cls._structured_response = structured
    service_cls._korgan_invariant_structured = True


def _install_research_determinism(service_cls: type[Any]) -> None:
    if getattr(service_cls, "_korgan_invariant_research", False):
        return
    original_research = service_cls.research_case

    async def research_case(self: Any, case_context: str, language: str = "ru") -> LegalResearch:
        key = _research_key(case_context, language)
        cached = _RESEARCH_CACHE.get(key)
        if cached is not None:
            LOGGER.info("RESEARCH_DETERMINISTIC_CACHE hit=1 key=%s", key[:12])
            return copy.deepcopy(cached)
        result = await original_research(self, case_context, language=language)
        result = canonicalize_research(result)
        if len(_RESEARCH_CACHE) >= _RESEARCH_CACHE_MAX:
            _RESEARCH_CACHE.pop(next(iter(_RESEARCH_CACHE)))
        _RESEARCH_CACHE[key] = copy.deepcopy(result)
        LOGGER.info("RESEARCH_DETERMINISTIC_CACHE hit=0 key=%s", key[:12])
        return result

    service_cls.research_case = research_case
    service_cls._korgan_invariant_research = True


def _install_finalization_before_repair() -> None:
    from korgan import fast_professional_litigation as litigation
    from korgan.professional_claim_finalizer import finalize_professional_claim

    if getattr(litigation, "_korgan_invariant_prefinalize", False):
        return
    original_preqa = litigation._deterministic_pre_qa

    def preqa(case_context: str, research: LegalResearch, draft: ClaimDraft) -> None:
        before = list(draft.legal_basis)
        finalize_professional_claim(case_context, research, draft)
        after = list(draft.legal_basis)
        removed = [line for line in before if line not in after]
        added = [line for line in after if line not in before]
        if added and removed:
            LOGGER.info(
                "CLAIM_MATERIAL_LAW_RESCUE added=%s removed=%s rewritten=[]",
                added[:6],
                removed[:6],
            )
        elif added:
            LOGGER.info("CLAIM_MATERIAL_LAW_ENRICH added=%s", added[:6])
        original_preqa(case_context, research, draft)

    litigation._deterministic_pre_qa = preqa
    litigation._korgan_invariant_prefinalize = True


def _install_pretrial_runtime() -> None:
    from korgan import bot as base_bot
    from korgan import pretrial_runtime as runtime
    from korgan.pretrial import build_pretrial_docx, pretrial_quality_issues
    from korgan.ui import main_menu

    async def generate(message: Message, state: Any) -> None:
        await runtime._save_text(message, state)
        lang = await runtime._lang(state)
        context = await base_bot._case_context(state)
        menu = main_menu(lang)

        if not context.strip():
            await state.update_data(mode="pretrial_waiting")
            await message.answer(
                "Чтобы подготовить досудебную претензию, опишите ситуацию одним сообщением или приложите документы."
                if lang != "kk" else
                "Сотқа дейінгі талапты дайындау үшін жағдайды бір хабарламада сипаттаңыз немесе құжаттарды тіркеңіз.",
                reply_markup=menu,
            )
            return

        service = base_bot.service
        research_method = getattr(service, "research_pretrial", None) if service is not None else None
        draft_method = getattr(service, "draft_pretrial", None) if service is not None else None
        if research_method is None or draft_method is None:
            await message.answer("Модуль досудебной претензии не загружен.", reply_markup=menu)
            return

        await state.update_data(mode="main")
        await message.answer(
            "Формирую досудебную претензию и проверяю правовую основу…"
            if lang != "kk" else "Сотқа дейінгі талапты дайындап, құқықтық негізін тексеріп жатырмын…",
            reply_markup=menu,
        )
        await message.bot.send_chat_action(message.chat.id, "typing")

        try:
            research = await research_method(context, language=lang)
            draft = await draft_method(context, research, language=lang)
            apply_money_ledger_to_pretrial(context, draft)
            issues_before = pretrial_quality_issues(draft, research)
            user_before, internal_before = classify_issues(issues_before)

            if user_before:
                LOGGER.warning(
                    "UNIVERSAL_WORD_QUALITY kind=pretrial issues_after=%d blocker_class=NEEDS_USER_DATA",
                    len(issues_before),
                )
                await message.answer(_user_data_message("досудебная претензия", user_before), reply_markup=menu)
                return

            if internal_before:
                draft = await _repair_pretrial_once(
                    service, context, research, draft, [x.text for x in internal_before], lang
                )
                apply_money_ledger_to_pretrial(context, draft)

            issues_after = pretrial_quality_issues(draft, research)
            user_after, internal_after = classify_issues(issues_after)
            if user_after:
                LOGGER.warning(
                    "UNIVERSAL_WORD_QUALITY kind=pretrial issues_after=%d blocker_class=NEEDS_USER_DATA",
                    len(issues_after),
                )
                await message.answer(_user_data_message("досудебная претензия", user_after), reply_markup=menu)
                return

            file_bytes = build_pretrial_docx(draft, language=lang)
            if internal_after:
                file_bytes = append_review_markers(file_bytes, [x.text for x in internal_after], language=lang)
                draft.status = VerificationStatus.NEEDS_VERIFICATION
            LOGGER.info(
                "UNIVERSAL_WORD_QUALITY kind=pretrial issues_before=%d issues_after=%d delivered=1 internal_markers=%d",
                len(issues_before),
                len(issues_after),
                len(internal_after),
            )
        except asyncio.CancelledError:
            raise
        except Exception:
            LOGGER.exception("Pretrial demand generation failed")
            await message.answer("Не удалось безопасно сформировать досудебную претензию.", reply_markup=menu)
            return

        if not _mark_delivery_once("pretrial"):
            return
        caption = "✅ Досудебная претензия сформирована в Word (.docx)."
        if internal_after:
            caption = (
                "⚠️ Досудебная претензия сформирована с явными отметками [СВЕРИТЬ].\n\n"
                "Что осталось проверить:\n" + _internal_message(internal_after)
            )
        filename = "KORGAN_sotqa_deyingi_talap.docx" if lang == "kk" else "KORGAN_dosudebnaya_pretenziya.docx"
        await message.answer_document(
            BufferedInputFile(file_bytes, filename=filename),
            caption=caption,
            reply_markup=menu,
        )

    runtime._generate = generate


def _install_claim_runtime() -> None:
    from korgan import bot as base_bot
    from korgan import universal_claim_runtime as runtime
    from korgan.claim_docx import build_claim_docx
    from korgan.claim_failure import ClaimStage, failure_from_exception
    from korgan.document_quality import assess_document_quality, rendered_docx_blockers
    from korgan.telegram_text import fit_caption

    async def send_claim(message: Message, state: Any, *, context: str, research: LegalResearch, draft: ClaimDraft) -> None:
        # Money propagation is checked before release/scoring.  A monetary input
        # can therefore never disappear merely because the drafting model omitted
        # it from the prayer for relief.
        apply_money_ledger_to_claim(context, draft)

        fit = runtime.enforce_legal_basis_fit(draft)
        if fit:
            for item in fit:
                note = f"Правовое основание требует проверки: {item}"
                if note not in draft.verification_notes:
                    draft.verification_notes.append(note)
            draft.status = VerificationStatus.NEEDS_VERIFICATION

        release = runtime._downgrade_unverified_citations_live(draft, research)
        release_issues = [x.as_note() for x in release.citations.blocking] + [x.as_note() for x in release.integrity]
        quality = assess_document_quality("claim", context, research, draft)
        all_issues = list(dict.fromkeys([*release_issues, *quality.repair_issues()]))
        user_issues, internal_issues = classify_issues(all_issues)

        if user_issues:
            LOGGER.warning(
                "UNIVERSAL_WORD_QUALITY kind=claim issues_after=%d blocker_class=NEEDS_USER_DATA",
                len(all_issues),
            )
            await message.answer(_user_data_message("исковое заявление", user_issues), reply_markup=base_bot.MENU)
            return

        # Internal defects never become an impossible instruction to the user.
        # The service already had its targeted repair pass; any residue is
        # disclosed in-document instead of triggering a dead-end gate.
        for item in internal_issues:
            marker = review_marker(item.text)
            if marker not in draft.verification_notes:
                draft.verification_notes.append(marker)
        if internal_issues:
            draft.status = VerificationStatus.NEEDS_VERIFICATION

        try:
            file_bytes = build_claim_docx(draft)
        except Exception as exc:
            await base_bot._report_claim_failure(message, failure_from_exception(exc, stage=ClaimStage.RENDER))
            return

        export_blockers = rendered_docx_blockers(file_bytes, ready_expected=quality.ready and not internal_issues)
        if export_blockers:
            export_classified = [classify_issue(f"экспорт Word: {item}") for item in export_blockers]
            internal_issues.extend(x for x in export_classified if x.blocker_class == BlockerClass.INTERNAL_QUALITY)
        if internal_issues:
            file_bytes = append_review_markers(file_bytes, [x.text for x in internal_issues])

        LOGGER.info(
            "UNIVERSAL_WORD_QUALITY kind=claim issues_after=%d delivered=1 internal_markers=%d",
            len(all_issues) + len(export_blockers),
            len(internal_issues),
        )
        await state.update_data(mode="main", gate_issues=[], claim_draft=None, pending_fields=[])
        if not _mark_delivery_once("claim"):
            return

        if internal_issues:
            caption = (
                f"⚠️ KORGAN QUALITY {quality.score:.1f}/10 · ПРОЕКТ С ОТМЕТКАМИ [СВЕРИТЬ]\n"
                "Система не скрыла оставшиеся внутренние замечания:\n"
                + _internal_message(internal_issues)
            )
        else:
            caption = f"✅ KORGAN QUALITY {quality.score:.1f}/10\nИск сформирован в Word (.docx)."
        await message.answer_document(
            BufferedInputFile(file_bytes, filename="KORGAN_iskovoe_zayavlenie.docx"),
            caption=fit_caption(caption),
            reply_markup=base_bot.MENU,
        )

    runtime._send_claim = send_claim


def install_production_invariants_v2() -> None:
    """Install the cross-cutting invariant layer on the strict production runtime."""
    from korgan.pretrial import PretrialProductionService

    if getattr(PretrialProductionService, "_korgan_invariants_v2_installed", False):
        return

    _install_repair_progress_guard(PretrialProductionService)
    _install_high_context_research(PretrialProductionService)
    _install_research_determinism(PretrialProductionService)
    _install_finalization_before_repair()
    _install_pretrial_runtime()
    _install_claim_runtime()
    PretrialProductionService._korgan_invariants_v2_installed = True
    LOGGER.info("Installed KORGAN production invariants v2 (I1-I10 runtime layer)")

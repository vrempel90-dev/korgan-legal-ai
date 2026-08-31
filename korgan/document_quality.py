from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Any, Literal

from docx import Document

from korgan.contract_preamble import preamble_defects
from korgan.contract_type_safety import misclassification_blockers
from korgan.document_release import review_lines
from korgan.legal_basis_fit import enforce_legal_basis_fit
from korgan.legal_calc import parse_all_amounts_kzt
from korgan.legal_types import ClaimDraft, ContractDraft, LegalResearch
from korgan.response_types import ResponseToClaimDraft
from korgan.text_integrity import integrity_findings

DocumentKind = Literal["claim", "contract", "response_to_claim", "pretrial", "pretrial_response"]
MIN_READY_SCORE = 8.5

_PLACEHOLDER_RE = re.compile(
    r"\[(?:ТРЕБУЕТ УТОЧНЕНИЯ|ТРЕБУЕТ ПРОВЕРКИ|ТРЕБУЕТ РАСЧ[ЕЁ]ТА|ТРЕБУЕТ ДОБАВИТЬ)[^\]]*\]",
    re.IGNORECASE,
)
_ARTICLE_RE = re.compile(r"(?i)(?:стать(?:я|и|е|ю|ёй|ей)|ст\.)\s*\d+(?:-\d+)?")
_ENTITY_RE = re.compile(
    r"\b(?:ТОО|АО|РГП|РГУ|КГУ|КГП|ИП|ЖК|КСК|ОО)\b|\bБИН\b|"
    r"товариществ\w*\s+с\s+ограниченн\w*\s+ответственност\w*|"
    r"акционерн\w*\s+обществ\w*",
    re.IGNORECASE,
)
_SERVICE_MARKERS = (
    "korgan qa status",
    "preliminary draft",
    "lawyer-review draft",
    "официальные источники",
    "могу доработать",
    "если нужно",
    "###",
    "**",
    # Внутренняя терминология конвейера. Она нужна внутри проверок и в
    # verification_notes, но в теле документа, который читает суд или
    # контрагент, её быть не может.
    "needs_verification",
    "korgan quality",
    "senior_preflight_score",
    "filing_action",
    "legal_grounding",
    "korgan pipeline",
)
_GENERIC_COURT_MARKERS = (
    "требует уточнения",
    "наименование суда",
    "компетентного рассматривать",
    "компетентный суд",
    "по месту нахождения",
    "по месту жительства",
    "надлежащий суд",
)


@dataclass(slots=True)
class DocumentQualityReport:
    kind: DocumentKind
    score: float
    hard_blockers: list[str] = field(default_factory=list)
    issues: list[str] = field(default_factory=list)
    category_scores: dict[str, float] = field(default_factory=dict)

    @property
    def ready(self) -> bool:
        return self.score >= MIN_READY_SCORE and not self.hard_blockers

    def repair_issues(self, limit: int = 12) -> list[str]:
        return list(dict.fromkeys([*self.hard_blockers, *self.issues]))[:limit]


def _clean_lines(values: Any) -> list[str]:
    if not values:
        return []
    if isinstance(values, str):
        return [values.strip()] if values.strip() else []
    result: list[str] = []
    for value in values:
        text = str(value or "").strip()
        if text:
            result.append(text)
    return result


def _claim_lines(draft: ClaimDraft) -> list[str]:
    return _clean_lines(
        [
            draft.title,
            draft.court,
            *draft.claimant,
            *draft.defendant,
            draft.price_of_claim,
            draft.state_duty,
            draft.late_interest,
            *draft.facts,
            *draft.legal_basis,
            *draft.requests,
            *draft.attachments,
        ]
    )


def _claim_citation_lines(draft: ClaimDraft) -> list[str]:
    """Model/legal prose only; deterministic state duty is checked separately."""
    return _clean_lines(
        [
            draft.title,
            *draft.facts,
            *draft.legal_basis,
            draft.late_interest,
            *draft.requests,
            *draft.attachments,
        ]
    )


def _contract_lines(draft: ContractDraft) -> list[str]:
    return _clean_lines(draft.body_lines())


def _response_lines(draft: ResponseToClaimDraft) -> list[str]:
    return _clean_lines(draft.body_lines())


def document_lines(kind: DocumentKind, draft: Any) -> list[str]:
    if kind == "claim":
        return _claim_lines(draft)
    if kind == "contract":
        return _contract_lines(draft)
    if kind == "response_to_claim":
        return _response_lines(draft)
    raise ValueError(f"Unsupported document kind: {kind}")


def _common_hygiene(
    kind: DocumentKind,
    lines: list[str],
    blockers: list[str],
    issues: list[str],
    *,
    case_context: str = "",
    verified_claims: list[str] | None = None,
    verification_notes: list[str] | None = None,
) -> float:
    text = "\n".join(lines)
    lowered = text.lower()
    score = 1.0

    placeholders = _PLACEHOLDER_RE.findall(text)
    if placeholders:
        blockers.append(f"в документе остались незаполненные обязательные/проверочные поля: {len(placeholders)}")
        score -= min(0.6, 0.12 * len(placeholders))

    for marker in _SERVICE_MARKERS:
        if marker in lowered:
            blockers.append(f"в юридическом тексте осталась служебная фраза: {marker}")
            score -= 0.4
            break

    # Пометки проверки — это ПРОДУКТ работы, а не её дефект: ими документ
    # сообщает юристу, что сверить перед подачей («FILING_ACTION: указать
    # банковские реквизиты истца», «сверить редакцию статьи»). Пока они были
    # жёстким блокером, выпуск был невозможен в принципе: у любого реального
    # дела такая пометка есть хотя бы одна, а hard blocker обнуляет ready
    # независимо от оценки. Промпт составления прямо просит модель их писать —
    # то есть чем добросовестнее работал документ, тем вернее он не выходил.
    #
    # Теперь это замечание со снижением оценки. Пользователь видит пометки в
    # самом документе: _document_status ставит на него штамп LAWYER-REVIEW
    # DRAFT и подвал «перед подачей необходимо проверить…». Настоящие дефекты
    # (незаполненные поля, служебный текст, неподтверждённые ссылки,
    # повреждение текста) остаются блокирующими и проверяются рядом.
    unresolved = _clean_lines(verification_notes)
    if unresolved:
        issues.append("вопросы к проверке перед подачей: " + unresolved[0][:180])
        score -= 0.35

    # Ошибка в виде договора обесценивает весь раздел правового обоснования:
    # нормы будут реальными и процитированными верно, но не о тех отношениях.
    for finding in misclassification_blockers(case_context, lines):
        blockers.append(finding)
        score -= 0.6

    integrity = integrity_findings(text)
    if integrity:
        blockers.append("нарушена целостность текста: " + integrity[0].description)
        score -= 0.6

    release = review_lines(lines, verified_claims=verified_claims)
    if release.citations.blocking:
        blockers.append("есть правовая ссылка, не прошедшая source-bound/corpus проверку")
        score -= 0.6
    if release.integrity:
        blockers.append("release-check обнаружил повреждение текста")
        score -= 0.6

    return max(0.0, score)


# Возражения, которые нельзя заявлять «на всякий случай»: исковая давность и
# процессуальные нарушения. Каждое из них либо подтверждено конкретными датами
# и нормой, либо не заявляется вовсе — иначе документ раздувается доводами,
# которые оппонент разобьёт первым же абзацем.
_UNSUPPORTABLE_OBJECTION_RE = re.compile(
    r"(?i)(?:исков\w*\s+давност\w*|срок\w*\s+давност\w*|"
    r"пропущен\w*\s+срок|"
    r"нарушен\w*\s+(?:процессуальн\w*|порядок\s+подач\w*)|"
    r"подсудност\w*\s+наруш\w*|"
    r"талап\s+қою\s+мерзім\w*)"
)

# Опора, без которой такое возражение не проверяемо: конкретная дата, период
# или названная норма.
# Длительности («3 года») здесь намеренно нет: она не опора. Довод об
# исковой давности проверяем только тогда, когда названы даты, из которых
# срок вычисляется, либо норма, которая его устанавливает. Пока длительность
# засчитывалась, «Истёк срок исковой давности — 3 года» проходило шлюз,
# не сообщая ни того, ни другого.
_OBJECTION_ANCHOR_RE = re.compile(
    r"(?i)(?:\d{2}\.\d{2}\.\d{4}|"
    r"\b\d{4}\s*год|"
    r"\bстать\w*\s*\d|\bст\.\s*\d)"
)


def unsupported_objections(objections: list[str]) -> list[str]:
    """Возражения об исковой давности/процессуальных нарушениях без опоры.

    Опора требуется в самом возражении, а не где-то ещё в документе. Довод об
    исковой давности профессионально всегда называет даты, из которых срок
    вычисляется: когда началось течение и когда истекло. Дата, случайно
    оказавшаяся в соседнем разделе, такой довод не подтверждает — иначе любое
    возражение «на всякий случай» проходило бы за счёт чужих фактов.
    """
    findings: list[str] = []
    for item in _clean_lines(objections):
        if not _UNSUPPORTABLE_OBJECTION_RE.search(item):
            continue
        if _OBJECTION_ANCHOR_RE.search(item):
            continue
        findings.append(
            "возражение заявлено без подтверждающих дат или нормы: " + item[:120]
        )
    return findings


def _preserve_known_identifiers(case_context: str, lines: list[str], blockers: list[str]) -> float:
    score = 1.0
    body = "\n".join(lines)
    for value in sorted(set(re.findall(r"(?<!\d)\d{12}(?!\d)", case_context or ""))):
        if value not in body:
            blockers.append(f"потерян известный ИИН/БИН {value}")
            score -= 0.3
    return max(0.0, score)


def _normalized(value: str) -> str:
    return re.sub(r"[^0-9a-zа-яё]+", "", (value or "").lower())


def _court_supported(case_context: str, research: LegalResearch, court: str) -> bool:
    """Court name must come from user materials or verified official research."""
    norm = _normalized(court)
    if len(norm) < 8:
        return False
    if norm in _normalized(case_context):
        return True
    for note in research.notes:
        if not str(note).startswith("VERIFIED_COURT:"):
            continue
        verified = str(note).split(":", 1)[1].strip()
        verified_norm = _normalized(verified)
        if verified_norm and (verified_norm in norm or norm in verified_norm):
            return True
    return False


def _court_is_concrete(court: str) -> bool:
    lowered = (court or "").lower()
    if not lowered or "суд" not in lowered:
        return False
    if _PLACEHOLDER_RE.search(court):
        return False
    return not any(marker in lowered for marker in _GENERIC_COURT_MARKERS)


def _score_claim(case_context: str, research: LegalResearch, draft: ClaimDraft) -> DocumentQualityReport:
    blockers: list[str] = []
    issues: list[str] = []
    lines = _claim_lines(draft)

    parties = 2.0
    if not _clean_lines(draft.claimant):
        blockers.append("не заполнены данные истца")
        parties -= 0.8
    if not _clean_lines(draft.defendant):
        blockers.append("не заполнены данные ответчика")
        parties -= 0.8
    parties -= 1.0 - _preserve_known_identifiers(case_context, lines, blockers)

    defendant_text = "\n".join(draft.defendant)
    claimant_text = "\n".join(draft.claimant)
    if _ENTITY_RE.search(defendant_text) and re.search(r"фио\s+ответчика", defendant_text, re.I):
        blockers.append("у юридического лица ошибочно запрошено ФИО")
        parties -= 0.5
    if not _ENTITY_RE.search(claimant_text) and re.search(r"банковск\w*\s+реквизит", claimant_text, re.I):
        blockers.append("банковские реквизиты ошибочно сделаны обязательными для истца-физлица")
        parties -= 0.35

    facts = 2.0
    if len(_clean_lines(draft.facts)) < 3:
        issues.append("фактическая часть недостаточно раскрывает хронологию и нарушение")
        facts -= 0.5
    if not _clean_lines(draft.requests):
        blockers.append("отсутствует просительная часть")
        facts -= 1.0
    if not (draft.price_of_claim or "").strip() and re.search(r"\b(?:тенге|₸|тг)\b", case_context, re.I):
        blockers.append("не указана цена денежного иска при известной сумме")
        facts -= 0.5

    law = 2.5
    basis = "\n".join(draft.legal_basis)
    if not basis.strip():
        blockers.append("отсутствует правовое обоснование")
        law = 0.0
    else:
        if research.verified_claims and not _ARTICLE_RE.search(basis):
            blockers.append("есть VERIFIED-нормы, но документ не содержит конкретной статьи")
            law -= 1.0
        fit = enforce_legal_basis_fit(draft)
        if fit:
            blockers.append("правовое основание не поддерживает заявленное требование")
            issues.extend(str(item) for item in fit[:4])
            law -= min(1.5, 0.65 + 0.2 * len(fit))
        if not research.verified_claims:
            blockers.append("нет source-bound подтвержденной материально-правовой основы")
            law -= 1.2

    procedure = 1.5
    court = (draft.court or "").strip()
    if not _court_is_concrete(court):
        blockers.append("не определено конкретное наименование суда")
        procedure -= 0.75
    elif not _court_supported(case_context, research, court):
        blockers.append("наименование суда не подтверждено материалами дела или официальным source-bound исследованием")
        procedure -= 0.55

    duty = (draft.state_duty or "").strip()
    if not duty or _PLACEHOLDER_RE.search(duty):
        blockers.append("не определена госпошлина или подтвержденная льгота")
        procedure -= 0.5

    evidence = 1.0
    attachments = "\n".join(draft.attachments).lower()
    for marker in ("договор", "квитанц", "претензи", "расписк", "акт"):
        if marker in (case_context or "").lower() and marker not in attachments:
            issues.append(f"в приложениях потеряно упомянутое доказательство: {marker}")
            evidence -= 0.18

    hygiene = _common_hygiene(
        "claim",
        _claim_citation_lines(draft),
        blockers,
        issues,
        case_context=case_context,
        verified_claims=research.verified_claims,
        verification_notes=draft.verification_notes,
    )

    categories = {
        "fact_role_lock": max(0.0, parties),
        "facts_relief": max(0.0, facts),
        "legal_basis": max(0.0, law),
        "procedure": max(0.0, procedure),
        "evidence": max(0.0, evidence),
        "hygiene": hygiene,
    }
    score = round(sum(categories.values()), 1)
    if blockers:
        score = min(score, 8.4)
    return DocumentQualityReport("claim", score, list(dict.fromkeys(blockers)), list(dict.fromkeys(issues)), categories)


def _score_contract(case_context: str, research: LegalResearch, draft: ContractDraft) -> DocumentQualityReport:
    blockers: list[str] = []
    issues: list[str] = []
    lines = _contract_lines(draft)

    identity = 2.0
    if not _clean_lines(draft.party_a) or not _clean_lines(draft.party_b):
        blockers.append("не идентифицированы обе стороны договора")
        identity -= 1.0
    for defect in preamble_defects(draft.preamble):
        blockers.append("дефект преамбулы: " + str(defect))
        identity -= 0.25
    identity -= 1.0 - _preserve_known_identifiers(case_context, lines, blockers)

    terms = 3.0
    if not draft.sections:
        blockers.append("отсутствуют условия договора")
        terms -= 2.0
    clause_count = sum(len(section.clauses) for section in draft.sections)
    if clause_count < 5:
        issues.append("договор недостаточно полно раскрывает права, обязанности и исполнение")
        terms -= 0.6
    if not (draft.title or draft.contract_type or "").strip():
        blockers.append("не определен вид/название договора")
        terms -= 0.5

    legal = 2.0
    if not research.verified_claims or not research.source_urls:
        blockers.append("правовая конструкция договора не подтверждена source-bound исследованием")
        legal -= 1.2

    execution = 1.0
    if not _clean_lines(draft.requisites_a) or not _clean_lines(draft.requisites_b):
        blockers.append("нет реквизитов/подписного блока обеих сторон")
        execution -= 0.6
    if not (draft.place_and_date or "").strip():
        blockers.append("не заполнены место/дата заключения договора")
        execution -= 0.4

    hygiene = _common_hygiene(
        "contract",
        lines,
        blockers,
        issues,
        case_context=case_context,
        verified_claims=research.verified_claims,
        verification_notes=draft.verification_notes,
    )
    fact_lock = _preserve_known_identifiers(case_context, lines, blockers)

    categories = {
        "identity_preamble": max(0.0, identity),
        "essential_terms": max(0.0, terms),
        "legal_validity": max(0.0, legal),
        "execution": max(0.0, execution),
        "fact_lock": fact_lock,
        "hygiene": hygiene,
    }
    score = round(sum(categories.values()), 1)
    if blockers:
        score = min(score, 8.4)
    return DocumentQualityReport("contract", score, list(dict.fromkeys(blockers)), list(dict.fromkeys(issues)), categories)


def _score_response(case_context: str, research: LegalResearch, draft: ResponseToClaimDraft) -> DocumentQualityReport:
    blockers: list[str] = []
    issues: list[str] = []
    lines = _response_lines(draft)

    identity = 2.0
    court = (draft.court or "").strip()
    if not _court_is_concrete(court):
        blockers.append("в отзыве не указан конкретный суд")
        identity -= 0.6
    elif not _court_supported(case_context, research, court):
        blockers.append("суд в отзыве не подтвержден исходным иском/материалами")
        identity -= 0.4
    if not _clean_lines(draft.claimant) or not _clean_lines(draft.defendant):
        blockers.append("в отзыве не идентифицированы истец и ответчик")
        identity -= 0.8
    identity -= 1.0 - _preserve_known_identifiers(case_context, lines, blockers)

    position = 3.0
    if not _clean_lines(draft.claim_summary):
        blockers.append("не отражены требования исходного иска")
        position -= 0.8
    if not _clean_lines(draft.position):
        blockers.append("не сформулирована позиция ответчика")
        position -= 0.8
    if not draft.objections:
        blockers.append("нет возражений по существу требований")
        position -= 1.0
    if not _clean_lines(draft.requests):
        blockers.append("нет процессуальной просительной части")
        position -= 0.4
    # Схема разрешает вынести даты и норму в subclauses/prose, а в text
    # оставить заголовок довода. Проверять один заголовок значит блокировать
    # полностью обоснованное возражение за то, что опора лежит строкой ниже.
    objection_texts = ["\n".join(item.body_lines()) for item in draft.objections] if draft.objections else []
    unsupported = unsupported_objections(objection_texts)
    if unsupported:
        blockers.extend(unsupported)
        position -= 0.6

    # Схема требует ключ calculation_review, но не его содержимое: пустой
    # массив ей соответствует. Без этой проверки отзыв на денежный иск
    # выходил 10/10, не разобрав расчёт истца, и раунд правки не запускался.
    # Тот же критерий, что у ответа на претензию: разбор нужен там, где
    # оппонент предъявил сумму.
    if parse_all_amounts_kzt("\n".join(draft.claim_summary)) and not _clean_lines(
        getattr(draft, "calculation_review", [])
    ):
        blockers.append("расчёт истца не разобран")
        position -= 0.5

    law = 2.5
    basis = "\n".join(draft.legal_basis)
    if not basis.strip():
        blockers.append("у отзыва отсутствует правовое обоснование")
        law -= 1.3
    if research.verified_claims and not _ARTICLE_RE.search(basis):
        blockers.append("VERIFIED-нормы не перенесены в правовое обоснование отзыва")
        law -= 0.8
    if not research.verified_claims or not research.source_urls:
        blockers.append("правовая позиция отзыва не подтверждена source-bound исследованием")
        law -= 1.0

    evidence = 0.5
    if not draft.attachments:
        issues.append("не сформирован перечень приложений к отзыву")
        evidence -= 0.2

    completeness = 1.0
    if not draft.case_number:
        issues.append("номер дела не указан; допустимо только если он действительно неизвестен")
        completeness -= 0.15

    hygiene = _common_hygiene(
        "response_to_claim",
        lines,
        blockers,
        issues,
        case_context=case_context,
        verified_claims=research.verified_claims,
        verification_notes=draft.verification_notes,
    )

    categories = {
        "identity_procedure": max(0.0, identity),
        "position_objections": max(0.0, position),
        "legal_basis": max(0.0, law),
        "evidence": max(0.0, evidence),
        "completeness": max(0.0, completeness),
        "hygiene": hygiene,
    }
    score = round(sum(categories.values()), 1)
    if blockers:
        score = min(score, 8.4)
    return DocumentQualityReport("response_to_claim", score, list(dict.fromkeys(blockers)), list(dict.fromkeys(issues)), categories)


def _score_pretrial(case_context: str, research: LegalResearch, draft: Any) -> DocumentQualityReport:
    """Численная оценка досудебной претензии.

    До этого претензия была единственным клиентским документом без порога:
    проверялся только список замечаний, поэтому «готова» и «сойдёт» ничем не
    различались. Категории повторяют то, что читает адресат: кто кому пишет,
    из чего возник долг, чем требование обосновано в праве, как получена сумма,
    что именно требуется и к какому сроку.
    """
    from korgan.pretrial import has_money_demand, pretrial_quality_issues

    blockers: list[str] = []
    issues: list[str] = []
    lines = _clean_lines(draft.body_lines())

    identity = 2.0
    if not _clean_lines(draft.sender):
        blockers.append("в претензии не указан отправитель")
        identity -= 1.0
    if not _clean_lines(draft.recipient):
        blockers.append("в претензии не указан адресат")
        identity -= 1.0
    identity -= 1.0 - _preserve_known_identifiers(case_context, lines, blockers)

    facts = 2.0
    if not _clean_lines(draft.facts):
        blockers.append("не изложено фактическое основание требований")
        facts -= 1.5

    law = 2.5
    basis = "\n".join(draft.legal_basis)
    if not basis.strip():
        blockers.append("у претензии отсутствует правовое обоснование")
        law -= 1.5
    elif research.verified_claims and not _ARTICLE_RE.search(basis):
        blockers.append("VERIFIED-нормы не перенесены в правовое обоснование претензии")
        law -= 1.0

    calculation = 1.5
    if has_money_demand(draft) and not _clean_lines(draft.calculation):
        blockers.append("денежное требование не раскрыто расчётом")
        calculation -= 1.5

    demand = 1.0
    if not _clean_lines(draft.demands):
        blockers.append("в претензии нет сформулированного требования")
        demand -= 0.7
    if not str(draft.deadline or "").strip():
        blockers.append("не указан срок добровольного исполнения требований")
        demand -= 0.3

    for issue in pretrial_quality_issues(draft, research):
        if issue not in blockers:
            blockers.append(issue)

    hygiene = _common_hygiene(
        "pretrial",
        lines,
        blockers,
        issues,
        case_context=case_context,
        verified_claims=research.verified_claims,
        verification_notes=draft.verification_notes,
    )

    categories = {
        "identity": max(0.0, identity),
        "facts": max(0.0, facts),
        "legal_basis": max(0.0, law),
        "calculation": max(0.0, calculation),
        "demand": max(0.0, demand),
        "hygiene": hygiene,
    }
    score = round(sum(categories.values()), 1)
    if blockers:
        score = min(score, 8.4)
    return DocumentQualityReport("pretrial", score, list(dict.fromkeys(blockers)), list(dict.fromkeys(issues)), categories)


def _score_pretrial_response(case_context: str, research: LegalResearch, draft: Any) -> DocumentQualityReport:
    """Численная оценка ответа на досудебную претензию.

    Ответ обязан разобрать требования контрагента по существу: отделить
    признаваемое от оспариваемого, проверить расчёт и обосновать позицию.
    Шаблонное «с требованиями не согласны» без разбора — не ответ.
    """
    from korgan.pretrial_response import money_claimed, pretrial_response_quality_issues

    blockers: list[str] = []
    issues: list[str] = []
    lines = _clean_lines(draft.body_lines())

    identity = 1.5
    if not _clean_lines(draft.sender):
        blockers.append("в ответе не указан отправитель")
        identity -= 0.75
    if not _clean_lines(draft.recipient):
        blockers.append("в ответе не указан адресат")
        identity -= 0.75
    identity -= 1.0 - _preserve_known_identifiers(case_context, lines, blockers)

    engagement = 3.0
    if not _clean_lines(draft.claim_summary):
        blockers.append("не отражены требования исходной претензии")
        engagement -= 1.2
    if not _clean_lines(draft.position):
        blockers.append("не сформулирована позиция получателя претензии")
        engagement -= 0.9
    if not _clean_lines(draft.objections) and not _clean_lines(draft.response_terms):
        blockers.append("нет содержательного ответа на требования претензии")
        engagement -= 0.9

    law = 2.0
    basis = "\n".join(draft.legal_basis)
    if research.verified_claims and not basis.strip():
        blockers.append("VERIFIED-нормы не перенесены в правовое обоснование ответа")
        law -= 1.2

    unsupported = unsupported_objections(list(draft.objections))
    if unsupported:
        blockers.extend(unsupported)
        engagement -= 0.6

    review = 1.5
    # Тот же предикат, что и в pretrial_response_quality_issues: у неденежной
    # претензии расчёта нет, и требовать его разбор значит понижать документ
    # за отсутствие раздела, которого в нём быть не должно.
    if (
        money_claimed(draft)
        and _clean_lines(draft.objections)
        and not _clean_lines(getattr(draft, "calculation_review", []))
    ):
        issues.append("расчёт контрагента не разобран построчно")
        review -= 0.5

    completeness = 1.0
    if not str(draft.reference or "").strip():
        issues.append("не указана ссылка на исходную претензию (дата/номер)")
        completeness -= 0.3

    for issue in pretrial_response_quality_issues(draft, research):
        if issue not in blockers:
            blockers.append(issue)

    hygiene = _common_hygiene(
        "pretrial_response",
        lines,
        blockers,
        issues,
        case_context=case_context,
        verified_claims=research.verified_claims,
        verification_notes=draft.verification_notes,
    )

    categories = {
        "identity": max(0.0, identity),
        "engagement": max(0.0, engagement),
        "legal_basis": max(0.0, law),
        "calculation_review": max(0.0, review),
        "completeness": max(0.0, completeness),
        "hygiene": hygiene,
    }
    score = round(sum(categories.values()), 1)
    if blockers:
        score = min(score, 8.4)
    return DocumentQualityReport(
        "pretrial_response", score, list(dict.fromkeys(blockers)), list(dict.fromkeys(issues)), categories
    )


def assess_document_quality(
    kind: DocumentKind,
    case_context: str,
    research: LegalResearch,
    draft: Any,
) -> DocumentQualityReport:
    if kind == "claim":
        return _score_claim(case_context, research, draft)
    if kind == "contract":
        return _score_contract(case_context, research, draft)
    if kind == "response_to_claim":
        return _score_response(case_context, research, draft)
    if kind == "pretrial":
        return _score_pretrial(case_context, research, draft)
    if kind == "pretrial_response":
        return _score_pretrial_response(case_context, research, draft)
    raise ValueError(f"Unsupported document kind: {kind}")


def docx_text(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            chunks.append(paragraph.text.strip())
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip():
                chunks.append(paragraph.text.strip())
        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                chunks.append(paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    chunks.append(cell.text.strip())
    return "\n".join(chunks)


def rendered_docx_blockers(data: bytes, *, ready_expected: bool) -> list[str]:
    """Final export check shared by every Word-producing document type."""
    text = docx_text(data)
    lowered = text.lower()
    blockers: list[str] = []

    if integrity_findings(text):
        blockers.append("финальный DOCX содержит поврежденный/склеенный текст")
    if ready_expected and _PLACEHOLDER_RE.search(text):
        blockers.append("финальный DOCX содержит незаполненные [ТРЕБУЕТ ...] поля")
    if ready_expected and ("preliminary draft" in lowered or "lawyer-review draft" in lowered):
        blockers.append("финальный DOCX помечен как предварительный, хотя прошел quality gate")
    if "http://" in lowered or "https://" in lowered or "###" in text or "**" in text:
        blockers.append("в финальный DOCX попал служебный URL/Markdown")
    return list(dict.fromkeys(blockers))

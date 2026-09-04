from __future__ import annotations

import asyncio
import io
import logging
import os
import re
import time
from dataclasses import dataclass
from typing import Any, Iterator

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table
from docx.text.paragraph import Paragraph

from korgan import citation_audit
from korgan import miniapp_api_v2 as core
from korgan.legal.corpus import (
    ACT_CONSUMER,
    ACT_GK_GENERAL,
    ACT_GK_SPECIAL,
    ACT_GPK,
    ACT_LABOR,
    ACT_TAX_DUTY,
)
from korgan.legal.corpus_refresh import fetch_adilet
from korgan.legal_types import VerificationStatus
from korgan.provision_check import _norm_claim_only, paraphrase_defects
from korgan.provision_corpus import normalize_text
from scripts.load_corpus import act_url, parse_provisions, strip_html

LOGGER = logging.getLogger(__name__)

FLAG_ENV = "KORGAN_LIVE_ARTICLE_VERIFY"
_CACHE_SECONDS = 60 * 30
_INSTALLED = False
_ORIGINAL_GENERATE = core._generate

# Civil-document production currently owns these official acts. A citation to a
# different statute is not guessed. The final Word is still delivered, but a
# citation that cannot be deterministically verified is removed from the
# filing-ready version and the result is downgraded to a review draft.
_ACT_CANDIDATES: dict[str, tuple[str, ...]] = {
    "ГПК РК": (ACT_GPK,),
    "ГК РК": (ACT_GK_GENERAL, ACT_GK_SPECIAL),
    "НК РК": (ACT_TAX_DUTY,),
    "ТК РК": (ACT_LABOR,),
    "ЗПП РК": (ACT_CONSUMER,),
}

_EDITION_RE = re.compile(r"Дата\s+редакции\s+(\d{2}\.\d{2}\.\d{4})", re.IGNORECASE)
_REPAIR_ATTEMPTS = 12


@dataclass(frozen=True)
class LiveAct:
    act_id: str
    source_url: str
    edition_date: str
    articles: dict[str, dict[str, str]]


_CACHE: dict[str, tuple[float, LiveAct]] = {}


class LiveArticleVerificationError(RuntimeError):
    pass


def live_article_verification_enabled() -> bool:
    return str(os.getenv(FLAG_ENV, "") or "").strip().lower() in {"1", "true", "yes", "on"}


def _docx_text(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        text = " ".join(str(paragraph.text or "").split()).strip()
        if text:
            chunks.append(text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = " ".join(str(cell.text or "").split()).strip()
                if text:
                    chunks.append(text)
    return "\n".join(chunks)


def _build_live_act(act_id: str, html: str, source_url: str) -> LiveAct:
    text = strip_html(html)
    parsed = parse_provisions(text)
    if not parsed:
        raise LiveArticleVerificationError(f"официальный акт {act_id} не содержит распознаваемых статей")

    articles: dict[str, dict[str, str]] = {}
    for item in parsed:
        article = articles.setdefault(item.article_no, {})
        key = str(item.item_no or "")
        body = " ".join(f"{item.heading} {item.body}".split()).strip()
        if key in article and body not in article[key]:
            article[key] = (article[key] + " " + body).strip()
        else:
            article[key] = body

    edition = ""
    match = _EDITION_RE.search(text)
    if match:
        edition = match.group(1)
    return LiveAct(act_id=act_id, source_url=source_url, edition_date=edition, articles=articles)


async def _live_act(act_id: str) -> LiveAct:
    cached = _CACHE.get(act_id)
    now = time.monotonic()
    if cached is not None and now - cached[0] <= _CACHE_SECONDS:
        return cached[1]

    url = act_url(act_id)

    def fetch() -> LiveAct:
        html, final_url = fetch_adilet(url, timeout=35)
        return _build_live_act(act_id, html, final_url)

    try:
        act = await asyncio.to_thread(fetch)
    except Exception as exc:
        raise LiveArticleVerificationError(
            f"не удалось открыть актуальную официальную редакцию {act_id} на Adilet"
        ) from exc
    _CACHE[act_id] = (now, act)
    return act


def _article_variants(act: LiveAct, article: str, part: str) -> list[tuple[str, str]]:
    """Return independently verifiable subdivisions of one cited article.

    An article-level citation (for example ``статья 350 ГК РК``) does not mean
    that every qualifier from every paragraph of that article must be repeated
    in the sentence that cites it. The old verifier concatenated all paragraphs
    before checking wording. Article 350 therefore inherited ``только`` from
    paragraph 2 even when the claim relied on paragraph 1, and a correct claim
    was blocked in production.

    Explicit paragraph citations still resolve to that paragraph only. For an
    article-level citation each parsed paragraph is checked independently; an
    unsplit article remains one variant. This narrows verification to the legal
    proposition actually asserted without weakening the qualifier checks inside
    the relevant paragraph.
    """
    items = act.articles.get(article)
    if not items:
        return []
    if part:
        exact = items.get(part)
        if exact:
            return [(part, exact)]
        # Some one-part articles are not split by the parser because the only
        # numbered paragraph is also the entire body. Only part 1 may safely
        # fall back to the whole single entry.
        if part == "1" and len(items) == 1 and "" in items:
            return [(part, items[""])]
        return []
    return [(key, value) for key, value in items.items() if value]


def _citation_paragraph(text: str, reference: citation_audit.ProvisionReference) -> str:
    for match in citation_audit._REFERENCE_RE.finditer(text or ""):
        article = match.group("article")
        part = (match.group("part") or "").strip()
        window = match.group(0) + " " + citation_audit._paragraph_around(text, match.start())
        act = citation_audit._detect_act(window)
        candidate = citation_audit.ProvisionReference(act, article, part)
        if reference.matches(candidate):
            return citation_audit._paragraph_around(text, match.start())
    return ""


def _verify_wording_variants(
    paragraph: str,
    variants: list[tuple[str, str]],
    reference: citation_audit.ProvisionReference,
) -> None:
    if not variants:
        raise LiveArticleVerificationError(
            f"{reference.label()} не содержит текста, пригодного для live-проверки"
        )

    quote = citation_audit._quote_in(paragraph)
    if quote:
        normalized_quote = normalize_text(quote)
        if any(normalized_quote in normalize_text(live_text) for _, live_text in variants):
            return
        raise LiveArticleVerificationError(
            f"дословная цитата {reference.label()} не совпадает с живым текстом Adilet"
        )

    legal_statement = _norm_claim_only(paragraph)
    checks = [
        (part, paraphrase_defects(legal_statement, live_text))
        for part, live_text in variants
    ]
    if any(not defects for _, defects in checks):
        return

    # Every subdivision rejected the proposition. Report the least noisy
    # deterministic finding instead of a qualifier copied from an unrelated
    # paragraph of the same article.
    _, best_defects = min(checks, key=lambda item: (len(item[1]), item[0]))
    detail = best_defects[0] if best_defects else "правовой вывод не подтверждён"
    raise LiveArticleVerificationError(
        f"правовой вывод по {reference.label()} не подтверждается живым текстом нормы: {detail}"
    )


def _unique_references(
    references: list[citation_audit.ProvisionReference],
) -> list[citation_audit.ProvisionReference]:
    """Verify every legal proposition once even if the Word repeats its citation."""
    result: list[citation_audit.ProvisionReference] = []
    seen: set[tuple[str, str, str]] = set()
    for reference in references:
        key = (reference.act, reference.article, reference.part)
        if key in seen:
            continue
        seen.add(key)
        result.append(reference)
    return result


async def _load_required_acts(
    references: list[citation_audit.ProvisionReference],
) -> dict[str, LiveAct]:
    """Load every official act needed by the final Word concurrently.

    The old verifier awaited each act one after another. A claim that cited the
    Civil Code, Civil Procedure Code, Tax Code and consumer law could therefore
    spend several independent 35-second network windows after drafting had
    already finished. Unsupported acts and failed source reads still fail this
    verification pass. The delivery wrapper below then removes the unverified
    proposition and returns the already generated Word instead of discarding it.
    """
    act_ids: list[str] = []
    seen: set[str] = set()
    for reference in references:
        candidates = _ACT_CANDIDATES.get(reference.act)
        if not candidates:
            raise LiveArticleVerificationError(
                f"для {reference.label()} нет детерминированного live-verifier; ссылка не выпускается"
            )
        for act_id in candidates:
            if act_id not in seen:
                seen.add(act_id)
                act_ids.append(act_id)

    loaded = await asyncio.gather(*(_live_act(act_id) for act_id in act_ids))
    return dict(zip(act_ids, loaded, strict=True))


async def verify_document_articles(file_bytes: bytes) -> None:
    text = _docx_text(file_bytes)
    references = _unique_references(citation_audit.extract_references(text))
    if not references:
        return

    acts = await _load_required_acts(references)
    for reference in references:
        candidates = _ACT_CANDIDATES[reference.act]
        matches: list[tuple[LiveAct, list[tuple[str, str]]]] = []
        for act_id in candidates:
            live = acts[act_id]
            variants = _article_variants(live, reference.article, reference.part)
            if variants:
                matches.append((live, variants))

        if len(matches) != 1:
            if not matches:
                raise LiveArticleVerificationError(
                    f"{reference.label()} не найдена в актуальном официальном тексте Adilet"
                )
            raise LiveArticleVerificationError(
                f"{reference.label()} неоднозначно сопоставилась с несколькими актуальными актами"
            )

        live, variants = matches[0]
        paragraph = _citation_paragraph(text, reference)
        if not paragraph:
            raise LiveArticleVerificationError(f"не удалось локализовать {reference.label()} в финальном Word")
        _verify_wording_variants(paragraph, variants, reference)
        LOGGER.info(
            "LIVE_ARTICLE_VERIFIED reference=%s act_id=%s edition=%s source=%s",
            reference.label(),
            live.act_id,
            live.edition_date or "not_exposed",
            live.source_url,
        )


def _table_paragraphs(table: Table) -> Iterator[Paragraph]:
    seen: set[int] = set()
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                marker = id(paragraph._p)
                if marker not in seen:
                    seen.add(marker)
                    yield paragraph
            for nested in cell.tables:
                yield from _table_paragraphs(nested)


def _document_paragraphs(document: DocumentObject) -> Iterator[Paragraph]:
    seen: set[int] = set()
    for paragraph in document.paragraphs:
        marker = id(paragraph._p)
        if marker not in seen:
            seen.add(marker)
            yield paragraph
    for table in document.tables:
        for paragraph in _table_paragraphs(table):
            marker = id(paragraph._p)
            if marker not in seen:
                seen.add(marker)
                yield paragraph


def _error_references(detail: str) -> list[citation_audit.ProvisionReference]:
    """Extract the exact provision named by a verifier error when available."""
    return _unique_references(citation_audit.extract_references(str(detail or "")))


def _matches_target(
    reference: citation_audit.ProvisionReference,
    targets: list[citation_audit.ProvisionReference],
) -> bool:
    return any(target.matches(reference) or reference.matches(target) for target in targets)


def _strip_reference_paragraphs(
    file_bytes: bytes,
    *,
    targets: list[citation_audit.ProvisionReference] | None,
) -> tuple[bytes, list[str]]:
    """Remove only legal propositions that cannot be released safely.

    If the verifier identifies a concrete article, only paragraphs containing
    that article are removed. If the source itself is unavailable and no single
    provision can be identified, every detected statutory-reference paragraph
    is removed. User facts, calculations, party details, attachments and the
    prayer for relief are left untouched unless they are in that same unsafe
    legal-reference paragraph.
    """
    document = Document(io.BytesIO(file_bytes))
    removed: list[str] = []
    for paragraph in list(_document_paragraphs(document)):
        text = " ".join(str(paragraph.text or "").split()).strip()
        if not text:
            continue
        references = _unique_references(citation_audit.extract_references(text))
        if not references:
            continue
        selected = references if targets is None else [ref for ref in references if _matches_target(ref, targets)]
        if not selected:
            continue
        for reference in selected:
            label = reference.label()
            if label not in removed:
                removed.append(label)
        paragraph.text = ""

    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue(), removed


def _prepend_live_review_notice(file_bytes: bytes) -> bytes:
    """Make the post-render downgrade visible inside the Word itself."""
    document = Document(io.BytesIO(file_bytes))
    message = (
        "KORGAN — ПЕРЕД ПОДАЧЕЙ: одна или несколько ссылок на нормы права "
        "автоматически исключены, потому что не прошли сверку с актуальной "
        "официальной редакцией Adilet. Проверьте правовое основание документа."
    )
    if document.paragraphs:
        paragraph = document.paragraphs[0].insert_paragraph_before()
    else:
        paragraph = document.add_paragraph()
    run = paragraph.add_run(message)
    run.bold = True
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


async def _repair_live_article_release(
    file_bytes: bytes,
    first_error: LiveArticleVerificationError,
) -> tuple[bytes, list[str], list[str]]:
    """Keep verified citations, remove failing ones, and guarantee a Word result.

    The verifier is intentionally strict. Release is now resilient: a bad quote,
    stale article, unsupported statute or temporary Adilet outage may downgrade
    the document, but it may not erase a DOCX that was already generated.
    """
    current = file_bytes
    details: list[str] = [str(first_error)]
    removed: list[str] = []
    error: LiveArticleVerificationError = first_error

    for _ in range(_REPAIR_ATTEMPTS):
        targets = _error_references(str(error))
        current, just_removed = _strip_reference_paragraphs(
            current,
            targets=targets or None,
        )
        if not just_removed and targets:
            # Defensive fallback for a citation whose surrounding paragraph was
            # formatted in a way that made the exact-error match impossible.
            current, just_removed = _strip_reference_paragraphs(current, targets=None)
        for label in just_removed:
            if label not in removed:
                removed.append(label)

        # If nothing can be removed, the generated Word still exists. Return it
        # as a review draft rather than reviving the former 95%/HTTP-422 dead end.
        if not just_removed:
            return current, details, removed

        try:
            await verify_document_articles(current)
            return _prepend_live_review_notice(current), details, removed
        except LiveArticleVerificationError as exc:
            error = exc
            detail = str(exc)
            if detail not in details:
                details.append(detail)

    # A document cannot contain more than a small number of independent legal
    # references in normal use, but keep the final branch deterministic even for
    # an adversarially large draft: strip all detected citations once, never loop
    # forever, and deliver the Word as a review draft.
    current, final_removed = _strip_reference_paragraphs(current, targets=None)
    for label in final_removed:
        if label not in removed:
            removed.append(label)
    if final_removed:
        current = _prepend_live_review_notice(current)
    return current, details, removed


def _downgrade_release_meta(
    meta: dict[str, Any],
    *,
    details: list[str],
    removed: list[str],
) -> dict[str, Any]:
    updated = dict(meta or {})
    notes = [str(item) for item in list(updated.get("verification_notes") or []) if str(item).strip()]
    note = (
        "FILING_ACTION: перед подачей проверить правовое основание: ссылки на нормы, "
        "не прошедшие live-сверку с актуальной редакцией Adilet, были исключены из Word"
        if removed
        else
        "FILING_ACTION: перед подачей проверить ссылки на нормы права по актуальной редакции Adilet"
    )
    if note not in notes:
        notes.append(note)
    updated["verification_notes"] = notes
    updated["filing_ready"] = False
    updated["release_status"] = "preliminary"
    updated["live_article_verification"] = "repaired" if removed else "review_required"
    updated["live_article_removed"] = list(removed)
    updated["live_article_errors"] = list(details[:8])
    return updated


async def _guarded_generate(document_type: str, context: str, language: str) -> tuple[Any, bytes, str, dict[str, Any]]:
    draft, file_bytes, filename, meta = await _ORIGINAL_GENERATE(document_type, context, language)
    if not live_article_verification_enabled():
        return draft, file_bytes, filename, meta

    try:
        await verify_document_articles(file_bytes)
        return draft, file_bytes, filename, meta
    except LiveArticleVerificationError as exc:
        LOGGER.warning(
            "LIVE_ARTICLE_RELEASE_REPAIR document_type=%s filename=%s reason=%s",
            document_type,
            filename,
            exc,
        )
        try:
            file_bytes, details, removed = await _repair_live_article_release(file_bytes, exc)
        except Exception as repair_exc:
            # The original DOCX was already produced successfully. A bug in the
            # optional repair path must never recreate the old behaviour where a
            # client reached 95% and received no file at all.
            LOGGER.exception(
                "LIVE_ARTICLE_RELEASE_REPAIR_FAILED document_type=%s filename=%s",
                document_type,
                filename,
            )
            details = [str(exc), f"repair_error={type(repair_exc).__name__}"]
            removed = []

        meta = _downgrade_release_meta(meta, details=details, removed=removed)
        try:
            draft.status = VerificationStatus.NEEDS_VERIFICATION
        except Exception:
            pass
        return draft, file_bytes, filename, meta


def install_live_article_release_runtime() -> None:
    global _INSTALLED
    if _INSTALLED:
        return
    core._generate = _guarded_generate  # type: ignore[assignment]
    _INSTALLED = True
    LOGGER.info(
        "Installed live article release runtime enabled=%s",
        live_article_verification_enabled(),
    )


install_live_article_release_runtime()

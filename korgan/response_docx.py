from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Mm, Pt

from korgan.docx_blocks import AutoNumberedList, Block, Heading, NumberedItem, Prose, render_blocks
from korgan.legal_types import ResponseDraft
from korgan.response_types import ResponseToClaimDraft


def _setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    return doc


def _response_blocks(draft: ResponseToClaimDraft) -> list[Block]:
    blocks: list[Block] = []
    if draft.claim_summary:
        blocks.append(Heading("Краткое содержание заявленных требований"))
        blocks.extend(Prose(item) for item in draft.claim_summary)
    if draft.position:
        blocks.append(Heading("Позиция ответчика"))
        blocks.extend(Prose(item) for item in draft.position)

    blocks.append(Heading("Возражения по существу заявленных требований"))
    if draft.objections:
        for objection in draft.objections:
            blocks.append(NumberedItem(objection.text, level=0))
            blocks.extend(NumberedItem(sub, level=1) for sub in objection.subclauses)
            blocks.extend(Prose(item, indent_levels=1) for item in objection.prose)
    else:
        blocks.append(Prose("[ТРЕБУЕТ УТОЧНЕНИЯ: конкретные возражения ответчика по требованиям истца]"))

    if draft.legal_basis:
        blocks.append(Heading("Правовое обоснование"))
        blocks.extend(Prose(item) for item in draft.legal_basis)

    blocks.append(Heading("На основании изложенного ПРОШУ СУД:"))
    blocks.append(AutoNumberedList(list(draft.requests) or ["[ТРЕБУЕТ УТОЧНЕНИЯ: процессуальная просьба ответчика]"]))
    blocks.append(Heading("Приложения:"))
    blocks.append(AutoNumberedList(list(draft.attachments) or ["[ТРЕБУЕТ УТОЧНЕНИЯ: перечень прилагаемых документов]"], restart=True))
    return blocks


def build_response_to_claim_docx(draft: ResponseToClaimDraft) -> bytes:
    """Production exporter used by the Telegram response-to-claim workflow."""
    doc = _setup_doc()

    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.add_run(f"В {draft.court or '[ТРЕБУЕТ УТОЧНЕНИЯ: наименование суда]'}")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(
        f"Гражданское дело № {draft.case_number}"
        if draft.case_number
        else "Гражданское дело № [ТРЕБУЕТ УТОЧНЕНИЯ: номер дела, если присвоен]"
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Истец:\n" + ("\n".join(draft.claimant) if draft.claimant else "[ТРЕБУЕТ УТОЧНЕНИЯ: ФИО/наименование и адрес]"))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Ответчик:\n" + ("\n".join(draft.defendant) if draft.defendant else "[ТРЕБУЕТ УТОЧНЕНИЯ: ФИО/наименование и адрес]"))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(draft.title or "ОТЗЫВ НА ИСК")
    run.bold = True
    run.font.size = Pt(14)

    render_blocks(doc, _response_blocks(draft))

    doc.add_paragraph()
    sign = doc.add_paragraph()
    sign.add_run("Ответчик / представитель: ____________________    ")
    sign.add_run("Дата: ____________________")

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def _legacy_blocks(draft: ResponseDraft) -> list[Block]:
    """Claude golden-fixture exporter retained for regression coverage."""
    blocks: list[Block] = [Prose(item) for item in draft.introduction]
    if draft.circumstances:
        blocks.append(Heading("Фактические обстоятельства"))
        blocks.extend(Prose(item) for item in draft.circumstances)
    if draft.objections:
        blocks.append(Heading("Возражения по существу заявленных требований"))
        for objection in draft.objections:
            blocks.append(NumberedItem(objection.text, level=0))
            blocks.extend(NumberedItem(sub, level=1) for sub in objection.subclauses)
            blocks.extend(Prose(item, indent_levels=1) for item in objection.prose)
    if draft.legal_basis:
        blocks.append(Heading("Правовое обоснование"))
        blocks.extend(Prose(item) for item in draft.legal_basis)
    blocks.append(Heading("На основании изложенного ПРОШУ СУД:"))
    blocks.append(AutoNumberedList(list(draft.requests) or ["[ТРЕБУЕТ УТОЧНЕНИЯ: процессуальные требования ответчика]"]))
    blocks.append(Heading("Приложения:"))
    blocks.append(AutoNumberedList(list(draft.attachments) or ["[ТРЕБУЕТ ДОБАВИТЬ: перечень приложений]"], restart=True))
    return blocks


def build_response_docx(draft: ResponseDraft) -> bytes:
    """Compatibility exporter for Claude's golden-document regression suite."""
    doc = _setup_doc()
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.add_run(f"В суд: {draft.court or '[ТРЕБУЕТ УТОЧНЕНИЯ: наименование суда]'}\n").bold = True
    if draft.case_reference:
        header.add_run(f"Дело № {draft.case_reference}\n")
    header.add_run("Истец:\n").bold = True
    for item in draft.plaintiff or ["[ТРЕБУЕТ УТОЧНЕНИЯ: данные истца]"]:
        header.add_run(f"{item}\n")
    header.add_run("Ответчик:\n").bold = True
    for item in draft.defendant or ["[ТРЕБУЕТ УТОЧНЕНИЯ: данные ответчика]"]:
        header.add_run(f"{item}\n")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run((draft.title or "ОТЗЫВ НА ИСКОВОЕ ЗАЯВЛЕНИЕ").upper())
    run.bold = True
    run.font.size = Pt(14)

    render_blocks(doc, _legacy_blocks(draft))
    doc.add_paragraph()
    doc.add_paragraph("Дата: ____________________")
    doc.add_paragraph("Подпись: ____________________")

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()

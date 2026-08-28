from __future__ import annotations

import io
from datetime import datetime
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from korgan.docx_blocks import AutoNumberedList, Block, Heading, Prose, render_blocks
from korgan.i18n import KK
from korgan.language_context import current_language
from korgan.legal_calc import NEEDS_CALCULATION_MARKER
from korgan.legal_types import ClaimDraft, VerificationStatus
from korgan.pro_claim_sections import pro_text


DRAFT_NOTICE = (
    "Проект сформирован KORGAN Legal AI на основании материалов пользователя. "
    "Перед подачей необходимо проверить реквизиты, доказательства, подсудность, госпошлину и отмеченные системой вопросы. "
    "Формирование проекта не гарантирует принятие документа или исход дела."
)
DRAFT_NOTICE_KK = (
    "Жоба KORGAN Legal AI арқылы пайдаланушы материалдарының негізінде қалыптастырылды. "
    "Сотқа берер алдында деректемелерді, дәлелдемелерді, соттылықты және мемлекеттік бажды тексеру қажет. "
    "Жобаның қалыптастырылуы құжаттың қабылдануына немесе істің нәтижесіне кепілдік бермейді."
)

QA_PRELIMINARY = "PRELIMINARY DRAFT"
QA_LAWYER_REVIEW = "LAWYER-REVIEW DRAFT"
QA_READY = "READY FOR FINAL HUMAN REVIEW"

REQUIRED_DOCUMENT_FIELDS: tuple[tuple[str, str], ...] = (
    ("claimant", "данные истца (ФИО/наименование, идентификатор, адрес)"),
    ("defendant", "данные ответчика (ФИО/наименование, адрес)"),
    ("facts", "обстоятельства дела"),
    ("requests", "требования к ответчику (просительная часть)"),
)


def _is_blank(value: object) -> bool:
    if isinstance(value, str):
        return not value.strip()
    if isinstance(value, list):
        return not [item for item in value if str(item).strip()]
    return not value


def missing_required_fields(draft: ClaimDraft) -> list[str]:
    return [label for attribute, label in REQUIRED_DOCUMENT_FIELDS if _is_blank(getattr(draft, attribute, None))]


def _strip_label(value: str, label: str) -> str:
    text = value.strip()
    if text.lower().startswith(label.lower()):
        text = text[len(label):].lstrip(" : \t")
    return text.strip()


def _party_lines(items: list[str], labels: tuple[str, ...], fallback: str) -> list[str]:
    lines: list[str] = []
    for item in items:
        if not item or not item.strip():
            continue
        value = item.strip()
        for label in labels:
            value = _strip_label(value, label)
        if value:
            lines.append(value)
    return lines or [fallback]


def _document_status(draft: ClaimDraft) -> str:
    court_text = "\n".join(
        [
            draft.court,
            *draft.claimant,
            *draft.defendant,
            draft.price_of_claim,
            draft.state_duty,
            draft.late_interest,
            *draft.facts,
            *draft.legal_basis,
            *draft.requests,
            *draft.attachments,
            # Маркер «требует уточнения» в расчёте или ходатайстве — такой же
            # признак предварительного проекта, как и в фактах.
            *pro_text(draft),
        ]
    ).upper()
    if (
        "[ТРЕБУЕТ УТОЧНЕНИЯ" in court_text
        or "[ТРЕБУЕТ ДОБАВИТЬ" in court_text
        or "[ТРЕБУЕТ ПРОВЕРКИ" in court_text
        or "[НАҚТЫЛАУ ҚАЖЕТ" in court_text
        or "[ТЕКСЕРУ ҚАЖЕТ" in court_text
        or NEEDS_CALCULATION_MARKER.upper() in court_text
    ):
        return QA_PRELIMINARY
    if draft.status == VerificationStatus.NEEDS_VERIFICATION or draft.verification_notes:
        return QA_LAWYER_REVIEW
    return QA_READY


def _kazakhstan_today() -> str:
    return datetime.now(ZoneInfo("Asia/Almaty")).strftime("%d.%m.%Y")


def _kk_line(value: str) -> str:
    """Localize deterministic Russian calculator wording without changing numbers."""
    text = value or ""
    replacements = (
        ("[ТРЕБУЕТ РАСЧЁТА ГОСПОШЛИНЫ]", "[МЕМЛЕКЕТТІК БАЖДЫ ЕСЕПТЕУ ҚАЖЕТ]"),
        ("от цены иска", "талап қою бағасынан"),
        ("статья 665 Налогового кодекса РК", "ҚР Салық кодексінің 665-бабы"),
        ("статья 353 Гражданского кодекса РК (Общая часть)", "ҚР Азаматтық кодексінің (Жалпы бөлім) 353-бабы"),
        ("за период", "кезеңі үшін"),
        ("базовая ставка НБ РК", "ҚР Ұлттық Банкінің базалық мөлшерлемесі"),
        ("дн.", "күн"),
    )
    for source, target in replacements:
        text = text.replace(source, target)
    return text


def _body_blocks(draft: ClaimDraft, *, kk: bool) -> list[Block]:
    """Разделы иска в порядке, в котором их читает суд.

    Профессиональные разделы (расчёт, подсудность, досудебный порядок,
    примирение, исковая давность, снятие возражений, ходатайства) печатаются
    только при наличии материала: пустой раздел хуже отсутствующего.
    """
    blocks: list[Block] = [Prose(fact) for fact in draft.facts]

    if draft.calculation:
        blocks.append(Heading("Өндіріп алынатын сомалардың есебі" if kk else "Расчёт взыскиваемых сумм"))
        blocks.extend(Prose(_kk_line(line) if kk else line) for line in draft.calculation if line.strip())

    if draft.late_interest:
        blocks.append(Heading("ҚР АК 353-бабы бойынша есеп" if kk else "Расчёт неустойки по статье 353 ГК РК"))
        blocks.append(Prose(_kk_line(draft.late_interest) if kk else draft.late_interest))

    procedural = [
        value.strip()
        for value in (
            draft.jurisdiction_reason,
            draft.pretrial_compliance,
            draft.reconciliation_measures,
            draft.limitation_period,
        )
        if value and value.strip()
    ]
    if draft.legal_basis or procedural:
        blocks.append(Heading("Құқықтық негіздеме" if kk else "Правовое обоснование"))
        blocks.extend(Prose(basis) for basis in draft.legal_basis)
        blocks.extend(Prose(_kk_line(value) if kk else value) for value in procedural)

    defenses = [item.strip() for item in draft.anticipated_defenses if item and item.strip()]
    if defenses:
        blocks.append(
            Heading(
                "Жауапкердің ықтимал қарсылықтары және оларға жауап"
                if kk
                else "Возражения ответчика и ответ на них"
            )
        )
        blocks.extend(Prose(item) for item in defenses)

    blocks.append(Heading("Жоғарыда баяндалғандардың негізінде СОТТАН СҰРАЙМЫН:" if kk else "На основании изложенного ПРОШУ СУД:"))
    blocks.append(AutoNumberedList(list(draft.requests)))

    motions = [item.strip() for item in draft.motions if item and item.strip()]
    if motions:
        blocks.append(Heading("Өтінішхаттар:" if kk else "Ходатайства:"))
        blocks.append(AutoNumberedList(motions, restart=True))

    blocks.append(Heading("Қосымшалар:" if kk else "Приложения:"))
    blocks.append(AutoNumberedList(list(draft.attachments), restart=True))
    return blocks


def build_claim_docx(draft: ClaimDraft) -> bytes:
    """Build a clean court-facing DOCX in the current Telegram session language."""
    kk = current_language() == KK
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for style_name in ("Title", "Heading 1", "Heading 2"):
        if style_name in styles:
            styles[style_name].font.name = "Times New Roman"

    document_status = _document_status(draft)
    if document_status != QA_READY:
        for current_section in doc.sections:
            footer = current_section.footer.paragraphs[0]
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = footer.add_run(DRAFT_NOTICE_KK if kk else DRAFT_NOTICE)
            run.font.name = "Times New Roman"
            run.font.size = Pt(8)
        qa = doc.add_paragraph()
        qa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        qa_run = qa.add_run(f"KORGAN QA STATUS: {document_status}")
        qa_run.bold = True
        qa_run.font.name = "Times New Roman"
        qa_run.font.size = Pt(9)

    if kk:
        court = _strip_label(_strip_label(draft.court, "Сот"), "В суд") or "[НАҚТЫЛАУ ҚАЖЕТ: соттың нақты атауы]"
        price = _strip_label(_strip_label(draft.price_of_claim, "Талап қою бағасы"), "Цена иска") or "[НАҚТЫЛАУ ҚАЖЕТ: талап қою бағасы]"
        claimant_labels = ("Талап қоюшы", "Истец")
        defendant_labels = ("Жауапкер", "Ответчик")
        claimant_fallback = "[НАҚТЫЛАУ ҚАЖЕТ: талап қоюшының деректері]"
        defendant_fallback = "[НАҚТЫЛАУ ҚАЖЕТ: жауапкердің деректері]"
    else:
        court = _strip_label(draft.court, "В суд") or "[ТРЕБУЕТ УТОЧНЕНИЯ: точное наименование суда]"
        price = _strip_label(draft.price_of_claim, "Цена иска") or "[ТРЕБУЕТ УТОЧНЕНИЯ: цена иска]"
        claimant_labels = ("Истец",)
        defendant_labels = ("Ответчик",)
        claimant_fallback = "[ТРЕБУЕТ УТОЧНЕНИЯ: данные истца]"
        defendant_fallback = "[ТРЕБУЕТ УТОЧНЕНИЯ: данные ответчика]"

    right = doc.add_paragraph()
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.add_run((f"Сот: {court}\n" if kk else f"В суд: {court}\n")).bold = True
    right.add_run("Талап қоюшы:\n" if kk else "Истец:\n").bold = True
    for item in _party_lines(draft.claimant, claimant_labels, claimant_fallback):
        right.add_run(f"{item}\n")
    right.add_run("Жауапкер:\n" if kk else "Ответчик:\n").bold = True
    for item in _party_lines(draft.defendant, defendant_labels, defendant_fallback):
        right.add_run(f"{item}\n")
    right.add_run(f"Талап қою бағасы: {price}\n" if kk else f"Цена иска: {price}\n")
    duty = _strip_label(_strip_label(draft.state_duty, "Мемлекеттік баж"), "Госпошлина") or NEEDS_CALCULATION_MARKER
    if kk:
        duty = _kk_line(duty)
    right.add_run(f"Мемлекеттік баж: {duty}" if kk else f"Госпошлина: {duty}")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(draft.title or ("ТАЛАП ҚОЮ АРЫЗЫ" if kk else "ИСКОВОЕ ЗАЯВЛЕНИЕ"))
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(14)

    render_blocks(doc, _body_blocks(draft, kk=kk))

    doc.add_paragraph()
    doc.add_paragraph(("Күні: " if kk else "Дата: ") + _kazakhstan_today())
    doc.add_paragraph("Қолы: ____________________" if kk else "Подпись: ____________________")

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()

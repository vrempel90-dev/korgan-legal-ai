"""Secure document preview and native Telegram download support for the Mini App.

The Telegram Android WebView does not reliably honor blob URL + <a download>.
This module issues short-lived HMAC-signed HTTPS links so Telegram.WebApp.downloadFile
can use the client's native file-download flow. A separate HTML preview keeps legal
documents inside KORGAN infrastructure and avoids third-party document viewers.

The module owns an isolated APIRouter. Importing it must never import or mutate the
payment/runtime app; the recovery composition layer includes this router only after
the already-tested Mini App payment stack has been constructed.
"""

from __future__ import annotations

import base64
import hashlib
import hmac
import html
import io
import json
import time
from typing import Any
from urllib.parse import quote

from docx import Document
from fastapi import APIRouter, Header, HTTPException, Request
from fastapi.responses import HTMLResponse, Response

from korgan.config import get_settings

router = APIRouter()

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_LINK_TTL_SECONDS = 180
_MAX_DOCUMENT_BYTES = 45 * 1024 * 1024


def _b64url_encode(data: bytes) -> str:
    return base64.urlsafe_b64encode(data).decode("ascii").rstrip("=")


def _b64url_decode(value: str) -> bytes:
    padding = "=" * (-len(value) % 4)
    return base64.urlsafe_b64decode(value + padding)


def _make_token(user_id: str, case_id: str, expires_at: int, secret: str) -> str:
    payload = json.dumps(
        {"u": str(user_id), "c": str(case_id), "e": int(expires_at)},
        ensure_ascii=True,
        separators=(",", ":"),
        sort_keys=True,
    ).encode("utf-8")
    encoded = _b64url_encode(payload)
    signature = hmac.new(secret.encode("utf-8"), encoded.encode("ascii"), hashlib.sha256).digest()
    return f"{encoded}.{_b64url_encode(signature)}"


def _read_token(token: str, secret: str, *, now: int | None = None) -> tuple[str, str]:
    try:
        encoded, provided = token.split(".", 1)
        expected = hmac.new(secret.encode("utf-8"), encoded.encode("ascii"), hashlib.sha256).digest()
        if not hmac.compare_digest(expected, _b64url_decode(provided)):
            raise ValueError("signature")
        payload = json.loads(_b64url_decode(encoded).decode("utf-8"))
        user_id = str(payload["u"])
        case_id = str(payload["c"])
        expires_at = int(payload["e"])
    except (ValueError, KeyError, TypeError, json.JSONDecodeError) as exc:
        raise HTTPException(status_code=403, detail="Ссылка на документ недействительна") from exc

    current = int(time.time()) if now is None else int(now)
    if expires_at < current or expires_at > current + 10 * 60:
        raise HTTPException(status_code=403, detail="Ссылка на документ истекла")
    return user_id, case_id


def _secret() -> str:
    value = (get_settings().telegram_bot_token or "").strip()
    if not value:
        raise HTTPException(status_code=503, detail="Защищённая выдача документов не настроена")
    return value


def _decode_document(case: dict[str, Any]) -> bytes:
    encoded = str(case.get("document_base64") or "")
    if not encoded:
        raise HTTPException(status_code=404, detail="Документ по этому делу ещё не готов")
    try:
        payload = base64.b64decode(encoded, validate=True)
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=500, detail="Файл документа повреждён, сгенерируйте заново") from exc
    if not payload:
        raise HTTPException(status_code=404, detail="Документ по этому делу ещё не готов")
    if len(payload) > _MAX_DOCUMENT_BYTES:
        raise HTTPException(status_code=413, detail="Документ слишком большой")
    return payload


def _safe_filename(value: str) -> str:
    raw = str(value or "KORGAN_document.docx").strip()
    cleaned = "".join(ch for ch in raw if ch.isascii() and (ch.isalnum() or ch in "._-"))
    if not cleaned:
        cleaned = "KORGAN_document.docx"
    if not cleaned.lower().endswith(".docx"):
        cleaned += ".docx"
    return cleaned[:120]


def _unicode_filename(value: str) -> str:
    """Имя файла как его видит пользователь, без разделителей пути и управляющих
    символов.

    Кириллица сохраняется: ASCII-вариант собирается отбрасыванием не-ASCII
    символов, и от «Исковое заявление.docx» в нём не остаётся ничего.
    """
    raw = " ".join(str(value or "").split()).strip()
    cleaned = "".join(
        ch for ch in raw if ch.isprintable() and ch not in '"\\/\r\n\t:*?<>|'
    ).strip(" .")
    if not cleaned:
        return "KORGAN_document.docx"
    if not cleaned.lower().endswith(".docx"):
        cleaned = f"{cleaned}.docx"
    return cleaned[:120]


def _content_disposition(value: str) -> str:
    """Заголовок вложения по RFC 6266.

    ``filename*`` несёт настоящее имя в UTF-8 — его берут браузеры и Telegram
    WebView. ``filename`` остаётся запасным ASCII-вариантом для клиентов,
    которые ``filename*`` не понимают.
    """
    ascii_name = _safe_filename(value)
    unicode_name = _unicode_filename(value)
    encoded = quote(unicode_name, safe="")
    return f"attachment; filename=\"{ascii_name}\"; filename*=UTF-8''{encoded}"


def _external_base(request: Request) -> str:
    proto = (request.headers.get("x-forwarded-proto") or request.url.scheme or "https").split(",", 1)[0].strip()
    host = (request.headers.get("x-forwarded-host") or request.headers.get("host") or request.url.netloc).split(",", 1)[0].strip()
    if host.endswith(".up.railway.app"):
        proto = "https"
    return f"{proto}://{host}".rstrip("/")


async def _case_for_identity(user_id: str, case_id: str) -> dict[str, Any]:
    # Deliberately lazy: the router module stays import-isolated from the Mini App
    # runtime and therefore cannot perturb route identity/order during test collection.
    from korgan import miniapp_api_v2 as core

    state = await core.legacy._require_consent(user_id)
    case = state.get("cases", {}).get(case_id)
    if case is None:
        raise HTTPException(status_code=404, detail="Дело не найдено")
    _decode_document(case)
    return case


def _render_docx_html(payload: bytes, title: str = "Документ KORGAN") -> str:
    try:
        document = Document(io.BytesIO(payload))
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=500, detail="Не удалось открыть документ для просмотра") from exc

    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            blocks.append('<div class="spacer"></div>')
            continue
        style_name = str(getattr(paragraph.style, "name", "") or "").lower()
        escaped = html.escape(text)
        if "title" in style_name:
            blocks.append(f"<h1>{escaped}</h1>")
        elif "heading 1" in style_name or "заголовок 1" in style_name:
            blocks.append(f"<h2>{escaped}</h2>")
        elif "heading" in style_name or "заголовок" in style_name:
            blocks.append(f"<h3>{escaped}</h3>")
        else:
            blocks.append(f"<p>{escaped}</p>")

    for table in document.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = "".join(f"<td>{html.escape(cell.text.strip())}</td>" for cell in row.cells)
            rows.append(f"<tr>{cells}</tr>")
        if rows:
            blocks.append(f"<div class=\"table-wrap\"><table>{''.join(rows)}</table></div>")

    safe_title = html.escape(str(title or "Документ KORGAN"))
    body = "".join(blocks) or "<p>Документ не содержит отображаемого текста.</p>"
    return f"""<!doctype html>
<html lang=\"ru\">
<head>
<meta charset=\"utf-8\">
<meta name=\"viewport\" content=\"width=device-width,initial-scale=1,viewport-fit=cover\">
<meta name=\"color-scheme\" content=\"dark\">
<title>{safe_title}</title>
<style>
:root{{color-scheme:dark}}*{{box-sizing:border-box}}body{{margin:0;background:#0b0f14;color:#ede9df;font-family:Inter,system-ui,-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;padding:18px 14px 40px}}.shell{{max-width:820px;margin:0 auto}}.brand{{display:flex;align-items:center;gap:9px;color:#d4ae2b;font-weight:800;letter-spacing:.14em;font-size:13px;margin:4px 0 18px}}.brand i{{width:22px;height:22px;border:2px solid #d4ae2b;border-radius:7px;display:grid;place-items:center;font-style:normal;font-size:12px}}.paper{{background:#fff;color:#181818;border-radius:16px;padding:28px 22px;box-shadow:0 16px 48px rgba(0,0,0,.35)}}h1{{font-family:Georgia,serif;font-size:23px;line-height:1.25;text-align:center;margin:4px 0 22px}}h2{{font-size:17px;margin:24px 0 10px}}h3{{font-size:15px;margin:18px 0 8px}}p{{font-family:Georgia,'Times New Roman',serif;font-size:16px;line-height:1.55;margin:8px 0;white-space:pre-wrap}}.spacer{{height:8px}}.table-wrap{{overflow:auto;margin:14px 0}}table{{border-collapse:collapse;width:100%;font-size:14px}}td{{border:1px solid #aaa;padding:7px;vertical-align:top}}.note{{color:#8f99a4;font-size:12px;line-height:1.45;margin:14px 4px 0;text-align:center}}@media(max-width:420px){{body{{padding:12px 10px 30px}}.paper{{padding:22px 16px;border-radius:13px}}p{{font-size:15px}}}}
</style>
</head><body><main class=\"shell\"><div class=\"brand\"><i>✓</i>KORGAN</div><article class=\"paper\">{body}</article><div class=\"note\">Защищённый просмотр. Ссылка действует ограниченное время и не передаёт документ сторонним сервисам.</div></main></body></html>"""


@router.post("/miniapp/cases/{case_id}/document/access")
async def create_document_access(
    case_id: str,
    request: Request,
    x_telegram_init_data: str = Header(default=""),
) -> dict[str, Any]:
    from korgan import miniapp_api_v2 as core

    user_id = core.legacy._identity(x_telegram_init_data)
    case = await _case_for_identity(user_id, case_id)
    expires_at = int(time.time()) + _LINK_TTL_SECONDS
    token = _make_token(user_id, case_id, expires_at, _secret())
    base = _external_base(request)
    encoded_token = quote(token, safe="")
    filename = _safe_filename(str(case.get("filename") or "KORGAN_document.docx"))
    return {
        "ok": True,
        "case_id": case_id,
        "filename": filename,
        "expires_at": expires_at,
        "download_url": f"{base}/miniapp/document/download?token={encoded_token}",
        "preview_url": f"{base}/miniapp/document/preview?token={encoded_token}",
    }


@router.get("/miniapp/document/download")
async def download_document(token: str) -> Response:
    user_id, case_id = _read_token(token, _secret())
    case = await _case_for_identity(user_id, case_id)
    payload = _decode_document(case)
    return Response(
        content=payload,
        media_type=_DOCX_MIME,
        headers={
            "Content-Disposition": _content_disposition(
                str(case.get("filename") or "KORGAN_document.docx")
            ),
            "Access-Control-Allow-Origin": "https://web.telegram.org",
            "Cache-Control": "private, no-store, max-age=0",
            "X-Content-Type-Options": "nosniff",
            "Referrer-Policy": "no-referrer",
        },
    )


@router.get("/miniapp/document/preview")
async def preview_document(token: str) -> HTMLResponse:
    user_id, case_id = _read_token(token, _secret())
    case = await _case_for_identity(user_id, case_id)
    payload = _decode_document(case)
    response = HTMLResponse(_render_docx_html(payload, str(case.get("title") or "Документ KORGAN")))
    response.headers["Cache-Control"] = "private, no-store, max-age=0"
    response.headers["Referrer-Policy"] = "no-referrer"
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["Content-Security-Policy"] = "default-src 'none'; style-src 'unsafe-inline'; base-uri 'none'; form-action 'none'"
    return response

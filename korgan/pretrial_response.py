"""Professional response to a pre-trial demand for RU/KK."""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from korgan.document_release import review_lines
from korgan.legal_types import LegalResearch, VerificationStatus
from korgan.pretrial import PretrialProductionService
from korgan.stable_legal_release import clean_language_labels, sanitize_research_sources

_PRETRIAL_RESPONSE_SCHEMA: dict[str, Any] = {
    "type": "object",
    "properties": {
        "title": {"type": "string"},
        "sender": {"type": "array", "items": {"type": "string"}},
        "recipient": {"type": "array", "items": {"type": "string"}},
        "reference": {"type": "string"},
        "claim_summary": {"type": "array", "items": {"type": "string"}},
        "position": {"type": "array", "items": {"type": "string"}},
        "objections": {"type": "array", "items": {"type": "string"}},
        "legal_basis": {"type": "array", "items": {"type": "string"}},
        "response_terms": {"type": "array", "items": {"type": "string"}},
        "attachments": {"type": "array", "items": {"type": "string"}},
        "verification_notes": {"type": "array", "items": {"type": "string"}},
    },
    "required": [
        "title", "sender", "recipient", "reference", "claim_summary", "position",
        "objections", "legal_basis", "response_terms", "attachments", "verification_notes"
    ],
    "additionalProperties": False,
}

_INTENT_RU = re.compile(
    r"(?i)(?:\bотзыв\w*\s+на\s+(?:досудебн\w*\s+)?претензи\w*|"
    r"\bответ\w*\s+на\s+(?:досудебн\w*\s+)?претензи\w*|"
    r"\bвозражен\w*\s+на\s+(?:досудебн\w*\s+)?претензи\w*)"
)
_INTENT_KK = re.compile(
    r"(?i)(?:сотқа\s+дейінгі\s+талап\w*\s+(?:жауап\w*|пікір\w*)|"
    r"талап\s+хат\w*\s+(?:жауап\w*|пікір\w*))"
)
_ACTION = re.compile(
    r"(?i)\b(?:подготов\w*|состав\w*|сформир\w*|сдел\w*|напиш\w*|напис\w*|"
    r"дайында\w*|жаса\w*|әзірле\w*|құрастыр\w*|жаз\w*)\b"
)
_ADVICE = re.compile(r"(?i)^\s*как\b|\bқалай\b")


def is_pretrial_response_request(text: str | None) -> bool:
    value = " ".join((text or "").split())
    if not value or _ADVICE.search(value):
        return False
    return bool((_INTENT_RU.search(value) or _INTENT_KK.search(value)) and _ACTION.search(value))


@dataclass(slots=True)
class PretrialResponseDraft:
    status: VerificationStatus
    title: str
    sender: list[str]
    recipient: list[str]
    reference: str
    claim_summary: list[str]
    position: list[str]
    objections: list[str]
    legal_basis: list[str]
    response_terms: list[str]
    attachments: list[str]
    verification_notes: list[str] = field(default_factory=list)
    source_urls: list[str] = field(default_factory=list)

    def body_lines(self) -> list[str]:
        return [
            self.title,
            *self.sender,
            *self.recipient,
            self.reference,
            *self.claim_summary,
            *self.position,
            *self.objections,
            *self.legal_basis,
            *self.response_terms,
            *self.attachments,
        ]


def _dedupe(lines: list[str]) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    for raw in lines:
        line = clean_language_labels(str(raw or "")).strip()
        key = re.sub(r"\W+", "", line.lower())
        if line and key and key not in seen:
            seen.add(key)
            result.append(line)
    return result


def normalize_pretrial_response(draft: PretrialResponseDraft) -> None:
    draft.claim_summary = _dedupe(draft.claim_summary)
    draft.position = _dedupe(draft.position)
    draft.objections = _dedupe(draft.objections)
    draft.legal_basis = _dedupe(draft.legal_basis)
    draft.response_terms = _dedupe(draft.response_terms)
    draft.attachments = _dedupe(draft.attachments)


def pretrial_response_quality_issues(draft: PretrialResponseDraft, research: LegalResearch) -> list[str]:
    issues: list[str] = []
    if not draft.sender:
        issues.append("не указан отправитель ответа")
    if not draft.recipient:
        issues.append("не указан адресат ответа")
    if not draft.claim_summary:
        issues.append("не отражены требования исходной претензии")
    if not draft.position:
        issues.append("не сформулирована позиция получателя претензии")
    if not draft.objections and not draft.response_terms:
        issues.append("нет содержательного ответа на требования претензии")
    if research.verified_claims and not draft.legal_basis:
        issues.append("VERIFIED нормы не перенесены в правовое обоснование")
    report = review_lines(draft.body_lines(), verified_claims=research.verified_claims)
    issues.extend(report.blocking)
    return list(dict.fromkeys(issues))


class PretrialResponseProductionService(PretrialProductionService):
    async def research_pretrial_response(self, case_context: str, language: str = "ru") -> LegalResearch:
        research = await self.research_case(case_context, language=language)
        return sanitize_research_sources(research)

    async def draft_pretrial_response(
        self,
        case_context: str,
        research: LegalResearch,
        language: str = "ru",
    ) -> PretrialResponseDraft:
        verified = "\n".join(f"- {x}" for x in research.verified_claims) or "- нет подтвержденных норм"
        prompt = (
            "Подготовь профессиональный ОТВЕТ НА ДОСУДЕБНУЮ ПРЕТЕНЗИЮ по праву Республики Казахстан. "
            "Это не отзыв на иск, не иск и не встречная претензия. Документ готовится непосредственно от имени получателя претензии.\n\n"
            "ЭТАЛОН ПОДАЧИ И СТРУКТУРЫ:\n"
            "Документ должен выглядеть как реальный деловой ответ на претензию, подготовленный практикующим юристом, а не как AI-анализ. "
            "Ориентируйся на такую последовательность: реквизиты сторон → ссылка на исходящую претензию → сразу чёткая позиция адресата (несогласие, частичное признание либо иная позиция строго по материалам) → описание договора и существенных условий → последовательный разбор каждого спорного довода → относящиеся к спору соглашения, зачёты, гарантийные удержания, акты, исполнительная документация и другие обстоятельства только если они реально есть в деле → договорные пункты и расчёты неустойки/удержаний только при наличии оснований и исходных данных → VERIFIED-нормы → итоговая позиция, встречное требование/зачёт или предложение урегулирования только если это подтверждено материалами → подпись.\n"
            "Пиши связными деловыми абзацами. Не превращай документ в меморандум с искусственными разделами «Содержание претензии», «Позиция», «Возражения и пояснения», «Правовое обоснование», «Ответ на требования». "
            "Каждый элемент claim_summary, position, objections, legal_basis и response_terms формулируй как готовый абзац письма, который естественно продолжает предыдущий.\n\n"
            "ОБЯЗАТЕЛЬНЫЕ ПРАВИЛА:\n"
            "1. Не копируй никакие факты, суммы, даты, названия компаний, договоры, проценты, сроки, пункты договора или статьи из примера оформления. Пример задаёт только форму и уровень юридической подачи.\n"
            "2. Сначала точно выдели требования исходной претензии и не подменяй их другими требованиями.\n"
            "3. Позицию получателя формулируй только из фактов пользователя: признание, частичное признание или несогласие нельзя придумывать.\n"
            "4. Пиши от имени адресата напрямую: «сообщаем», «не признаём», «считаем», «обращаем внимание», «предлагаем», «готовы рассмотреть». Не описывай собственную позицию в третьем лице как «ТОО считает», «ответчик сообщает», «со стороны ТОО отсутствует позиция».\n"
            "5. Каждое возражение связывай с конкретным требованием/фактом претензии и имеющимися доказательствами.\n"
            "6. Конкретные статьи и юридические выводы используй только из VERIFIED. Никаких норм по памяти.\n"
            "7. Не применяй процессуальные правила об отзыве на иск как основание для этого документа: это внесудебный ответ на претензию.\n"
            "8. Не придумывай ФИО/БИН/ИИН, адреса, договоры, даты, суммы, платежи, переписку, доказательства, сроки или факт отправки ответа.\n"
            "9. Если в материалах есть договорные удержания, зачёт, неустойка, штраф, встречные требования или право на одностороннее удержание, объясни их последовательно и рассчитай только при наличии всех исходных данных. Если таких фактов нет — не добавляй их ради сходства с образцом.\n"
            "10. Не обещай оплату/исполнение и не устанавливай новый срок, если пользователь этого не сообщил и обязанность не подтверждена.\n"
            "11. Если разумно предложить урегулирование, делай это нейтрально и только если оно не противоречит позиции пользователя.\n"
            "12. Приложения перечисляй только из реально имеющихся материалов.\n"
            "13. Не включай в тело документа внутренние рассуждения вроде «позиция не определена», «нет подтверждённого согласия», «правовая оценка не проведена». Недостающие критичные сведения оставляй только в verification_notes.\n"
            f"14. Язык документа: {'казахский' if language == 'kk' else 'русский'}.\n\n"
            f"МАТЕРИАЛЫ:\n{case_context[:self.settings.max_case_text_chars]}\n\n"
            f"VERIFIED:\n{verified}"
        )
        payload, _ = await self._structured_response(
            model=self.settings.openai_model,
            instructions=(
                "Ты практикующий юрист KORGAN в Республике Казахстан. Составляй полноценный деловой ответ на досудебную претензию уровня юридической практики: прямую позицию адресата, связный разбор договора и доводов, договорные расчёты и правовые основания только там, где они подтверждены материалами. "
                "Не смешивай его с судебным отзывом на иск, не выдумывай факты, не копируй содержание образцов оформления и не добавляй непроверенное право."
            ),
            content=[{"role": "user", "content": [{"type": "input_text", "text": prompt}]}],
            schema_name="korgan_pretrial_response",
            schema=_PRETRIAL_RESPONSE_SCHEMA,
        )
        draft = PretrialResponseDraft(
            status=research.status,
            source_urls=list(research.source_urls),
            **payload,
        )
        normalize_pretrial_response(draft)
        issues = pretrial_response_quality_issues(draft, research)
        if issues:
            draft.status = VerificationStatus.NEEDS_VERIFICATION
            draft.verification_notes.extend(x for x in issues if x not in draft.verification_notes)
        return draft


def _today() -> str:
    return datetime.now(ZoneInfo("Asia/Almaty")).strftime("%d.%m.%Y")


def _body_paragraph(doc: Document, text: str) -> None:
    value = str(text or "").strip()
    if not value:
        return
    p = doc.add_paragraph(value)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(6)


def _reference_line(reference: str, kk: bool) -> str:
    value = str(reference or "").strip()
    if not value:
        return ""
    lower = value.lower()
    if kk:
        if "шығыс" in lower or "№" in value:
            return value
        return f"Сотқа дейінгі талапқа қатысты: {value}"
    if "исх" in lower or "№" in value:
        return value if lower.startswith("на ") else f"На {value}"
    return f"На досудебную претензию: {value}"


def build_pretrial_response_docx(draft: PretrialResponseDraft, language: str = "ru") -> bytes:
    """Render the response as a continuous business letter matching the approved reference style."""
    kk = language == "kk"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    head.add_run(("Алушы:\n" if kk else "Кому:\n")).bold = True
    for value in draft.recipient or [("[Алушы деректері]" if kk else "[Данные адресата]")]:
        head.add_run(str(value) + "\n")
    head.add_run(("Жіберуші:\n" if kk else "От:\n")).bold = True
    for value in draft.sender or [("[Жіберуші деректері]" if kk else "[Данные отправителя]")]:
        head.add_run(str(value) + "\n")

    reference = _reference_line(draft.reference, kk)
    if reference:
        p = doc.add_paragraph(reference)
        p.paragraph_format.space_after = Pt(8)
    else:
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(draft.title or ("Сотқа дейінгі талапқа жауап" if kk else "Ответ на досудебную претензию"))
        run.bold = True
        run.font.size = Pt(14)

    for item in draft.claim_summary:
        _body_paragraph(doc, item)

    for item in draft.position:
        _body_paragraph(doc, item)

    for item in draft.objections:
        _body_paragraph(doc, item)

    for item in draft.legal_basis:
        _body_paragraph(doc, item)

    for item in draft.response_terms:
        _body_paragraph(doc, item)

    if draft.attachments:
        p = doc.add_paragraph()
        p.add_run("Қосымшалар:" if kk else "Приложения:").bold = True
        for index, item in enumerate(draft.attachments, 1):
            doc.add_paragraph(f"{index}. {item}")

    doc.add_paragraph()
    doc.add_paragraph(("Күні: " if kk else "Дата: ") + _today())
    doc.add_paragraph("Қолы: ____________________" if kk else "Подпись: ____________________")

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()

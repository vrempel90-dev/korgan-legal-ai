from __future__ import annotations

import io
from typing import Any
from urllib.parse import quote

from aiogram import Bot
from aiogram.types import (
    BufferedInputFile,
    InlineKeyboardButton,
    InlineKeyboardMarkup,
    KeyboardButton,
    ReplyKeyboardMarkup,
)
from docx import Document

from korgan.client_safe_ui import ClientSafeBot, _clean_upload, sanitize_client_text
from korgan.i18n import BUTTONS, KK, RU, tr
from korgan.language_context import current_language

_BUTTON_MAP = {value: BUTTONS[KK][key] for key, value in BUTTONS[RU].items()}
_EXTRA_BUTTON_MAP = {
    "✅ Согласен, сформировать документ": "✅ Келісемін, құжатты дайындау",
    "↩️ Отмена": "↩️ Болдырмау",
}

_DOCX_REPLACEMENTS = (
    ("В суд:", "Сотқа:"),
    ("Истец:", "Талап қоюшы:"),
    ("Ответчик:", "Жауапкер:"),
    ("Цена иска:", "Талап қою бағасы:"),
    ("Госпошлина:", "Мемлекеттік баж:"),
    ("Правовое обоснование", "Құқықтық негіздеме"),
    ("Расчёт неустойки по статье 353 ГК РК", "ҚР АК 353-бабы бойынша есеп"),
    ("На основании изложенного ПРОШУ СУД:", "Жоғарыда баяндалғандардың негізінде СОТТАН СҰРАЙМЫН:"),
    ("Приложения:", "Қосымшалар:"),
    ("Дата:", "Күні:"),
    ("Подпись:", "Қолы:"),
    ("ИСКОВОЕ ЗАЯВЛЕНИЕ", "ТАЛАП ҚОЮ АРЫЗЫ"),
    ("ОТЗЫВ НА ИСКОВОЕ ЗАЯВЛЕНИЕ", "ТАЛАП ҚОЮ АРЫЗЫНА ПІКІР"),
    ("ОТЗЫВ НА ИСК", "ТАЛАП ҚОЮ АРЫЗЫНА ПІКІР"),
    ("Краткое содержание заявленных требований", "Мәлімделген талаптардың қысқаша мазмұны"),
    ("Позиция ответчика", "Жауапкердің ұстанымы"),
    ("Возражения по существу заявленных требований", "Мәлімделген талаптардың мәні бойынша қарсылықтар"),
    ("Гражданское дело №", "Азаматтық іс №"),
    ("Ответчик / представитель:", "Жауапкер / өкіл:"),
    ("РЕКВИЗИТЫ И ПОДПИСИ СТОРОН", "ТАРАПТАРДЫҢ ДЕРЕКТЕМЕЛЕРІ МЕН ҚОЛДАРЫ"),
    ("Сторона 1", "1-тарап"),
    ("Сторона 2", "2-тарап"),
)

_REVIEW_PHONE = "77005000553"
_CLAIM_REVIEW_PHONE = _REVIEW_PHONE
_CLAIM_FILENAME = "korgan_iskovoe_zayavlenie.docx"

_DOCUMENT_KINDS = {
    "korgan_iskovoe_zayavlenie.docx": "claim",
    "korgan_otzyv_na_isk.docx": "response",
    "korgan_dogovor.docx": "contract",
    "korgan_dosudebnaya_pretenziya.docx": "pretrial",
    "korgan_sotqa_deyingi_talap.docx": "pretrial",
}


def _generated_document_kind(document: Any) -> str | None:
    """Return the client-facing document category for a KORGAN-generated file."""
    filename = str(getattr(document, "filename", "") or "").rsplit("/", 1)[-1].lower()
    known = _DOCUMENT_KINDS.get(filename)
    if known:
        return known
    if filename.startswith("korgan_") and filename.endswith((".docx", ".pdf")):
        return "document"
    return None


def _document_review_text(kind: str, language: str) -> str:
    if language == KK:
        intro = {
            "claim": (
                "⚠️ Бұл талап қою арызы KORGAN AI көмегімен дайындалды. "
                "Сотқа берер алдында оны заңгерге тексертуге кеңес береміз."
            ),
            "pretrial": (
                "⚠️ Бұл сотқа дейінгі талап KORGAN AI көмегімен дайындалды. "
                "Оны адресатқа жіберер алдында заңгерге тексертуге кеңес береміз."
            ),
            "response": (
                "⚠️ Бұл талап қою арызына пікір KORGAN AI көмегімен дайындалды. "
                "Сотқа ұсынар алдында оны заңгерге тексертуге кеңес береміз."
            ),
            "contract": (
                "⚠️ Бұл шарт жобасы KORGAN AI көмегімен дайындалды. "
                "Қол қояр алдында оны заңгерге тексертуге кеңес береміз."
            ),
            "document": (
                "⚠️ Бұл құжат KORGAN AI көмегімен дайындалды. "
                "Пайдаланар алдында оны заңгерге тексертуге кеңес береміз."
            ),
        }.get(kind, "⚠️ Бұл құжат KORGAN AI көмегімен дайындалды. Оны заңгерге тексертуге кеңес береміз.")
        scope = {
            "claim": "Тексеру тек осы талап қою арызына қатысты.",
            "pretrial": "Тексеру тек осы сотқа дейінгі талапқа қатысты.",
            "response": "Тексеру тек осы пікірге қатысты.",
            "contract": "Тексеру тек осы шарт жобасына қатысты.",
            "document": "Тексеру тек осы құжатқа қатысты.",
        }.get(kind, "Тексеру тек осы құжатқа қатысты.")
        return (
            f"{intro}\n\n"
            f"{scope} Қосымша құжаттарды дайындау немесе тексеру және өзге заңгерлік жұмыс — бөлек ақылы қызмет.\n\n"
            "WhatsApp ашылғаннан кейін осы чатта алған Word/PDF файлын заңгерге тіркеңіз."
        )

    intro = {
        "claim": "⚠️ Этот иск подготовлен с использованием KORGAN AI. Перед подачей в суд рекомендуем проверить его у юриста.",
        "pretrial": "⚠️ Эта досудебная претензия подготовлена с использованием KORGAN AI. Перед направлением адресату рекомендуем проверить её у юриста.",
        "response": "⚠️ Этот отзыв на иск подготовлен с использованием KORGAN AI. Перед подачей в суд рекомендуем проверить его у юриста.",
        "contract": "⚠️ Этот проект договора подготовлен с использованием KORGAN AI. Перед подписанием рекомендуем проверить его у юриста.",
        "document": "⚠️ Этот документ подготовлен с использованием KORGAN AI. Перед использованием рекомендуем проверить его у юриста.",
    }.get(kind, "⚠️ Этот документ подготовлен с использованием KORGAN AI. Рекомендуем проверить его у юриста.")
    scope = {
        "claim": "Проверка относится только к этому иску.",
        "pretrial": "Проверка относится только к этой досудебной претензии.",
        "response": "Проверка относится только к этому отзыву на иск.",
        "contract": "Проверка относится только к этому проекту договора.",
        "document": "Проверка относится только к этому документу.",
    }.get(kind, "Проверка относится только к этому документу.")
    return (
        f"{intro}\n\n"
        f"{scope} Подготовка или проверка дополнительных документов и иная юридическая работа — отдельная платная услуга.\n\n"
        "После открытия WhatsApp прикрепите полученный в этом чате Word/PDF-файл."
    )


def _document_review_markup(kind: str, language: str) -> InlineKeyboardMarkup:
    if language == KK:
        label = {
            "claim": "👨‍⚖️ Талапты WhatsApp-та тексеру",
            "pretrial": "👨‍⚖️ Сотқа дейінгі талапты тексеру",
            "response": "👨‍⚖️ Пікірді WhatsApp-та тексеру",
            "contract": "👨‍⚖️ Шартты WhatsApp-та тексеру",
            "document": "👨‍⚖️ Құжатты WhatsApp-та тексеру",
        }.get(kind, "👨‍⚖️ Құжатты WhatsApp-та тексеру")
        subject = {
            "claim": "талап қою арызын",
            "pretrial": "сотқа дейінгі талапты",
            "response": "талап қою арызына пікірді",
            "contract": "шарт жобасын",
            "document": "құжатты",
        }.get(kind, "құжатты")
        prefill = (
            f"Сәлеметсіз бе! KORGAN-да дайындалған {subject} заңгерге тексеруге бергім келеді. "
            "Қазір Word/PDF файлын тіркеймін."
        )
    else:
        label = {
            "claim": "👨‍⚖️ Проверить иск в WhatsApp",
            "pretrial": "👨‍⚖️ Проверить претензию в WhatsApp",
            "response": "👨‍⚖️ Проверить отзыв в WhatsApp",
            "contract": "👨‍⚖️ Проверить договор в WhatsApp",
            "document": "👨‍⚖️ Проверить документ в WhatsApp",
        }.get(kind, "👨‍⚖️ Проверить документ в WhatsApp")
        subject = {
            "claim": "иск",
            "pretrial": "досудебную претензию",
            "response": "отзыв на иск",
            "contract": "проект договора",
            "document": "документ",
        }.get(kind, "документ")
        prefill = (
            f"Здравствуйте! Хочу передать {subject}, подготовленный в KORGAN, на проверку юристу. "
            "Сейчас прикреплю Word/PDF-файл."
        )
    url = f"https://wa.me/{_REVIEW_PHONE}?text={quote(prefill)}"
    return InlineKeyboardMarkup(inline_keyboard=[[InlineKeyboardButton(text=label, url=url)]])


# Backward-compatible claim helpers kept for existing tests and callers.
def _claim_review_text(language: str) -> str:
    return _document_review_text("claim", language)


def _claim_review_markup(language: str) -> InlineKeyboardMarkup:
    return _document_review_markup("claim", language)


def _is_claim_document(document: Any) -> bool:
    return _generated_document_kind(document) == "claim"


def _localize_markup(markup: Any) -> Any:
    if current_language() != KK or markup is None:
        return markup
    mapping = {**_BUTTON_MAP, **_EXTRA_BUTTON_MAP}
    if isinstance(markup, ReplyKeyboardMarkup):
        rows: list[list[KeyboardButton]] = []
        for row in markup.keyboard:
            rows.append([button.model_copy(update={"text": mapping.get(button.text, button.text)}) for button in row])
        return markup.model_copy(update={"keyboard": rows, "input_field_placeholder": "Әрекетті таңдаңыз немесе хабарлама жазыңыз…"})
    if isinstance(markup, InlineKeyboardMarkup):
        rows = []
        for row in markup.inline_keyboard:
            rows.append([button.model_copy(update={"text": mapping.get(button.text, button.text)}) for button in row])
        return markup.model_copy(update={"inline_keyboard": rows})
    return markup


def _localize_text(text: str | None) -> str | None:
    clean = sanitize_client_text(text)
    if clean is None or current_language() != KK:
        return clean
    if clean.startswith("📄 Что нужно подготовить?"):
        return tr(KK, "documents")
    if clean.startswith("⚖️ Сначала опишите обстоятельства дела"):
        return "⚖️ Алдымен істің мән-жайын сипаттаңыз немесе құжаттарды/скандарды жіберіңіз. Содан кейін, мысалы, «қарызды өндіріп алу туралы талап қою арызын дайында» деп жазыңыз."
    if clean.startswith("🤝 Опишите договор одним сообщением:"):
        return "🤝 Шартты бір хабарламада сипаттаңыз: қандай шарт керек, тараптар мен олардың рөлдері, шарттың мәні, бағасы/төлемі, мерзімі және маңызды талаптары. Құжаттарды немесе хат-хабарды да тіркеуге болады."
    if clean.startswith("🛡 Чтобы подготовить отзыв на иск"):
        return "🛡 Талап қою арызына пікір дайындау үшін талап қою арызының өзін (PDF/DOCX/фото) жіберіңіз немесе оның негізгі талаптарын мәтінмен енгізіңіз. Жауапкердің ұстанымы мен дәлелдерін де көрсетіңіз."
    if clean.startswith("📦 Моё дело"):
        return (clean.replace("📦 Моё дело", "📦 Менің ісім").replace("Документов / сканов:", "Құжаттар / скандар:").replace("Текстовых описаний:", "Мәтіндік сипаттамалар:").replace("Добавляйте файлы, фото или текст.", "Файл, фото немесе мәтін қосыңыз.").replace("Затем нажмите «📄 Документ» и выберите иск или договор — либо напишите запрос обычным сообщением.", "Содан кейін «📄 Құжат» батырмасын басып, қажетті құжатты таңдаңыз немесе сұрауды жай хабарлама ретінде жазыңыз."))
    if clean.startswith("💰 Акционные цены KORGAN"):
        return tr(KK, "prices")
    if clean.startswith("❓ Как работать с KORGAN:"):
        return tr(KK, "help")
    if clean.startswith("✅ Материал разобран и добавлен в дело"):
        return clean.replace("✅ Материал разобран и добавлен в дело", "✅ Материал талданып, іске қосылды").replace("Если всё верно, можно добавить ещё документы или попросить подготовить иск — он придёт файлом Word (.docx).", "Барлығы дұрыс болса, қосымша құжаттар жіберуге немесе талап қою арызын дайындауды сұрауға болады — ол Word (.docx) файлы түрінде келеді.")

    exact = {
        "⚖️ Опишите ситуацию одним сообщением. Если есть документы или сканы — просто отправьте их в этот чат.": tr(KK, "consult_prompt"),
        "👨‍⚖️ Раздел подключения живого юриста будет доступен после запуска клиентского режима.": tr(KK, "lawyer"),
        "🆘 Если файл не принимается или документ формируется некорректно, пришлите описание ошибки сюда.": tr(KK, "support"),
        "⭐ Напишите отзыв следующим сообщением, начав его со слова «Отзыв:».": tr(KK, "feedback"),
        "Не удалось разобрать документ. Проверьте формат/качество и попробуйте ещё раз.": tr(KK, "upload_error"),
        "Не удалось выполнить юридический поиск. Попробуйте повторить вопрос.": tr(KK, "consult_error"),
        "Проверяю вид договора и актуальные нормы РК, затем формирую Word-документ…": tr(KK, "contract_progress"),
        "Проверяю требования иска и актуальные нормы РК, затем формирую отзыв в Word…": tr(KK, "response_progress"),
        "🏠 Главное меню": "🏠 Басты мәзір",
    }
    if clean in exact:
        return exact[clean]
    for source, target in (
        ("Официальные источники:", "Ресми дереккөздер:"),
        ("Готовый проект иска — файл Word (.docx).", "Талап қою арызының дайын жобасы — Word (.docx) файлы."),
        ("Проект иска сформирован в Word (.docx).", "Талап қою арызының жобасы Word (.docx) форматында дайын."),
        ("Отзыв на иск сформирован в Word (.docx).", "Талап қою арызына пікір Word (.docx) форматында дайын."),
        ("Проект договора сформирован в Word (.docx).", "Шарт жобасы Word (.docx) форматында дайын."),
        ("Перед подачей проверьте:", "Сотқа берер алдында тексеріңіз:"),
        ("Перед подписанием проверьте:", "Қол қою алдында тексеріңіз:"),
        ("Не удалось безопасно сформировать договор.", "Шартты қауіпсіз қалыптастыру мүмкін болмады."),
        ("Не удалось сформировать отзыв из текущих материалов.", "Ағымдағы материалдар бойынша пікірді қалыптастыру мүмкін болмады."),
    ):
        clean = clean.replace(source, target)
    return clean


def _replace_runs(paragraph: Any) -> None:
    for run in paragraph.runs:
        text = run.text
        for source, target in _DOCX_REPLACEMENTS:
            text = text.replace(source, target)
        if text.startswith("Проект сформирован KORGAN Legal AI"):
            text = "Жоба KORGAN Legal AI жүйесінде пайдаланушы материалдары және тексерілген құқықтық дереккөздер негізінде қалыптастырылды. Қол қою немесе сотқа беру алдында деректемелерді және маңызды талаптарды тексеріңіз."
        run.text = text


def _localize_docx(document: Any) -> Any:
    if current_language() != KK or not isinstance(document, BufferedInputFile):
        return document
    filename = str(getattr(document, "filename", "") or "")
    payload = getattr(document, "data", None)
    if not filename.lower().endswith(".docx") or not isinstance(payload, (bytes, bytearray)):
        return document
    try:
        doc = Document(io.BytesIO(bytes(payload)))
        for paragraph in doc.paragraphs:
            _replace_runs(paragraph)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_runs(paragraph)
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                _replace_runs(paragraph)
            for paragraph in section.footer.paragraphs:
                _replace_runs(paragraph)
        output = io.BytesIO()
        doc.save(output)
        return BufferedInputFile(output.getvalue(), filename=filename)
    except Exception:
        return document


class LocalizedClientSafeBot(ClientSafeBot):
    async def send_message(self, chat_id: Any, text: str, *args: Any, **kwargs: Any) -> Any:
        if "reply_markup" in kwargs:
            kwargs["reply_markup"] = _localize_markup(kwargs.get("reply_markup"))
        return await Bot.send_message(self, chat_id, _localize_text(text) or "", *args, **kwargs)

    async def send_document(self, chat_id: Any, document: Any, *args: Any, **kwargs: Any) -> Any:
        document_kind = _generated_document_kind(document)
        if "caption" in kwargs:
            kwargs["caption"] = _localize_text(kwargs.get("caption"))
        if "reply_markup" in kwargs:
            kwargs["reply_markup"] = _localize_markup(kwargs.get("reply_markup"))
        result = await Bot.send_document(self, chat_id, _localize_docx(_clean_upload(document)), *args, **kwargs)
        if document_kind is not None:
            language = current_language()
            try:
                await Bot.send_message(
                    self,
                    chat_id,
                    _document_review_text(document_kind, language),
                    reply_markup=_document_review_markup(document_kind, language),
                )
            except Exception:
                # The optional CTA must never turn a successfully delivered legal document into a failed request.
                pass
        return result

    async def edit_message_text(self, text: str, *args: Any, **kwargs: Any) -> Any:
        if "reply_markup" in kwargs:
            kwargs["reply_markup"] = _localize_markup(kwargs.get("reply_markup"))
        return await Bot.edit_message_text(self, _localize_text(text) or "", *args, **kwargs)

    async def edit_message_caption(self, *args: Any, **kwargs: Any) -> Any:
        if "caption" in kwargs:
            kwargs["caption"] = _localize_text(kwargs.get("caption"))
        if "reply_markup" in kwargs:
            kwargs["reply_markup"] = _localize_markup(kwargs.get("reply_markup"))
        return await Bot.edit_message_caption(self, *args, **kwargs)
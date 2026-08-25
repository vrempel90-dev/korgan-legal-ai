from __future__ import annotations

import io
import re
from datetime import datetime
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from korgan.contract_preamble import ensure_identified_preamble
from korgan.docx_blocks import Block, Heading, NumberedItem, Prose, render_blocks
from korgan.legal_types import ContractDraft, VerificationStatus


DRAFT_NOTICE = (
    "Проект сформирован KORGAN Legal AI на основании материалов пользователя и проверенных правовых источников. "
    "Перед подписанием необходимо проверить реквизиты сторон, коммерческие условия и отмеченные системой вопросы."
)

_REQUISITES_SECTION_RE = re.compile(
    r"(?i)^\s*(?:"
    r"реквизит\w*(?:\s+и\s+подпис\w*)?(?:\s+сторон)?|"
    r"адрес\w*\s+и\s+реквизит\w*(?:\s+сторон)?|"
    r"подпис\w*\s+сторон|"
    r"местонахождени\w*\s+и\s+(?:банковск\w*\s+)?реквизит\w*"
    r")\s*[.:;-]*\s*$"
)
_REQUISITES_HEADING_TOKEN_RE = re.compile(r"(?i)\b(?:реквизит\w*|подпис\w*)\b")
_REQUISITES_CLAUSE_RE = re.compile(
    r"(?i)(?:\bБИН\b|\bИИН\b|\bИИК\b|\bБИК\b|\bIBAN\b|\bКБе\b|"
    r"\bбанк\w*\b|\bюридическ\w*\s+адрес\b|\bпочтов\w*\s+адрес\b|"
    r"\bподпис\w*\b|\bм\.?\s*п\.?\b|\bдиректор\b\s*[:_])"
)
_SIGNATURE_LINE_RE = re.compile(
    r"(?i)^\s*(?:подпис\w*|signature|м\.?\s*п\.?)\s*[:_\-–—.\s]*$"
)
_PARTY_LABEL_RE = re.compile(r"(?i)^\s*(?:сторона\s*[12аб]|party\s*[ab12])\s*:?[\s.]*$")
_SPEC_REFERENCE_RE = re.compile(r"(?i)\bспецификаци\w*\b")
_SPEC_HEADING_RE = re.compile(r"(?i)^\s*(?:приложени\w*\s*№?\s*\d*\s*[:.-]?\s*)?спецификаци\w*\s*$")


def _status(draft: ContractDraft, preamble: list[str]) -> str:
    body = "\n".join([*draft.body_lines(), *preamble]).upper()
    if "[ТРЕБУЕТ УТОЧНЕНИЯ" in body:
        return "PRELIMINARY DRAFT"
    if draft.status == VerificationStatus.NEEDS_VERIFICATION or draft.verification_notes:
        return "LAWYER-REVIEW DRAFT"
    return "READY FOR FINAL HUMAN REVIEW"


def _today_kz() -> str:
    return datetime.now(ZoneInfo("Asia/Almaty")).strftime("%d.%m.%Y")


def _renderer_owned_requisites_section(heading: str) -> bool:
    """The DOCX renderer owns the single final requisites/signature block."""
    return bool(_REQUISITES_SECTION_RE.fullmatch(" ".join(str(heading or "").split())))


def _clean_combined_heading(heading: str) -> str:
    value = " ".join(str(heading or "").split()).strip()
    value = re.sub(
        r"(?i)\s+(?:и|,|/)\s*(?:адрес\w*\s+и\s+)?(?:реквизит\w*|подпис\w*).*$",
        "",
        value,
    ).strip(" ,;:-")
    return value or "Заключительные положения"


def _looks_like_requisites_clause(text: str) -> bool:
    value = " ".join(str(text or "").split()).strip()
    return bool(value and _REQUISITES_CLAUSE_RE.search(value))


def _clean_requisites(values: list[str], fallback: list[str]) -> list[str]:
    result: list[str] = []
    for raw in values or fallback:
        value = " ".join(str(raw or "").split()).strip()
        if not value or _SIGNATURE_LINE_RE.fullmatch(value) or _PARTY_LABEL_RE.fullmatch(value):
            continue
        key = re.sub(r"\W+", "", value.casefold())
        if key and not any(re.sub(r"\W+", "", item.casefold()) == key for item in result):
            result.append(value)
    return result


def _needs_specification(draft: ContractDraft) -> bool:
    body = "\n".join(draft.body_lines())
    if not _SPEC_REFERENCE_RE.search(body):
        return False
    return not any(_SPEC_HEADING_RE.fullmatch(" ".join(str(section.heading or "").split())) for section in draft.sections)


def _body_blocks(draft: ContractDraft, preamble: list[str]) -> list[Block]:
    """Describe the contract body; numbering is left entirely to the renderer."""
    blocks: list[Block] = [Prose(item) for item in preamble]

    for section in draft.sections:
        heading = " ".join(str(section.heading or "").split()).strip()
        if _renderer_owned_requisites_section(heading):
            continue

        combined_requisites = bool(_REQUISITES_HEADING_TOKEN_RE.search(heading))
        rendered_heading = _clean_combined_heading(heading) if combined_requisites else heading
        blocks.append(NumberedItem(rendered_heading, level=0, bold=True))
        for clause in section.clauses:
            if combined_requisites and _looks_like_requisites_clause(clause.text):
                continue
            blocks.append(NumberedItem(clause.text, level=1))
            blocks.extend(
                NumberedItem(sub, level=2)
                for sub in clause.subclauses
                if not (combined_requisites and _looks_like_requisites_clause(sub))
            )

    if not draft.sections:
        blocks.append(
            Heading("[ТРЕБУЕТ УТОЧНЕНИЯ: существенные и иные условия договора]")
        )
    return blocks


def _append_specification(doc: Document) -> None:
    doc.add_paragraph()
    app = doc.add_paragraph()
    app.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    app.add_run("Приложение № 1 к Договору").italic = True

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run("СПЕЦИФИКАЦИЯ").bold = True

    table = doc.add_table(rows=2, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ("Наименование", "Характеристики", "Количество", "Ед. изм.", "Цена", "Сумма")
    for idx, value in enumerate(headers):
        table.rows[0].cells[idx].paragraphs[0].add_run(value).bold = True
    placeholders = (
        "[ТРЕБУЕТ УТОЧНЕНИЯ]",
        "[ТРЕБУЕТ УТОЧНЕНИЯ]",
        "[ТРЕБУЕТ УТОЧНЕНИЯ]",
        "[ТРЕБУЕТ УТОЧНЕНИЯ]",
        "[ТРЕБУЕТ УТОЧНЕНИЯ]",
        "[ТРЕБУЕТ УТОЧНЕНИЯ]",
    )
    for idx, value in enumerate(placeholders):
        table.rows[1].cells[idx].text = value

    note = doc.add_paragraph(
        "Спецификация является неотъемлемой частью Договора и подлежит заполнению Сторонами до подписания, "
        "если соответствующие сведения не содержатся в материалах дела."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def build_contract_docx(draft: ContractDraft) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    for name in ("Title", "Heading 1", "Heading 2"):
        if name in styles:
            styles[name].font.name = "Times New Roman"

    preamble = ensure_identified_preamble(
        draft.preamble,
        party_a=draft.party_a,
        party_b=draft.party_b,
    )
    document_status = _status(draft, preamble)

    if document_status != "READY FOR FINAL HUMAN REVIEW":
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(DRAFT_NOTICE)
        run.font.name = "Times New Roman"
        run.font.size = Pt(8)

        qa = doc.add_paragraph()
        qa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        qa_run = qa.add_run(f"KORGAN QA STATUS: {document_status}")
        qa_run.bold = True
        qa_run.font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run((draft.title or draft.contract_type or "ДОГОВОР").upper())
    title_run.bold = True
    title_run.font.size = Pt(14)

    date_line = doc.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_line.add_run(draft.place_and_date or f"[ТРЕБУЕТ УТОЧНЕНИЯ: место заключения], {_today_kz()}")

    render_blocks(doc, _body_blocks(draft, preamble))

    if _needs_specification(draft):
        _append_specification(doc)

    doc.add_paragraph()
    req_heading = doc.add_paragraph()
    req_heading.add_run("РЕКВИЗИТЫ И ПОДПИСИ СТОРОН").bold = True
    req_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    left, right = table.rows[0].cells
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    left_values = _clean_requisites(
        draft.requisites_a,
        draft.party_a or ["[ТРЕБУЕТ УТОЧНЕНИЯ: реквизиты стороны 1]"],
    ) or ["[ТРЕБУЕТ УТОЧНЕНИЯ: реквизиты стороны 1]"]
    right_values = _clean_requisites(
        draft.requisites_b,
        draft.party_b or ["[ТРЕБУЕТ УТОЧНЕНИЯ: реквизиты стороны 2]"],
    ) or ["[ТРЕБУЕТ УТОЧНЕНИЯ: реквизиты стороны 2]"]

    left_p = left.paragraphs[0]
    left_p.add_run("Сторона 1\n").bold = True
    for item in left_values:
        left_p.add_run(f"{item}\n")
    left_p.add_run("\nПодпись: ____________________")

    right_p = right.paragraphs[0]
    right_p.add_run("Сторона 2\n").bold = True
    for item in right_values:
        right_p.add_run(f"{item}\n")
    right_p.add_run("\nПодпись: ____________________")

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()

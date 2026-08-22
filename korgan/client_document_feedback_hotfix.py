"""Targeted production hardening for verified client document feedback.

This module deliberately layers narrow, idempotent corrections over the current
production pipeline. It does not replace payment, request ownership, citation
release, or DOCX safety gates.
"""

from __future__ import annotations

import io
import logging
import re
from typing import Any, Awaitable, Callable

LOGGER = logging.getLogger(__name__)
_INSTALLED = False
_RUNTIME_GUIDANCE_INSTALLED = False

_PROCEDURAL_ACTS = {"ГПК РК", "АППК РК", "УПК РК"}
_PRIVATE_DISPUTE_RE = re.compile(
    r"(?i)(?:договор\w*|обязательств\w*|задолжен\w*|оплат\w*|аванс\w*|неустойк\w*|пен[яию]\b|"
    r"подряд\w*|работ\w*|услуг\w*|поставк\w*|товар\w*|возврат\w*|убыт\w*|взыск\w*|"
    r"шарт\w*|міндеттем\w*|берешек\w*|төлем\w*|жұмыс\w*|қызмет\w*|тауар\w*|өндір\w*)"
)
_AFFIRMATIVE_PENALTY_RE = re.compile(
    r"(?i)(?:"
    r"(?:неустойк\w*|пен[яию]\b|штраф\w*)[^\n.;]{0,100}(?:подлежит\s+взыскан\w*|начислен\w*|составля\w*|обязан\w*\s+уплат\w*|требу\w*\s+уплат\w*)|"
    r"(?:подлежит\s+взыскан\w*|начислен\w*|составля\w*)[^\n.;]{0,100}(?:неустойк\w*|пен[яию]\b|штраф\w*)|"
    r"(?:тұрақсыздық\s+айыб\w*|өсімпұл\w*|айыппұл\w*)[^\n.;]{0,100}(?:өндір\w*|есептел\w*|құрай\w*)"
    r")"
)

_RESEARCH_SUFFIX = (
    "\n\nМАТЕРИАЛЬНО-ПРАВОВОЕ ПОКРЫТИЕ ДОКУМЕНТА:\n"
    "25. Для гражданского, договорного, трудового или иного частноправового спора нельзя ограничиваться ГПК/процессуальными нормами. "
    "Сначала source-bound проверь материально-правовую основу самого требования: применимые положения общей части и, когда факты этого требуют, соответствующую особенную часть кодекса.\n"
    "26. Для каждого самостоятельного способа защиты или денежного требования (основной долг/возврат, неустойка или пеня, убытки, проценты, компенсация и т.п.) найди отдельную действующую VERIFIED-норму, прямо поддерживающую именно этот способ защиты. Процессуальная норма не заменяет материальную.\n"
    "27. Если характер правоотношения указывает на специальный закон, правила, постановление или иной подзаконный нормативный акт, проверь его применимость по официальному источнику и используй только после VERIFIED-подтверждения. Не добавляй специальные акты по памяти и не перечисляй их ради объёма.\n"
    "28. В итоговом документе разделяй материальные и процессуальные основания: ГПК используется для подсудности, формы, судебных расходов и иных процессуальных вопросов, но не как единственное основание договорного/денежного требования. Если материальная норма не подтверждена, оставь соответствующее требование предварительным вместо выдумывания статьи."
)

_CHECKLIST_RU = {
    "claim": (
        "📋 Чтобы получить максимально готовое исковое заявление, по возможности укажите:\n"
        "• истца и ответчика: ФИО/наименование, ИИН/БИН, адреса и контакты;\n"
        "• договор/событие: номер, дата, предмет и ключевые условия;\n"
        "• что произошло и в какие даты;\n"
        "• суммы и расчёт каждого денежного требования, включая неустойку/пеню, если вы её требуете;\n"
        "• что именно просите суд взыскать/признать/обязать сделать;\n"
        "• доказательства и приложения; суд и номер дела — если уже известны."
    ),
    "pretrial": (
        "📋 Чтобы получить максимально готовую претензию, по возможности укажите:\n"
        "• отправителя и адресата с реквизитами;\n"
        "• договор/правоотношение, номер и дату;\n"
        "• нарушение и хронологию;\n"
        "• основной долг/возврат и расчёт неустойки или пени, если она заявляется;\n"
        "• конкретное требование и желаемый срок исполнения;\n"
        "• подтверждающие документы и переписку."
    ),
    "pretrial_response": (
        "📋 Чтобы получить максимально готовый ОТВЕТ НА ПРЕТЕНЗИЮ, по возможности пришлите:\n"
        "• саму претензию, её дату/номер и все требования;\n"
        "• реквизиты отправителя и получателя;\n"
        "• вашу позицию по каждому требованию: признаёте, частично признаёте или оспариваете;\n"
        "• какие факты неверны и почему;\n"
        "• договор, акты, платежи, переписку и другие доказательства;\n"
        "• желаемый итог: отказ, частичное согласие, зачёт или предложение урегулирования — только если это ваша позиция."
    ),
    "contract": (
        "📋 Чтобы получить максимально готовый договор, по возможности укажите:\n"
        "• вид и цель договора; стороны, роли и реквизиты;\n"
        "• предмет договора;\n"
        "• цену, порядок и сроки оплаты;\n"
        "• срок исполнения/действия;\n"
        "• ответственность, неустойку и порядок приёмки — если нужны;\n"
        "• особые условия, расторжение, конфиденциальность и приложения."
    ),
    "response": (
        "📋 Чтобы получить максимально готовый ОТЗЫВ НА ИСК, по возможности пришлите:\n"
        "• сам иск, суд и номер дела;\n"
        "• стороны и требования истца с суммами;\n"
        "• вашу позицию по каждому требованию;\n"
        "• возражения, расчёты и спорные факты;\n"
        "• договоры, акты, платежи, переписку и иные доказательства;\n"
        "• какие процессуальные просьбы нужно заявить, если они вам известны."
    ),
}

_CHECKLIST_KK = {
    "claim": (
        "📋 Талап қою арызын барынша толық дайындау үшін мүмкіндігінше мыналарды көрсетіңіз:\n"
        "• талапкер мен жауапкердің деректері: аты/атауы, ЖСН/БСН, мекенжайы, байланысы;\n"
        "• шарт/оқиға: нөмірі, күні, пәні және негізгі талаптары;\n"
        "• не болғаны және негізгі күндер;\n"
        "• әр ақшалай талаптың сомасы мен есебі, талап етілсе өсімпұл/тұрақсыздық айыбы;\n"
        "• соттан нақты не сұрайсыз;\n"
        "• дәлелдер мен қосымшалар; сот пен іс нөмірі — белгілі болса."
    ),
    "pretrial": (
        "📋 Сотқа дейінгі талапты барынша толық дайындау үшін мүмкіндігінше көрсетіңіз:\n"
        "• жіберуші мен алушының деректері; шарттың нөмірі мен күні;\n"
        "• бұзушылық және оқиғалар реті;\n"
        "• негізгі сома және талап етілсе өсімпұл/тұрақсыздық айыбының есебі;\n"
        "• нақты талап пен орындау мерзімі;\n"
        "• растайтын құжаттар мен хат алмасу."
    ),
    "pretrial_response": (
        "📋 Сотқа дейінгі ТАЛАПҚА ЖАУАПТЫ барынша толық дайындау үшін:\n"
        "• талаптың өзін, күні/нөмірін және барлық талаптарын;\n"
        "• тараптардың деректерін;\n"
        "• әр талап бойынша ұстанымыңызды;\n"
        "• даулы фактілер мен себептерін;\n"
        "• шарт, акт, төлем, хат алмасу және өзге дәлелдерді;\n"
        "• қалаған нәтижені: бас тарту, ішінара келісу, есепке жатқызу не реттеу ұсынысын көрсетіңіз."
    ),
    "contract": (
        "📋 Шартты барынша толық дайындау үшін мүмкіндігінше көрсетіңіз:\n"
        "• шарттың түрі/мақсаты, тараптар, рөлдер және деректемелер;\n"
        "• пәні; баға және төлем тәртібі;\n"
        "• орындау және қолданылу мерзімі;\n"
        "• жауапкершілік, тұрақсыздық айыбы және қабылдау тәртібі — қажет болса;\n"
        "• ерекше талаптар, бұзу тәртібі, құпиялылық және қосымшалар."
    ),
    "response": (
        "📋 Талап қою арызына ПІКІРДІ барынша толық дайындау үшін:\n"
        "• талап арызын, сот пен іс нөмірін;\n"
        "• тараптар мен талапкердің барлық талаптарын/сомаларын;\n"
        "• әр талап бойынша ұстанымыңызды;\n"
        "• қарсылықтар, есептер және даулы фактілерді;\n"
        "• шарттар, актілер, төлемдер, хат алмасу және өзге дәлелдерді жіберіңіз."
    ),
}


def checklist_text(kind: str, language: str = "ru") -> str:
    """Return a document-specific client checklist without making fields mandatory."""
    mapping = _CHECKLIST_KK if language == "kk" else _CHECKLIST_RU
    body = mapping.get(kind, mapping.get("claim", ""))
    tail = (
        "\n\nҚұжаттарды алдымен тіркеп, содан кейін жетіспейтін мәліметтерді бір хабарламада жібере аласыз. Белгісіз деректерді KORGAN ойдан шығармайды."
        if language == "kk"
        else
        "\n\nМожно сначала приложить все файлы, затем одним сообщением отправить недостающие сведения. Неизвестные данные KORGAN не будет придумывать."
    )
    return body + tail


def progress_text(kind: str, language: str = "ru") -> str:
    """Explain a real post-prepayment generation wait to the client."""
    if language == "kk":
        return (
            "⏳ Құжат өңделуде. KORGAN қолданылатын құқықты, талаптарды және дәлелдерді тексеріп, Word-файлды қалыптастырып жатыр. "
            "Күрделі іс бірнеше минут алуы мүмкін."
        )
    return (
        "⏳ Документ в работе. KORGAN проверяет применимое право, требования и доказательства и формирует Word-файл. "
        "По сложному делу это может занять несколько минут."
    )


def _acts(lines: list[str]) -> set[str]:
    from korgan.citation_audit import extract_references

    result: set[str] = set()
    for line in lines or []:
        result.update(ref.act for ref in extract_references(str(line)))
    return result


def material_law_issue(
    legal_basis: list[str],
    verified_claims: list[str],
    *,
    context: str = "",
    require_for_private_dispute: bool = False,
) -> str | None:
    """Detect when procedure is being used as a substitute for substantive law."""
    basis_acts = _acts(legal_basis)
    verified_acts = _acts(verified_claims)
    material_basis = basis_acts - _PROCEDURAL_ACTS
    verified_material = verified_acts - _PROCEDURAL_ACTS

    if verified_material and not material_basis:
        return (
            "Подтвержденные материально-правовые нормы не перенесены в документ: процессуальные нормы не могут заменять основание самого требования."
        )
    if basis_acts and not material_basis and basis_acts.issubset(_PROCEDURAL_ACTS):
        return "Правовое обоснование содержит только процессуальные нормы; требуется VERIFIED материально-правовая опора требования."
    if require_for_private_dispute and _PRIVATE_DISPUTE_RE.search(context or "") and not material_basis and not verified_material:
        return (
            "Для частноправового/договорного требования source-bound исследование не подтвердило материально-правовую норму; документ нельзя считать полностью готовым только на основании ГПК."
        )
    return None


def _install_research_prompt_patch() -> None:
    from korgan import fast_professional_litigation as litigation

    current = litigation._professional_research_prompt
    if getattr(current, "_korgan_client_material_law", False):
        return

    def material_prompt(case_context: str, *, max_chars: int, checked_on: str, **kwargs: object) -> str:
        base = current(case_context, max_chars=max_chars, checked_on=checked_on, **kwargs)
        return base + _RESEARCH_SUFFIX

    material_prompt._korgan_client_material_law = True  # type: ignore[attr-defined]
    litigation._professional_research_prompt = material_prompt


def _install_pretrial_quality_patches() -> None:
    from korgan import pretrial, pretrial_response

    original_pretrial = pretrial.pretrial_quality_issues
    if not getattr(original_pretrial, "_korgan_material_law", False):
        def pretrial_issues(draft: Any, research: Any) -> list[str]:
            issues = list(original_pretrial(draft, research))
            issue = material_law_issue(
                list(draft.legal_basis),
                list(research.verified_claims),
                context="\n".join(draft.body_lines()),
                require_for_private_dispute=True,
            )
            if issue and issue not in issues:
                issues.append(issue)
            return issues

        pretrial_issues._korgan_material_law = True  # type: ignore[attr-defined]
        pretrial.pretrial_quality_issues = pretrial_issues

    original_response = pretrial_response.pretrial_response_quality_issues
    if not getattr(original_response, "_korgan_material_law", False):
        def response_issues(draft: Any, research: Any) -> list[str]:
            issues = list(original_response(draft, research))
            issue = material_law_issue(
                list(draft.legal_basis),
                list(research.verified_claims),
                context="\n".join(draft.body_lines()),
                require_for_private_dispute=True,
            )
            if issue and issue not in issues:
                issues.append(issue)
            return issues

        response_issues._korgan_material_law = True  # type: ignore[attr-defined]
        pretrial_response.pretrial_response_quality_issues = response_issues


def _install_claim_consistency_patch() -> None:
    from korgan import claim_consistency_guard as guard

    current = guard.claim_consistency_errors
    if getattr(current, "_korgan_client_feedback", False):
        return

    def client_consistency(case_context: str, draft: Any) -> list[str]:
        errors = list(current(case_context, draft))
        prayer = "\n".join(str(value) for value in list(draft.requests or []))
        body = "\n".join(
            [
                *[str(value) for value in list(draft.facts or [])],
                *[str(value) for value in list(draft.legal_basis or [])],
                str(getattr(draft, "late_interest", "") or ""),
            ]
        )
        if _AFFIRMATIVE_PENALTY_RE.search(body) and not guard._PENALTY_REQUEST_RE.search(prayer):
            errors.append(
                "Текст иска утверждает, что неустойка/пеня начислена или подлежит взысканию, но отдельного требования о ней в разделе «ПРОШУ СУД» нет. "
                "Нельзя оставлять внутренне противоречивый документ: если клиент действительно заявляет эту меру, добавьте исполнимое требование только при VERIFIED основании и полном расчете; иначе уберите или уточните утверждение в мотивировочной части."
            )

        legal_issue = material_law_issue(
            list(draft.legal_basis or []),
            [],
            context=f"{case_context}\n{body}",
            require_for_private_dispute=True,
        )
        if legal_issue and legal_issue not in errors:
            errors.append(legal_issue)
        return list(dict.fromkeys(errors))

    client_consistency._korgan_client_feedback = True  # type: ignore[attr-defined]
    guard.claim_consistency_errors = client_consistency


def _build_pretrial_response_with_required_title(draft: Any, language: str = "ru") -> bytes:
    """Render the existing response layout while always emitting the required heading."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt
    from korgan import pretrial_response as source

    kk = language == "kk"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    head.add_run(("Алушы:\n" if kk else "Кому:\n")).bold = True
    for value in draft.recipient or [("[Алушы деректері]" if kk else "[Данные адресата]")]:
        head.add_run(str(value) + "\n")
    head.add_run(("Жіберуші:\n" if kk else "От:\n")).bold = True
    for value in draft.sender or [("[Жіберуші деректері]" if kk else "[Данные отправителя]")]:
        head.add_run(str(value) + "\n")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("СОТҚА ДЕЙІНГІ ТАЛАПҚА ЖАУАП" if kk else "ОТВЕТ НА ПРЕТЕНЗИЮ")
    run.bold = True
    run.font.size = Pt(14)

    reference = source._reference_line(draft.reference, kk)
    if reference:
        p = doc.add_paragraph(reference)
        p.paragraph_format.space_after = Pt(8)

    for item in draft.claim_summary:
        source._body_paragraph(doc, item)
    for item in draft.position:
        source._body_paragraph(doc, item)
    for item in draft.objections:
        source._body_paragraph(doc, item)
    for item in draft.legal_basis:
        source._body_paragraph(doc, item)
    for item in draft.response_terms:
        source._body_paragraph(doc, item)

    if draft.attachments:
        p = doc.add_paragraph()
        p.add_run("Қосымшалар:" if kk else "Приложения:").bold = True
        for index, item in enumerate(draft.attachments, 1):
            doc.add_paragraph(f"{index}. {item}")

    doc.add_paragraph()
    doc.add_paragraph(("Күні: " if kk else "Дата: ") + source._today())
    doc.add_paragraph("Қолы: ____________________" if kk else "Подпись: ____________________")
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def _install_response_title_patch() -> None:
    from korgan import pretrial_response

    pretrial_response.build_pretrial_response_docx = _build_pretrial_response_with_required_title


def _guidance_wrapper(original: Callable[..., Awaitable[None]], kind: str):
    async def guided(*args: Any, **kwargs: Any) -> None:
        await original(*args, **kwargs)
        if len(args) < 2:
            return
        message, state = args[0], args[1]
        try:
            data = await state.get_data()
            language = "kk" if str(data.get("language") or "ru") == "kk" else "ru"
            from korgan.ui import main_menu
            await message.answer(checklist_text(kind, language), reply_markup=main_menu(language))
        except Exception:
            LOGGER.exception("CLIENT_CHECKLIST_FAILED kind=%s", kind)

    guided._korgan_client_guidance = True  # type: ignore[attr-defined]
    return guided


def _wrap_guidance(module: Any, name: str, kind: str) -> None:
    original = getattr(module, name)
    if getattr(original, "_korgan_client_guidance", False):
        return
    setattr(module, name, _guidance_wrapper(original, kind))


def _install_prepayment_and_runtime_guidance_patch() -> None:
    from korgan import prepayment_gate

    current_ensure = prepayment_gate.ensure_prepayment
    if not getattr(current_ensure, "_korgan_progress", False):
        async def ensure_with_progress(message: Any, state: Any, *, kind: str) -> bool:
            allowed = await current_ensure(message, state, kind=kind)
            if not allowed:
                return False
            try:
                data = await state.get_data()
                request_id = str(data.get("request_id") or "")
                request_kind = str(data.get("request_kind") or "")
                language = "kk" if str(data.get("language") or "ru") == "kk" else "ru"
                already = (
                    str(data.get("generation_progress_request_id") or "") == request_id
                    and str(data.get("generation_progress_kind") or "") == kind
                )
                if request_id and request_kind == kind and not already:
                    await state.update_data(
                        generation_progress_request_id=request_id,
                        generation_progress_kind=kind,
                    )
                    await message.answer(progress_text(kind, language))
            except Exception:
                LOGGER.exception("GENERATION_PROGRESS_NOTICE_FAILED kind=%s", kind)
            return True

        ensure_with_progress._korgan_progress = True  # type: ignore[attr-defined]
        prepayment_gate.ensure_prepayment = ensure_with_progress

    current_installer = prepayment_gate.install_generation_prepayment_gate
    if getattr(current_installer, "_korgan_client_guidance", False):
        return

    def install_with_guidance() -> None:
        global _RUNTIME_GUIDANCE_INSTALLED
        current_installer()
        if _RUNTIME_GUIDANCE_INSTALLED:
            return
        from korgan import pretrial_response_runtime, pretrial_runtime, universal_claim_runtime, universal_document_runtime

        _wrap_guidance(universal_claim_runtime, "begin_claim_request", "claim")
        _wrap_guidance(pretrial_runtime, "_ask_pretrial", "pretrial")
        _wrap_guidance(pretrial_response_runtime, "_ask_materials", "pretrial_response")
        _wrap_guidance(universal_document_runtime, "_ask_contract", "contract")
        _wrap_guidance(universal_document_runtime, "_ask_response", "response")

        # Runtime modules import renderer/quality functions by name. Refresh those
        # references after the modules are loaded so the hotfix is effective on
        # both direct calls and production router calls.
        from korgan import pretrial, pretrial_response
        pretrial_runtime.pretrial_quality_issues = pretrial.pretrial_quality_issues
        pretrial_response_runtime.pretrial_response_quality_issues = pretrial_response.pretrial_response_quality_issues
        pretrial_response_runtime.build_pretrial_response_docx = pretrial_response.build_pretrial_response_docx
        _RUNTIME_GUIDANCE_INSTALLED = True
        LOGGER.info("Installed KORGAN client document checklist/progress guidance")

    install_with_guidance._korgan_client_guidance = True  # type: ignore[attr-defined]
    prepayment_gate.install_generation_prepayment_gate = install_with_guidance


def install_client_document_feedback_hotfix() -> None:
    """Install narrow fixes without weakening existing payment or release gates."""
    global _INSTALLED
    if _INSTALLED:
        return
    _install_research_prompt_patch()
    _install_pretrial_quality_patches()
    _install_claim_consistency_patch()
    _install_response_title_patch()
    _install_prepayment_and_runtime_guidance_patch()
    _INSTALLED = True
    LOGGER.info("Installed KORGAN verified client document feedback hotfix")

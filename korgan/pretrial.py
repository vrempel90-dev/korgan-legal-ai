"""Professional pre-trial demand generation for RU/KK without field questionnaires."""

from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from korgan.document_release import review_lines
from korgan.legal_calc import parse_all_amounts_kzt
from korgan.legal_types import LegalResearch, VerificationStatus
from korgan.stable_legal_release import StableLegalProductionService, clean_language_labels, sanitize_research_sources

_PRETRIAL_SCHEMA: dict[str, Any] = {
    "type": "object",
    "properties": {
        "title": {"type": "string"},
        "sender": {"type": "array", "items": {"type": "string"}},
        "recipient": {"type": "array", "items": {"type": "string"}},
        "facts": {"type": "array", "items": {"type": "string"}},
        "legal_basis": {"type": "array", "items": {"type": "string"}},
        "calculation": {"type": "array", "items": {"type": "string"}},
        "demands": {"type": "array", "items": {"type": "string"}},
        "deadline": {"type": "string"},
        "consequences": {"type": "array", "items": {"type": "string"}},
        "attachments": {"type": "array", "items": {"type": "string"}},
        "verification_notes": {"type": "array", "items": {"type": "string"}},
    },
    "required": ["title", "sender", "recipient", "facts", "legal_basis", "calculation", "demands", "deadline", "consequences", "attachments", "verification_notes"],
    "additionalProperties": False,
}

_INTENT_RU = re.compile(r"(?i)\b(?:досудебн\w*\s+претензи\w*|претензи\w*)\b")
# Kazakh case endings are part of the same word: талап, талапты, талаптың, талапқа...
# Do not terminate the pattern at bare "талап" with a word boundary.
_INTENT_KK = re.compile(r"(?i)(?:сотқа\s+дейінгі\s+талап\w*|талап\s+хат\w*)")
_ACTION = re.compile(r"(?i)\b(?:подготов\w*|состав\w*|сформир\w*|сдел\w*|напиш\w*|дайында\w*|жаса\w*|әзірле\w*|құрастыр\w*)\b")
# Russian advice normally begins with «как». In Kazakh the interrogative «қалай»
# naturally follows the object: «Сотқа дейінгі талапты қалай дайындауға болады?».
# Treat either shape as advice so a how-to question never triggers a DOCX.
_ADVICE_RU = re.compile(r"(?i)^\s*как\b")
_ADVICE_KK = re.compile(r"(?i)\bқалай\b")

_LANG_VERSION_RE = re.compile(r"(?i)английск\w*\s+верси\w*|англ\.?\s+ст\.|русск\w*\s+редакц\w*|english\s+version|russian\s+version")

# Законный следующий шаг, который претензия вправе назвать: обращение в суд,
# в уполномоченный орган, начисление предусмотренной договором неустойки,
# односторонний отказ от договора. Всё остальное — «мы примем меры» — адресату
# ничего не сообщает и последствий не порождает.
_LAWFUL_CONSEQUENCE_RE = re.compile(
    r"(?i)(?:\bсуд\w*|исков\w*\s+заявлен\w*|уполномоченн\w*\s+орган\w*|"
    r"неустойк\w*|пен[яию]\b|штраф\w*|расторж\w*|отказ\w*\s+от\s+договор\w*|"
    r"судебн\w*\s+расход\w*|госпошлин\w*|"
    r"сот\w*|талап\s+арыз\w*|тұрақсыздық\s+айыб\w*|өсімпұл\w*)"
)


def is_pretrial_request(text: str | None) -> bool:
    value = " ".join((text or "").split())
    if not value or _ADVICE_RU.search(value) or _ADVICE_KK.search(value):
        return False
    return bool((_INTENT_RU.search(value) or _INTENT_KK.search(value)) and _ACTION.search(value))


@dataclass(slots=True)
class PretrialDraft:
    status: VerificationStatus
    title: str
    sender: list[str]
    recipient: list[str]
    facts: list[str]
    legal_basis: list[str]
    demands: list[str]
    deadline: str
    consequences: list[str]
    attachments: list[str]
    verification_notes: list[str] = field(default_factory=list)
    source_urls: list[str] = field(default_factory=list)
    # Расчёт денежного требования. По умолчанию пуст: старые сохранённые
    # черновики, собранные до появления раздела, продолжают открываться, а
    # экспортёр просто не печатает заголовок, для которого нет материала.
    calculation: list[str] = field(default_factory=list)

    def body_lines(self) -> list[str]:
        return [
            self.title,
            *self.sender,
            *self.recipient,
            *self.facts,
            *self.legal_basis,
            *self.calculation,
            *self.demands,
            self.deadline,
            *self.consequences,
            *self.attachments,
        ]


def pretrial_payload(draft: PretrialDraft) -> dict[str, Any]:
    """Черновик в форме payload схемы — для раунда правки качества.

    Собирается рядом со схемой и dataclass намеренно: раньше такие payload
    строились копиями внутри модулей правки, и добавление раздела в схему
    оставляло копии позади. Раунд правки не видел уже собранный расчёт и
    пересобирал его с нуля, из-за чего суммы между проходами расходились.
    """
    return {
        "title": draft.title,
        "sender": list(draft.sender),
        "recipient": list(draft.recipient),
        "facts": list(draft.facts),
        "legal_basis": list(draft.legal_basis),
        "calculation": list(draft.calculation),
        "demands": list(draft.demands),
        "deadline": draft.deadline,
        "consequences": list(draft.consequences),
        "attachments": list(draft.attachments),
        "verification_notes": list(draft.verification_notes),
    }


def _dedupe(lines: list[str]) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    for raw in lines:
        line = clean_language_labels(raw)
        key = re.sub(r"\W+", "", line.lower())
        if line and key not in seen:
            seen.add(key)
            result.append(line)
    return result


def normalize_pretrial(draft: PretrialDraft) -> None:
    draft.legal_basis = _dedupe(draft.legal_basis)
    draft.facts = _dedupe(draft.facts)
    draft.calculation = _dedupe(draft.calculation)
    draft.demands = _dedupe(draft.demands)
    draft.consequences = _dedupe(draft.consequences)
    if _LANG_VERSION_RE.search("\n".join(draft.body_lines())):
        draft.status = VerificationStatus.NEEDS_VERIFICATION
        draft.verification_notes.append("В документе обнаружена некорректная ссылка на языковую версию нормы.")


def has_money_demand(draft: PretrialDraft) -> bool:
    """Есть ли в требованиях денежная сумма, которую нужно раскрыть расчётом."""
    return bool(parse_all_amounts_kzt("\n".join(draft.demands)))


def pretrial_quality_issues(draft: PretrialDraft, research: LegalResearch) -> list[str]:
    issues: list[str] = []
    if not draft.sender:
        issues.append("не указан отправитель претензии")
    if not draft.recipient:
        issues.append("не указан адресат претензии")
    if not draft.facts:
        issues.append("нет фактического основания требований")
    if not draft.demands:
        issues.append("нет сформулированных требований")
    if not draft.legal_basis and research.verified_claims:
        issues.append("VERIFIED нормы не перенесены в правовое обоснование")

    # Денежное требование без расчёта — это сумма, которую адресат не может
    # проверить, а суд позже не сможет соотнести с иском.
    if has_money_demand(draft) and not draft.calculation:
        issues.append("денежное требование не раскрыто расчётом")

    # Срок добровольного исполнения — то, что превращает письмо в претензию:
    # от него считается момент, с которого спор можно передать в суд.
    if not str(draft.deadline or "").strip():
        issues.append("не указан срок добровольного исполнения требований")

    # Последствие неисполнения должно называть законный следующий шаг.
    # «Мы примем меры» ничего не сообщает адресату и не порождает последствий.
    consequences = " ".join(draft.consequences)
    if consequences.strip() and not _LAWFUL_CONSEQUENCE_RE.search(consequences):
        issues.append("последствия неисполнения не названы конкретным законным шагом")

    report = review_lines(draft.body_lines(), verified_claims=research.verified_claims)
    issues.extend(report.blocking)
    return list(dict.fromkeys(issues))


class PretrialProductionService(StableLegalProductionService):
    async def research_pretrial(self, case_context: str, language: str = "ru") -> LegalResearch:
        # Reuse the proven professional source-bound pass. It already decomposes
        # remedies and checks mandatory pre-trial procedure; the dedicated draft
        # below simply omits court-only material from the client document.
        research = await self.research_case(case_context, language=language)
        return sanitize_research_sources(research)

    async def draft_pretrial(self, case_context: str, research: LegalResearch, language: str = "ru") -> PretrialDraft:
        verified = "\n".join(f"- {x}" for x in research.verified_claims) or "- нет подтвержденных норм"
        prompt = (
            "Подготовь профессиональную досудебную претензию по праву Республики Казахстан. "
            "Это не иск и не анкета. Используй только факты пользователя и VERIFIED-нормы.\n\n"
            "ЭТАЛОН ПОДАЧИ И СТРУКТУРЫ:\n"
            "Документ должен выглядеть как реальная деловая досудебная претензия, подготовленная практикующим юристом, а не как AI-справка или юридическое заключение. "
            "Ориентируйся на такую последовательность: реквизиты сторон → название документа → договор/правоотношение и роли сторон → хронология исполнения и нарушения → конкретные денежные последствия, задолженность, аванс или иные суммы → относящиеся к спору пункты договора и расчёт неустойки/пени, если они подтверждены материалами → применимые VERIFIED-нормы → конкретное требование и срок добровольного исполнения → законное последствие неисполнения → подпись.\n"
            "Пиши связными деловыми абзацами. Не превращай документ в меморандум с искусственными разделами «Фактические обстоятельства», «Правовое обоснование», «Требования», «Последствия». "
            "Переходы должны быть естественными: например «В нарушение условий договора…», «Кроме того…», «В этой связи…», «Руководствуясь…», но используй их только когда они подходят фактам. "
            "Каждый элемент facts, legal_basis, demands и consequences формулируй как готовый абзац документа, а не как внутреннюю заметку, тезис или комментарий модели.\n\n"
            "ПРАВИЛА БЕЗОПАСНОСТИ И ТОЧНОСТИ:\n"
            "1. Не копируй никакие факты, суммы, даты, названия компаний, договоры, проценты, сроки или статьи из примера оформления. Пример задаёт только форму и уровень юридической подачи.\n"
            "2. Не придумывай ФИО/БИН/ИИН, адрес, договор, даты, суммы, доказательства или факт направления прежней претензии.\n"
            "3. Каждое требование должно вытекать из факта и иметь правовое основание, если оно VERIFIED.\n"
            "4. Если пользователь дал пункт договора о неустойке/пене и данные для расчёта, изложи договорное основание и понятный расчёт. Если данных нет — не рассчитывай и не придумывай.\n"
            "4a. calculation — отдельный раздел «Расчёт задолженности». Он ОБЯЗАТЕЛЕН, если хотя бы одно требование денежное. "
            "Каждую составляющую покажи отдельной строкой и раскрой до элементов: основание, база, ставка, период, количество дней, формула, итог. "
            "Для договорной неустойки используй именно пункт договора и согласованную сторонами ставку; не подменяй её расчётом по статье 353 ГК РК. "
            "Расходы на юридические услуги и иные расходы показывай отдельными строками. "
            "Если не хватает базы, ставки или периода — строку не считай, а укажи недостающий элемент в verification_notes.\n"
            "5. Не пиши 'английская версия', 'русская редакция' и не представляй переводы одного акта как разные нормы.\n"
            "6. Одну статью не пересказывай несколько раз: один точный абзац на одну норму.\n"
            "7. Срок добровольного исполнения указывай только если он дан пользователем или VERIFIED законом/договором; иначе сформулируй нейтрально без выдуманного числа дней.\n"
            "8. В последствиях укажи возможное обращение в суд/уполномоченный орган только как следующий законный шаг, без угроз и без гарантии результата.\n"
            "9. В приложениях перечисляй только реально имеющиеся материалы.\n"
            "10. Не пиши мета-фразы «правовая оценка не проведена», «позиция не определена», «необходимо дополнительно проанализировать» в тело претензии. Недостающие критичные сведения вынеси только в verification_notes.\n"
            f"11. Язык документа: {'казахский' if language == 'kk' else 'русский'}.\n\n"
            f"МАТЕРИАЛЫ:\n{case_context[:self.settings.max_case_text_chars]}\n\n"
            f"VERIFIED:\n{verified}"
        )
        payload, _ = await self._structured_response(
            model=self.settings.openai_model,
            instructions=(
                "Ты практикующий юрист KORGAN в Республике Казахстан. Составляй полноценную деловую досудебную претензию уровня юридической практики: связное письмо, конкретные факты, договорные основания, расчёты и требования только там, где они подтверждены материалами. "
                "Не добавляй непроверенное право, не копируй факты из образцов оформления и не задавай пользователю анкету."
            ),
            content=[{"role": "user", "content": [{"type": "input_text", "text": prompt}]}],
            schema_name="korgan_pretrial_demand",
            schema=_PRETRIAL_SCHEMA,
        )
        draft = PretrialDraft(
            status=research.status,
            source_urls=list(research.source_urls),
            **payload,
        )
        normalize_pretrial(draft)
        issues = pretrial_quality_issues(draft, research)
        if issues:
            draft.status = VerificationStatus.NEEDS_VERIFICATION
            draft.verification_notes.extend(x for x in issues if x not in draft.verification_notes)
        return draft


def _today() -> str:
    return datetime.now(ZoneInfo("Asia/Almaty")).strftime("%d.%m.%Y")


def _body_paragraph(doc: Document, text: str) -> None:
    value = str(text or "").strip()
    if not value:
        return
    p = doc.add_paragraph(value)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(6)


def _deadline_paragraph(value: str, kk: bool) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    lower = text.lower()
    if kk:
        if any(word in lower for word in ("мерзім", "күн ішінде", "орында", "талап ет")):
            return text
        return f"Жоғарыда көрсетілген талаптарды {text} ішінде орындауды сұраймыз."
    if any(word in lower for word in ("срок", "в течение", "просим", "требуем", "исполн")):
        return text
    return f"Просим исполнить указанные требования в течение {text}."


def build_pretrial_docx(draft: PretrialDraft, language: str = "ru") -> bytes:
    """Render a business-letter pre-trial demand using the approved reference structure.

    The renderer intentionally avoids memo-like body headings. Facts, law,
    demands and consequences are already drafted as complete legal paragraphs.
    """
    kk = language == "kk"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    head.add_run(("Алушы:\n" if kk else "Кому:\n")).bold = True
    for value in draft.recipient or [("[Алушы деректері]" if kk else "[Данные адресата]")]:
        head.add_run(str(value) + "\n")
    head.add_run(("Жіберуші:\n" if kk else "От:\n")).bold = True
    for value in draft.sender or [("[Жіберуші деректері]" if kk else "[Данные отправителя]")]:
        head.add_run(str(value) + "\n")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(draft.title or ("Сотқа дейінгі талап" if kk else "Досудебная претензия"))
    run.bold = True
    run.font.size = Pt(14)

    for fact in draft.facts:
        _body_paragraph(doc, fact)

    for basis in draft.legal_basis:
        _body_paragraph(doc, basis)

    # Расчёт печатается отдельным озаглавленным блоком: адресат должен видеть,
    # из чего сложилась сумма, а не искать её в сплошном тексте письма.
    if draft.calculation:
        heading = doc.add_paragraph()
        heading.add_run("Берешек есебі:" if kk else "Расчёт задолженности:").bold = True
        for line in draft.calculation:
            _body_paragraph(doc, line)

    for demand in draft.demands:
        _body_paragraph(doc, demand)

    deadline = _deadline_paragraph(draft.deadline, kk)
    if deadline:
        _body_paragraph(doc, deadline)

    for line in draft.consequences:
        _body_paragraph(doc, line)

    if draft.attachments:
        p = doc.add_paragraph()
        p.add_run("Қосымшалар:" if kk else "Приложения:").bold = True
        for index, item in enumerate(draft.attachments, 1):
            doc.add_paragraph(f"{index}. {item}")

    doc.add_paragraph()
    doc.add_paragraph(("Күні: " if kk else "Дата: ") + _today())
    doc.add_paragraph("Қолы: ____________________" if kk else "Подпись: ____________________")

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()

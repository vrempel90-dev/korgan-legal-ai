from __future__ import annotations

import base64
import io
import json
import logging
from typing import Any
from urllib.parse import urlparse

from docx import Document
from openai import AsyncOpenAI

from korgan.config import Settings
from korgan.legal_types import ClaimDraft, ExtractedDocument, LegalResearch, VerificationStatus
from korgan.pro_claim_sections import extend_claim_schema as _extend_claim_schema

LOGGER = logging.getLogger(__name__)

_EXTRACT_SCHEMA: dict[str, Any] = {
    "type": "object",
    "properties": {
        "document_type": {"type": "string"},
        "text_summary": {"type": "string"},
        "parties": {"type": "array", "items": {"type": "string"}},
        "identifiers": {"type": "array", "items": {"type": "string"}},
        "dates": {"type": "array", "items": {"type": "string"}},
        "amounts": {"type": "array", "items": {"type": "string"}},
        "obligations": {"type": "array", "items": {"type": "string"}},
        "violations": {"type": "array", "items": {"type": "string"}},
        "evidence": {"type": "array", "items": {"type": "string"}},
        "important_facts": {"type": "array", "items": {"type": "string"}},
        "missing_or_unclear": {"type": "array", "items": {"type": "string"}},
    },
    "required": [
        "document_type", "text_summary", "parties", "identifiers", "dates", "amounts",
        "obligations", "violations", "evidence", "important_facts", "missing_or_unclear"
    ],
    "additionalProperties": False,
}

_RESEARCH_SCHEMA: dict[str, Any] = {
    "type": "object",
    "properties": {
        "applicable_law": {"type": "array", "items": {"type": "string"}},
        "procedural_requirements": {"type": "array", "items": {"type": "string"}},
        "verified_claims": {"type": "array", "items": {"type": "string"}},
        "unverified_claims": {"type": "array", "items": {"type": "string"}},
        "source_urls": {"type": "array", "items": {"type": "string"}},
        "notes": {"type": "array", "items": {"type": "string"}},
    },
    "required": ["applicable_law", "procedural_requirements", "verified_claims", "unverified_claims", "source_urls", "notes"],
    "additionalProperties": False,
}

_CLAIM_SCHEMA: dict[str, Any] = {
    "type": "object",
    "properties": {
        "title": {"type": "string"},
        "court": {"type": "string"},
        "claimant": {"type": "array", "items": {"type": "string"}},
        "defendant": {"type": "array", "items": {"type": "string"}},
        "price_of_claim": {"type": "string"},
        "facts": {"type": "array", "items": {"type": "string"}},
        "legal_basis": {"type": "array", "items": {"type": "string"}},
        "requests": {"type": "array", "items": {"type": "string"}},
        "attachments": {"type": "array", "items": {"type": "string"}},
        "verification_notes": {"type": "array", "items": {"type": "string"}},
    },
    "required": ["title", "court", "claimant", "defendant", "price_of_claim", "facts", "legal_basis", "requests", "attachments", "verification_notes"],
    "additionalProperties": False,
}

# Профессиональные разделы: подсудность, расчёт, досудебный порядок, примирение,
# исковая давность, снятие возражений ответчика, ходатайства. Добавляются здесь,
# в единственном определении схемы, поэтому все двенадцать вызывающих модулей
# получают их автоматически.
_extend_claim_schema(_CLAIM_SCHEMA)

_VALIDATION_SCHEMA: dict[str, Any] = {
    "type": "object",
    "properties": {
        "critical_errors": {"type": "array", "items": {"type": "string"}},
        "unsupported_legal_claims": {"type": "array", "items": {"type": "string"}},
        "missing_required_fields": {"type": "array", "items": {"type": "string"}},
    },
    "required": ["critical_errors", "unsupported_legal_claims", "missing_required_fields"],
    "additionalProperties": False,
}


class OpenAILegalService:
    def __init__(self, settings: Settings):
        self.settings = settings
        self.client = AsyncOpenAI(api_key=settings.openai_api_key)

    @staticmethod
    def _json_schema(name: str, schema: dict[str, Any]) -> dict[str, Any]:
        return {"format": {"type": "json_schema", "name": name, "schema": schema, "strict": True}}

    @staticmethod
    def _docx_text(data: bytes) -> str:
        document = Document(io.BytesIO(data))
        chunks: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                chunks.append(text)
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    chunks.append(row_text)
        return "\n".join(chunks)

    def _allowed_source(self, url: str) -> bool:
        try:
            host = (urlparse(url).hostname or "").lower()
        except ValueError:
            return False
        return any(host == domain or host.endswith("." + domain) for domain in self.settings.legal_domains)

    @staticmethod
    def _annotation_urls(response: Any) -> list[str]:
        urls: list[str] = []
        for item in getattr(response, "output", []) or []:
            for content in getattr(item, "content", []) or []:
                for annotation in getattr(content, "annotations", []) or []:
                    url = getattr(annotation, "url", None)
                    if url and url not in urls:
                        urls.append(url)
            action = getattr(item, "action", None)
            for source in getattr(action, "sources", []) or []:
                url = getattr(source, "url", None)
                if url and url not in urls:
                    urls.append(url)
        return urls

    async def _structured_response(
        self,
        *,
        model: str,
        instructions: str,
        content: list[dict[str, Any]] | str,
        schema_name: str,
        schema: dict[str, Any],
        tools: list[dict[str, Any]] | None = None,
    ) -> tuple[dict[str, Any], Any]:
        kwargs: dict[str, Any] = {
            "model": model,
            "instructions": instructions,
            "input": content,
            "text": self._json_schema(schema_name, schema),
            "store": False,
        }
        if tools:
            kwargs["tools"] = tools
        response = await self.client.responses.create(**kwargs)
        payload = json.loads(response.output_text)
        return payload, response

    async def extract_document(self, data: bytes, filename: str, mime_type: str | None = None) -> ExtractedDocument:
        suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
        prompt = (
            "Извлеки из документа только факты, видимые в самом документе. Не делай правовых выводов из памяти. "
            "Особенно найди стороны, ИИН/БИН/номера, даты, суммы, обязательства, нарушения условий, доказательства, "
            "важные факты для возможного судебного спора и явно перечисли всё нечитабельное или отсутствующее."
        )

        if suffix in {"jpg", "jpeg", "png", "webp"} or (mime_type or "").startswith("image/"):
            media = mime_type or ("image/png" if suffix == "png" else "image/jpeg")
            encoded = base64.b64encode(data).decode("ascii")
            content: list[dict[str, Any]] = [{
                "role": "user",
                "content": [
                    {"type": "input_text", "text": prompt},
                    {"type": "input_image", "image_url": f"data:{media};base64,{encoded}", "detail": "high"},
                ],
            }]
        elif suffix == "docx":
            text = self._docx_text(data)
            content = [{"role": "user", "content": [{"type": "input_text", "text": f"{prompt}\n\nТекст DOCX:\n{text[:self.settings.max_case_text_chars]}"}]}]
        elif suffix == "txt":
            text = data.decode("utf-8", errors="replace")
            content = [{"role": "user", "content": [{"type": "input_text", "text": f"{prompt}\n\nТекст файла:\n{text[:self.settings.max_case_text_chars]}"}]}]
        elif suffix == "pdf" or (mime_type or "") == "application/pdf":
            encoded = base64.b64encode(data).decode("ascii")
            content = [{
                "role": "user",
                "content": [
                    {"type": "input_text", "text": prompt},
                    {"type": "input_file", "filename": filename, "file_data": encoded},
                ],
            }]
        else:
            raise ValueError("Поддерживаются PDF, DOCX, TXT, JPG, JPEG, PNG и WEBP.")

        payload, _ = await self._structured_response(
            model=self.settings.openai_vision_model,
            instructions=(
                "Ты модуль извлечения фактов KORGAN Legal AI для Республики Казахстан. "
                "Не придумывай отсутствующие реквизиты. Если поле не видно — оставь список пустым или добавь его в missing_or_unclear."
            ),
            content=content,
            schema_name="korgan_document_extract",
            schema=_EXTRACT_SCHEMA,
        )
        return ExtractedDocument(filename=filename, **payload)

    async def research_case(self, case_context: str, language: str = "ru") -> LegalResearch:
        tools = [{
            "type": "web_search",
            "filters": {"allowed_domains": self.settings.legal_domains},
            "search_context_size": "high",
        }]
        prompt = (
            "Исследуй правовую основу для возможного иска по описанным фактам. Проверяй по официальным источникам. "
            "Отдельно проверь: применимые нормы, требования к форме/содержанию иска, подсудность если её можно установить, "
            "досудебный порядок, срок, госпошлину и приложения. Точные статьи, суммы, сроки и подсудность нельзя утверждать без источника. "
            "В verified_claims помещай только то, что подтверждено найденным официальным источником. Всё остальное — в unverified_claims. "
            "В source_urls перечисли URL использованных официальных страниц.\n\n"
            f"Материалы дела:\n{case_context[:self.settings.max_case_text_chars]}"
        )
        payload, response = await self._structured_response(
            model=self.settings.openai_model,
            instructions=(
                "Ты юридический исследователь KORGAN по праву Республики Казахстан. Работай fail-closed: "
                "не подтверждено официальным источником — значит не считается установленным. "
                f"Язык ответа: {'казахский' if language == 'kk' else 'русский'}."
            ),
            content=[{"role": "user", "content": [{"type": "input_text", "text": prompt}]}],
            schema_name="korgan_legal_research",
            schema=_RESEARCH_SCHEMA,
            tools=tools,
        )
        annotated = [url for url in self._annotation_urls(response) if self._allowed_source(url)]
        claimed = [url for url in payload.get("source_urls", []) if self._allowed_source(url)]
        source_urls = list(dict.fromkeys(annotated + claimed))
        unverified = list(payload.get("unverified_claims", []))
        status = VerificationStatus.VERIFIED if source_urls and not unverified else VerificationStatus.NEEDS_VERIFICATION
        return LegalResearch(
            status=status,
            applicable_law=list(payload.get("applicable_law", [])),
            procedural_requirements=list(payload.get("procedural_requirements", [])),
            verified_claims=list(payload.get("verified_claims", [])),
            unverified_claims=unverified,
            source_urls=source_urls,
            notes=list(payload.get("notes", [])),
        )

    async def consult(self, question: str, case_context: str = "", language: str = "ru") -> tuple[str, list[str]]:
        tools = [{
            "type": "web_search",
            "filters": {"allowed_domains": self.settings.legal_domains},
            "search_context_size": "medium",
        }]
        prompt = (
            f"Вопрос пользователя:\n{question}\n\n"
            f"Контекст загруженных документов:\n{case_context[:30000] if case_context else 'нет'}\n\n"
            "Дай практичный ответ. Точные статьи, сроки, госпошлину и подсудность называй только если они подтверждены официальным источником. "
            "Если подтверждения недостаточно, прямо напиши NEEDS_VERIFICATION и что нужно уточнить."
        )
        response = await self.client.responses.create(
            model=self.settings.openai_model,
            instructions=(
                "Ты KORGAN Legal AI, юридический ассистент только по праву Республики Казахстан. "
                f"Отвечай на {'казахском' if language == 'kk' else 'русском'}. Не выдумывай право."
            ),
            input=prompt,
            tools=tools,
            store=False,
        )
        urls = [url for url in self._annotation_urls(response) if self._allowed_source(url)]
        return response.output_text.strip(), urls

    async def draft_claim(self, case_context: str, research: LegalResearch, language: str = "ru") -> ClaimDraft:
        verified_block = "\n".join(f"- {x}" for x in research.verified_claims) or "нет подтвержденных выводов"
        unverified_block = "\n".join(f"- {x}" for x in research.unverified_claims) or "нет"
        source_block = "\n".join(research.source_urls) or "нет"
        prompt = (
            "Подготовь проект искового заявления для суда Республики Казахстан. Используй факты только из материалов дела. "
            "Юридическое обоснование строй только на VERIFIED-выводах исследования. Не заполняй неизвестные реквизиты догадками. "
            "Вместо неизвестного суда, госпошлины, цены иска, адреса, ИИН/БИН или иной обязательной информации используй "
            "[ТРЕБУЕТ УТОЧНЕНИЯ: ...]. Каждое непроверенное правовое положение внеси в verification_notes.\n\n"
            f"МАТЕРИАЛЫ ДЕЛА:\n{case_context[:self.settings.max_case_text_chars]}\n\n"
            f"ПОДТВЕРЖДЕНО:\n{verified_block}\n\n"
            f"НЕ ПОДТВЕРЖДЕНО:\n{unverified_block}\n\n"
            f"ОФИЦИАЛЬНЫЕ ИСТОЧНИКИ:\n{source_block}"
        )
        payload, _ = await self._structured_response(
            model=self.settings.openai_model,
            instructions=(
                "Ты старший судебный юрист KORGAN по законодательству Республики Казахстан. "
                "Твоя задача — аккуратный проект иска, а не правовая фантазия. "
                f"Язык документа: {'казахский' if language == 'kk' else 'русский'}."
            ),
            content=[{"role": "user", "content": [{"type": "input_text", "text": prompt}]}],
            schema_name="korgan_claim_draft",
            schema=_CLAIM_SCHEMA,
        )
        draft = ClaimDraft(status=research.status, source_urls=research.source_urls, **payload)
        validation = await self.validate_claim(case_context, research, draft)
        problems = validation["critical_errors"] + validation["unsupported_legal_claims"] + validation["missing_required_fields"]
        if problems:
            draft.status = VerificationStatus.NEEDS_VERIFICATION
            draft.verification_notes.extend(problems)
        return draft

    async def validate_claim(self, case_context: str, research: LegalResearch, draft: ClaimDraft) -> dict[str, list[str]]:
        draft_text = json.dumps(
            {
                "title": draft.title,
                "court": draft.court,
                "claimant": draft.claimant,
                "defendant": draft.defendant,
                "price_of_claim": draft.price_of_claim,
                "facts": draft.facts,
                "legal_basis": draft.legal_basis,
                "requests": draft.requests,
                "attachments": draft.attachments,
            },
            ensure_ascii=False,
        )
        prompt = (
            "Проверь проект иска. Найди: (1) факты, которых нет в материалах; (2) правовые утверждения, которых нет в verified_claims; "
            "(3) отсутствующие обязательные поля, которые нельзя безопасно восстановить. Не исправляй сам и не добавляй новые нормы.\n\n"
            f"МАТЕРИАЛЫ:\n{case_context[:40000]}\n\n"
            f"VERIFIED CLAIMS:\n{json.dumps(research.verified_claims, ensure_ascii=False)}\n\n"
            f"ПРОЕКТ:\n{draft_text}"
        )
        payload, _ = await self._structured_response(
            model=self.settings.openai_validation_model,
            instructions="Ты независимый валидатор судебных документов KORGAN. Будь строгим и fail-closed.",
            content=[{"role": "user", "content": [{"type": "input_text", "text": prompt}]}],
            schema_name="korgan_claim_validation",
            schema=_VALIDATION_SCHEMA,
        )
        return {
            "critical_errors": list(payload.get("critical_errors", [])),
            "unsupported_legal_claims": list(payload.get("unsupported_legal_claims", [])),
            "missing_required_fields": list(payload.get("missing_required_fields", [])),
        }
